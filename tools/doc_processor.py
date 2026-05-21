"""
BRIDGE Document Processor v2.6
후보자 이력서/커버레터에서 PII 삭제 + 강사번호 입력

사용법:
  python doc_processor.py setup                              # 폴더 구조 확인
  python doc_processor.py batch [--dry]                      # incoming/ 일괄 처리
  python doc_processor.py process <파일> --number 3057       # 단일 파일 처리
  python doc_processor.py download 3057                      # S3 다운로드+처리
  python doc_processor.py lookup "이름 또는 이메일"            # 후보자 검색
  python doc_processor.py init-db                            # file_uploads 테이블 생성

워크플로우:
  1) incoming/ 폴더에 파일 넣기 (파일명에 강사번호 포함: 3057_resume.pdf)
  2) batch --dry  → 미리보기 확인
  3) batch        → 처리 실행 (processed/ 에 결과, originals/ 에 원본 백업)
  4) 처리 완료 파일은 file_uploads 테이블에 자동 기록 (이력 추적)

지원 형식: .docx (서식 보존), .pdf (인-플레이스 redaction)
Python: "Q:/Phtyon 3/python.exe"
의존성: python-docx, PyMuPDF (fitz), python-dotenv, cryptography, boto3(S3용)
"""

import argparse
import json
import os
import re
import shutil
import sqlite3
import sys
from datetime import datetime
from pathlib import Path

# ── python-docx 부동소수점 XML 속성 패치 ──────────────────────────────────
# 일부 DOCX (Google Docs 내보내기 등)의 XML 속성값이 "863.9999999999999" 형태의
# 소수점 문자열 → python-docx int(str_value) 변환 실패 방지.
# simpletypes 모듈의 모든 클래스를 순회하며 패치.
try:
    import inspect as _inspect
    from docx.oxml import simpletypes as _st_mod

    def _make_robust_classmethod(original_func):
        """int(str_value) 실패 시 round(float(str_value)) 폴백"""
        def _wrapped(cls, str_value: str):
            try:
                return int(str_value)
            except ValueError:
                try:
                    return round(float(str_value))
                except (ValueError, TypeError):
                    return original_func(cls, str_value)
        return classmethod(_wrapped)

    for _name, _cls in _inspect.getmembers(_st_mod, _inspect.isclass):
        if not hasattr(_cls, 'convert_from_xml'):
            continue
        _fn = getattr(_cls, 'convert_from_xml', None)
        if _fn is None:
            continue
        # 직접 정의된 convert_from_xml만 패치 (상속된 것 제외)
        if 'convert_from_xml' in _cls.__dict__:
            _orig = _cls.__dict__['convert_from_xml']
            _orig_fn = _orig.__func__ if hasattr(_orig, '__func__') else _orig
            # 함수 소스에 int(str_value) 패턴이 있는지 확인
            try:
                _src = _inspect.getsource(_orig_fn)
                if 'int(str_value)' in _src and '.' not in _src.split('int(str_value)')[0].split('\n')[-1]:
                    _cls.convert_from_xml = _make_robust_classmethod(_orig_fn)  # type: ignore
            except (OSError, TypeError):
                pass
except Exception:
    pass  # 패치 실패 시 무시 (치명적 아님)

# ── 경로 ──────────────────────────────────────────────
BASE_DIR = Path("Q:/Claudework/bridge base")
DB_PATH = BASE_DIR / "master.db"
ENV_PATH = BASE_DIR / ".env"
DEFAULT_OUTPUT = BASE_DIR / "tools" / "processed_docs" / "processed"
INCOMING_DIR = BASE_DIR / "tools" / "processed_docs" / "incoming"
BACKUP_DIR = BASE_DIR / "tools" / "processed_docs" / "originals"
LOG_DIR = BASE_DIR / "tools" / "processed_docs" / "logs"


# ══════════════════════════════════════════════════════
#  1. DB 접근 + 암호화 처리
# ══════════════════════════════════════════════════════

_db_conn = None  # 세션 내 연결 재사용


def _get_db():
    """DB 연결 (세션 내 재사용)"""
    global _db_conn
    if _db_conn is None:
        if not DB_PATH.exists():
            print(f"[ERROR] DB not found: {DB_PATH}")
            sys.exit(1)
        _db_conn = sqlite3.connect(str(DB_PATH))
        _db_conn.row_factory = sqlite3.Row
    return _db_conn


def _try_decrypt(value):
    """암호화 필드 복호화 시도. 실패 시 원문 반환."""
    if not value or not value.strip():
        return value or ""
    try:
        # security_vault가 로드 가능할 때만 복호화 시도
        sys.path.insert(0, str(BASE_DIR))
        from dotenv import load_dotenv
        load_dotenv(str(ENV_PATH))
        from security_vault import decrypt_field
        return decrypt_field(value)
    except (ImportError, EnvironmentError):
        return value
    except (ValueError, Exception):
        # base64 디코딩 실패 등 → 평문으로 간주
        return value


def lookup_candidate(query: str):
    """이름/이메일로 후보자 검색"""
    db = _get_db()
    if "@" in query:
        rows = db.execute(
            "SELECT sheet_number, full_name, nationality, email "
            "FROM candidates WHERE email LIKE ? LIMIT 10",
            (f"%{query}%",),
        ).fetchall()
    else:
        rows = db.execute(
            "SELECT sheet_number, full_name, nationality, email "
            "FROM candidates WHERE full_name LIKE ? LIMIT 10",
            (f"%{query}%",),
        ).fetchall()
    return [(r[0], _try_decrypt(r[1]), _try_decrypt(r[2]), _try_decrypt(r[3])) for r in rows]


def get_candidate(number: int):
    """sheet_number로 후보자 조회 → dict or None"""
    db = _get_db()
    row = db.execute(
        "SELECT sheet_number, full_name, nationality, email, "
        "mobile_phone, kakaotalk, current_location, dob, gender, "
        "candidate_id, photo_url "
        "FROM candidates WHERE sheet_number = ?",
        (number,),
    ).fetchone()
    if not row:
        return None
    return {
        "sheet_number": row[0],
        "full_name": _try_decrypt(row[1]),
        "nationality": _try_decrypt(row[2]),
        "email": _try_decrypt(row[3]),
        "mobile_phone": _try_decrypt(row[4]),
        "kakaotalk": _try_decrypt(row[5]),
        "current_location": _try_decrypt(row[6]),
        "dob": _try_decrypt(row[7]) if row[7] else "",
        "gender": _try_decrypt(row[8]) if row[8] else "",
        "candidate_id": row[9] or "",
        "photo_url": row[10] or "",
    }


# 국적 영→한 매핑
_NAT_KR = {
    "american": "미국", "us": "미국", "usa": "미국", "united states": "미국",
    "british": "영국", "uk": "영국", "united kingdom": "영국", "english": "영국",
    "irish": "아일랜드", "ireland": "아일랜드", "n. ireland": "아일랜드",
    "northern irish": "아일랜드",
    "canadian": "캐나다", "canada": "캐나다",
    "australian": "호주", "australia": "호주",
    "new zealander": "뉴질랜드", "new zealand": "뉴질랜드", "kiwi": "뉴질랜드",
    "south african": "남아공", "south africa": "남아공",
    "filipino": "필리핀", "philippines": "필리핀",
    "indian": "인도", "india": "인도",
}

_GENDER_KR = {
    "male": "남성", "m": "남성", "남": "남성", "남성": "남성",
    "female": "여성", "f": "여성", "여": "여성", "여성": "여성",
}


def _extract_nationality_from_text(text: str) -> str:
    """
    이력서 텍스트에서 국적 추출.
    "Nationality: South African" → "south african"
    Returns 소문자 국적 문자열 or ""
    """
    # "Nationality: XXX" 패턴
    pat = re.compile(r"\bnationality\s*:?\s*([A-Za-z][A-Za-z\s]{2,30})", re.I)
    m = pat.search(text)
    if m:
        val = m.group(1).strip().lower()
        # 짧은 불필요 단어 제거
        for noise in ("unknown", "n/a", "na", "none", "other"):
            if val == noise:
                return ""
        return val
    # "Citizenship: XXX" 패턴 fallback
    pat2 = re.compile(r"\bcitizenship\s*:?\s*([A-Za-z][A-Za-z\s]{2,30})", re.I)
    m2 = pat2.search(text)
    if m2:
        return m2.group(1).strip().lower()
    return ""


def _update_candidate_nationality(sheet_number: int, nationality_raw: str):
    """DB의 nationality가 비어있거나 '미상'이면 추출된 값으로 업데이트"""
    if not nationality_raw or not sheet_number:
        return
    try:
        db = _get_db()
        row = db.execute(
            "SELECT nationality FROM candidates WHERE sheet_number = ?", (sheet_number,)
        ).fetchone()
        if not row:
            return
        current = (_try_decrypt(row[0]) or "").strip()
        # 비어있거나 "미상"인 경우에만 업데이트
        if current in ("", "미상", "unknown", "Unknown"):
            db.execute(
                "UPDATE candidates SET nationality = ? WHERE sheet_number = ?",
                (nationality_raw, sheet_number),
            )
            db.commit()
    except Exception:
        pass  # DB 업데이트 실패는 무시 (처리 계속)


def _is_cover_letter(filepath: Path) -> bool:
    """파일명으로 커버레터 단독 파일 여부 판별.
    결합 파일(Resume + Cover Letter)은 False (이력서 처리 우선)"""
    s = filepath.stem.lower()
    has_cl = any(kw in s for kw in (
        "cover letter", "coverletter", "cover_letter", "cl_", "cl ",
        "cl-", "커버레터", "커버", "자기소개",
    ))
    if not has_cl:
        return False
    # "resume" / "cv" / "이력서" 키워드가 같이 있으면 결합 파일 → 이력서로 처리
    has_resume = any(kw in s for kw in (
        "resume", "_cv", "cv_", "cv-", " cv", "이력서", "履歷",
    ))
    return not has_resume


# ── OCR 기반 스캔 PDF redaction (이미지 전용 PDF 지원) ──
_OCR_READER = None  # lazy-init
_FACE_CASCADE = None  # cv2 얼굴 detect용 (스캔 PDF의 원본 사진 영역 탐색)


def _get_face_cascade():
    """OpenCV Haar cascade 얼굴 인식기 — 스캔된 CV의 원본 사진 영역 탐색용"""
    global _FACE_CASCADE
    if _FACE_CASCADE is None:
        try:
            import cv2
            xml_path = cv2.data.haarcascades + "haarcascade_frontalface_default.xml"
            _FACE_CASCADE = cv2.CascadeClassifier(xml_path)
            if _FACE_CASCADE.empty():
                _FACE_CASCADE = False
        except Exception:
            _FACE_CASCADE = False
    return _FACE_CASCADE if _FACE_CASCADE else None


def _detect_and_cover_existing_photo(page, page_idx, all_logs, dpi=150):
    """스캔 페이지에서 얼굴(원본 사진) 영역 탐색 → PDF 좌표로 흰박스 + 영역 반환
    Returns: list of fitz.Rect (PDF 좌표) covered, or [] if none"""
    import fitz
    cascade = _get_face_cascade()
    if cascade is None:
        return []
    try:
        import cv2
        import numpy as np
    except ImportError:
        return []

    pix = page.get_pixmap(dpi=dpi)
    img_arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        gray = cv2.cvtColor(img_arr, cv2.COLOR_RGBA2GRAY)
    elif pix.n == 3:
        gray = cv2.cvtColor(img_arr, cv2.COLOR_RGB2GRAY)
    else:
        gray = img_arr.squeeze()

    faces = cascade.detectMultiScale(
        gray, scaleFactor=1.1, minNeighbors=5,
        minSize=(60, 60), maxSize=(int(pix.width * 0.4), int(pix.height * 0.4))
    )
    if len(faces) == 0:
        return []

    scale = 72.0 / dpi
    covered = []
    for (x, y, w, h) in faces:
        # 얼굴 주변 여백(원형 사진 + 어깨 부분) 까지 확장
        pad_x = int(w * 1.0)
        pad_y_top = int(h * 0.8)
        pad_y_bot = int(h * 1.2)
        x0 = max(0, x - pad_x)
        y0 = max(0, y - pad_y_top)
        x1 = min(pix.width,  x + w + pad_x)
        y1 = min(pix.height, y + h + pad_y_bot)
        rect = fitz.Rect(x0 * scale, y0 * scale, x1 * scale, y1 * scale)
        page.draw_rect(rect, color=None, fill=(1, 1, 1), overlay=True)
        covered.append(rect)
        all_logs.append(f"[page{page_idx}] FACE_COVER: {x},{y} {w}x{h} → PDF rect=({rect.x0:.0f},{rect.y0:.0f})-({rect.x1:.0f},{rect.y1:.0f})")
    return covered


def _is_scanned_page(page) -> bool:
    """페이지가 스캔 이미지 기반인지 (텍스트 없음 + 이미지 있음) 판별"""
    if page.get_text().strip():
        return False
    return len(page.get_images()) > 0


def _get_ocr_reader():
    """EasyOCR Reader 싱글톤 — 첫 호출 시 ~250MB 모델 로드"""
    global _OCR_READER
    if _OCR_READER is None:
        try:
            import easyocr
            _OCR_READER = easyocr.Reader(["en"], gpu=False, verbose=False)
        except ImportError:
            return None
    return _OCR_READER


def _ocr_redact_scanned_page(page, page_idx, candidate, pii_patterns, all_logs, dpi=200):
    """OCR로 스캔 페이지의 PII 영역을 white-redact.
    Returns: True if any redactions applied"""
    import fitz

    reader = _get_ocr_reader()
    if reader is None:
        all_logs.append(f"[page{page_idx}] OCR_SKIP: easyocr 미설치")
        return False

    pix = page.get_pixmap(dpi=dpi)
    img_bytes = pix.tobytes("png")
    scale = 72.0 / dpi
    try:
        ocr_results = reader.readtext(img_bytes, detail=1)
    except Exception as e:
        all_logs.append(f"[page{page_idx}] OCR_FAIL: {e}")
        return False

    cand_name = ""
    cand_parts = []
    if candidate:
        cand_name = (candidate.get("full_name") or candidate.get("name") or "").strip()
        if cand_name:
            cand_parts = [w for w in cand_name.split() if len(w) >= 3]

    # Contact 섹션 박스 전체 흰박스 (아이콘/벡터 그래픽 제거용)
    # 전략: "Contact" 라벨의 y → 다음 섹션 헤더 y 사이의 좌측 컬럼을 통째로 덮기
    _LEFT_SECTION_HDRS = {
        "certifications", "skill", "skills", "language", "languange",
        "languages", "education", "experience",
    }
    _contact_y = None
    _contact_x = None
    for bbox_pts, text, conf in ocr_results:
        if text.strip().lower().rstrip(":") == "contact":
            _contact_y = bbox_pts[0][1]
            _contact_x = bbox_pts[0][0]
            break
    if _contact_y is not None:
        # 같은 컬럼(±100px) 에서 _contact_y 아래의 다음 섹션 헤더
        _next_sec_y = None
        for bbox_pts, text, conf in ocr_results:
            t = text.strip().lower().rstrip(":")
            if t not in _LEFT_SECTION_HDRS:
                continue
            by = bbox_pts[0][1]
            bx = bbox_pts[0][0]
            if by <= _contact_y or abs(bx - _contact_x) > 150:
                continue
            if _next_sec_y is None or by < _next_sec_y:
                _next_sec_y = by
        if _next_sec_y is not None:
            import fitz as _fitz
            # 좌측 컬럼 폭: contact_x 기준 좌측 마진~컬럼 중앙
            col_left = 0
            col_right = int((_contact_x + 280 / scale) * scale)  # PDF ~280pt
            col_right = min(col_right, int(595 * 0.40))  # 페이지 좌측 40% 이내
            top = int((_contact_y - 10) * scale)
            bot = int((_next_sec_y - 10) * scale)
            box = _fitz.Rect(col_left, top, col_right, bot)
            page.draw_rect(box, color=None, fill=(1, 1, 1), overlay=True)
            all_logs.append(
                f"[page{page_idx}] CONTACT_BOX_WIPE: y={top}-{bot} x={col_left}-{col_right}"
            )

    # Teaching Experience 섹션 시작 y 위치 탐색 (한국 학원명 detect용)
    _EXP_HEADERS = {
        "teaching experience", "work experience", "career", "experience",
        "employment", "employment history", "professional experience",
    }
    # 단어 경계 매칭 (substring 오탐 방지 — "esl"이 "ESL students" 본문에 매칭되는 문제)
    _JOB_TITLE_RE = re.compile(
        r'\b(?:teacher|tutor|instructor|lecturer|professor|trainer|head\s+of)\b',
        re.I,
    )
    _exp_y = None
    for bbox_pts, text, conf in ocr_results:
        if text.lower().strip().rstrip(":") in _EXP_HEADERS:
            _exp_y = bbox_pts[0][1]
            break

    # Y 기준 정렬 (직책 다음 줄 = 학원명 탐색용)
    _sorted_ocr = sorted(ocr_results, key=lambda r: r[0][0][1])
    _workplace_keys = set()  # (y, text) keys to mark as workplaces
    if _exp_y is not None:
        _SENTENCE_WORDS = {
            "the", "and", "with", "from", "for", "to", "is", "was",
            "are", "were", "have", "has", "had", "of", "in", "on",
            "at", "by", "as", "an", "led", "taught", "worked",
            "subjects", "instructed", "educated", "implementing",
        }
        for i, (bbox_i, txt_i, conf_i) in enumerate(_sorted_ocr):
            y_i = bbox_i[0][1]
            if y_i < _exp_y:
                continue
            if not _JOB_TITLE_RE.search(txt_i):
                continue
            job_x = bbox_i[0][0]
            # 다음 줄(같은 컬럼) 탐색
            for j in range(i + 1, len(_sorted_ocr)):
                bbox_j, txt_j, conf_j = _sorted_ocr[j]
                y_j = bbox_j[0][1]
                x_j = bbox_j[0][0]
                if y_j - y_i > 100:
                    break
                if abs(x_j - job_x) > 60:
                    continue
                t_strip = txt_j.strip()
                if len(t_strip) > 60:
                    continue
                # 연도/날짜 스킵
                if re.match(r'^\d{4}\s*[-–]', t_strip) or re.match(r'^\d{4}$', t_strip):
                    continue
                # 문장형(소문자 시작, 일반 어휘 포함) 스킵
                first_word = t_strip.split()[0].lower() if t_strip.split() else ""
                if first_word in _SENTENCE_WORDS:
                    continue
                # 학원명 후보로 마킹 (y좌표 + 텍스트로 key 생성)
                _workplace_keys.add((int(y_j), txt_j))
                break

    redacted = 0
    for bbox_pts, text, conf in ocr_results:
        if conf < 0.30 or not text.strip():
            continue
        text_lower = text.lower()
        why = None

        # 한국 학원/일터 (Teaching Experience 아래 직책 다음 줄)
        if (int(bbox_pts[0][1]), text) in _workplace_keys:
            why = f"WORKPLACE: {text}"

        if cand_name and cand_name.lower() in text_lower:
            why = f"NAME: {text}"
        elif cand_parts and any(p.lower() in text_lower for p in cand_parts):
            why = f"NAMEPART: {text}"
        else:
            for pat, label in pii_patterns:
                if pat.search(text):
                    why = f"{label}: {text}"
                    break
            if not why:
                lower_strip = text_lower.strip().rstrip(":")
                _PII_LABELS = {
                    "contact", "address", "phone", "email", "tel",
                    "mobile", "cell", "kakao", "kakaotalk", "wechat",
                    "skype", "line", "linkedin", "instagram",
                }
                if lower_strip in _PII_LABELS:
                    why = f"LABEL: {text}"

            # 한국 도시 + Korea (RE_KR_CITY_COUNTRY 보완: "Seoul, Korea" 단순형)
            if not why:
                _KR_CITIES_LOW = {
                    "seoul", "busan", "incheon", "daegu", "daejeon",
                    "gwangju", "ulsan", "suwon", "gyeonggi", "yongin",
                    "seongnam", "bundang", "ilsan", "ansan", "anyang",
                    "bucheon", "goyang", "guri", "namyangju", "uijeongbu",
                    "pyeongtaek", "siheung", "hwaseong", "gimpo",
                    "cheongju", "jeonju", "pohang", "changwon", "jeju",
                    "gangnam", "sinchon", "hongdae", "itaewon",
                }
                _has_kr = "korea" in text_lower or "korean" in text_lower
                _has_city = any(c in text_lower for c in _KR_CITIES_LOW)
                if _has_kr and _has_city:
                    why = f"KR_ADDR: {text}"
                elif text_lower.strip() in (
                    "seoul, korea", "seoul korea", "south korea",
                    "republic of korea",
                ):
                    why = f"KR_LOC: {text}"

        if why:
            xs = [p[0] for p in bbox_pts]
            ys = [p[1] for p in bbox_pts]
            rect = fitz.Rect(
                min(xs) * scale - 2, min(ys) * scale - 1,
                max(xs) * scale + 2, max(ys) * scale + 1,
            )
            page.add_redact_annot(rect, fill=(1, 1, 1))
            redacted += 1
            all_logs.append(f"[page{page_idx}] OCR_{why[:60]}")

    if redacted:
        page.apply_redactions()
        all_logs.append(f"[page{page_idx}] OCR_APPLIED: {redacted}건")
    return redacted > 0


def _merge_pdfs(cover_pdf: Path, resume_pdf: Path, output_pdf: Path):
    """커버레터 PDF + 이력서 PDF → 단일 PDF 병합 (최대 압축)"""
    import fitz
    merged = fitz.open()
    for src_path in [cover_pdf, resume_pdf]:
        if src_path and src_path.exists():
            src = fitz.open(str(src_path))
            merged.insert_pdf(src)
            src.close()
    merged.save(
        str(output_pdf),
        garbage=4, deflate=True, deflate_images=True,
        deflate_fonts=True, clean=True,
    )
    merged.close()


def _compress_pdf_images(pdf_path: Path, max_dim: int = 1200, jpeg_quality: int = 75) -> None:
    """PDF 내 이미지 다운샘플링 + JPEG 재압축 (용량 감소).
    스캔 PDF나 고해상도 사진 포함 PDF에 효과적.
    max_dim: 긴 변 최대 픽셀 (1200px = 200dpi A4 적정)
    jpeg_quality: JPEG 품질 (75 = 시각적 손실 거의 없음)
    """
    import fitz, io
    try:
        from PIL import Image
    except ImportError:
        return
    doc = fitz.open(str(pdf_path))
    changed = False
    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.width <= max_dim and pix.height <= max_dim:
                    pix = None
                    continue
                if pix.colorspace and pix.colorspace.name not in ("DeviceRGB", "DeviceGray"):
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                mode = "RGB" if pix.n >= 3 else "L"
                img_bytes = pix.samples
                pil = Image.frombytes(mode, (pix.width, pix.height), img_bytes)
                # 다운샘플
                ratio = max_dim / max(pil.width, pil.height)
                if ratio < 1:
                    new_size = (int(pil.width * ratio), int(pil.height * ratio))
                    pil = pil.resize(new_size, Image.LANCZOS)
                # JPEG 재인코딩
                buf = io.BytesIO()
                if mode == "L":
                    pil.save(buf, format="JPEG", quality=jpeg_quality, optimize=True)
                else:
                    pil.save(buf, format="JPEG", quality=jpeg_quality, optimize=True)
                new_bytes = buf.getvalue()
                # 새 이미지로 교체
                doc.update_stream(xref, new_bytes)
                changed = True
                pix = None
            except Exception:
                pass
    if changed:
        tmp = pdf_path.with_suffix(".compress.tmp.pdf")
        doc.save(
            str(tmp),
            garbage=4, deflate=True, deflate_images=True,
            deflate_fonts=True, clean=True,
        )
        doc.close()
        # 압축 결과가 더 작을 때만 대체
        if tmp.exists() and tmp.stat().st_size < pdf_path.stat().st_size:
            pdf_path.unlink()
            tmp.rename(pdf_path)
        else:
            tmp.unlink(missing_ok=True)
    else:
        doc.close()


def _build_output_filename(number: int, candidate: dict = None, ext: str = ".docx") -> str:
    """
    파일명 단일 규격 (SSoT v2.0):
        {sheet_number}{국적}_{성별}({YY}born){ext}
    예: 5739영국_여성(00born).pdf / 6121아일랜드_남성(00born).pdf

    누락 요소는 폴백 문자열로 대체 (파일명 생성 100% 보장):
        국적 → "국적미상", 성별 → "성별미상", 생년 → "미상"

    호출 지점 (유일 SSoT — 다른 경로에서 파일명 직접 조합 금지):
      1) tools.doc_processor.process_docx / process_pdf
      2) api_server._auto_process_resume (자동 처리)
      3) api_server._fetch_processed_cv_attachment (메일 첨부)
      4) /api/admin/resume/preview/{id}
    """
    nat_kr = "국적미상"
    gender_kr = "성별미상"
    birth_yr = "미상"

    if candidate:
        nat_raw = (candidate.get("nationality") or "").strip().lower()
        if nat_raw:
            mapped = _NAT_KR.get(nat_raw)
            if mapped:
                nat_kr = mapped
            elif re.match(r"^[가-힣]+$", nat_raw):
                # 이미 한글 국적이 들어온 경우 (e.g. "미국")
                nat_kr = nat_raw[:6]

        gen_raw = (candidate.get("gender") or "").strip().lower()
        if gen_raw:
            mapped_g = _GENDER_KR.get(gen_raw)
            if mapped_g:
                gender_kr = mapped_g

        dob = (candidate.get("dob") or "").strip()
        if dob:
            yr_match = re.search(r"(19|20)(\d{2})", dob)
            if yr_match:
                birth_yr = yr_match.group(2)
            elif re.match(r"^\d{2}$", dob):
                birth_yr = dob
            elif re.match(r"^\d{1}$", dob):
                # 한자리 dob (e.g. "3" → "03") zero-pad
                birth_yr = dob.zfill(2)

    # 규격: 항상 3요소 모두 포함 (누락 시 폴백 문자열)
    return f"{number}{nat_kr}_{gender_kr}({birth_yr}born){ext}"


# ── v2.0 전 파일타입 dispatcher ──────────────────────────────────────────
# 파일 확장자 → 처리 카테고리
_EXT_CATEGORY = {
    # PII 제거 대상 (PDF/워드)
    ".pdf":  "document",
    ".docx": "document",
    ".doc":  "document",
    ".rtf":  "document",
    ".odt":  "document",
    ".hwp":  "document",
    # 이미지 (PII 불필요)
    ".jpg":  "image",
    ".jpeg": "image",
    ".png":  "image",
    ".webp": "image",
    ".gif":  "image",
    ".bmp":  "image",
    ".tiff": "image",
    # 영상 (보관만)
    ".mp4":  "video",
    ".mov":  "video",
    ".avi":  "video",
    ".mkv":  "video",
    ".webm": "video",
    ".wmv":  "video",
    # 기타 (해시만)
}


def detect_file_category(filename: str) -> str:
    """파일 확장자로 카테고리 자동 판별.
    Returns: 'document' | 'image' | 'video' | 'other'
    """
    ext = Path(filename).suffix.lower()
    return _EXT_CATEGORY.get(ext, "other")


def process_file(
    filepath: Path,
    brj_number: int,
    candidate: dict = None,
    file_category: str = None,
    dry: bool = False,
    photo_path: Path = None,
):
    """
    v2.0 Universal file processor — 파일 타입에 따라 적절한 처리 라우팅.

    Returns: dict {
        "category": "document"|"image"|"video"|"other",
        "processed": bool,       # PII 제거 완료 여부
        "output_path": Path|None,# 처리 결과 저장 경로 (document만)
        "pii_count": int,        # 검출된 PII 수 (document만)
        "logs": list[str],
        "sha256": str|None,      # 출력 파일 무결성 해시
        "error": str|None,
    }
    """
    import hashlib as _hl

    if file_category is None:
        file_category = detect_file_category(filepath.name)

    result = {
        "category": file_category,
        "processed": False,
        "output_path": None,
        "pii_count": 0,
        "logs": [],
        "sha256": None,
        "error": None,
    }

    try:
        ext = filepath.suffix.lower()

        # ── DOCUMENT: PII 제거 (PDF/DOCX/DOC) ────────────────────────────
        if file_category == "document":
            if ext in (".docx", ".doc"):
                doc, logs = process_docx(
                    filepath, brj_number=brj_number, candidate=candidate,
                    dry=dry, photo_path=photo_path,
                )
                out_name = _build_output_filename(brj_number, candidate, ".docx")
                out_path = filepath.parent / out_name
                if not dry:
                    doc.save(str(out_path))
                    # DOCX → PDF 변환 (Word COM)
                    _pdf_result = _docx_to_pdf(out_path)
                    if _pdf_result and _pdf_result.exists():
                        out_path.unlink(missing_ok=True)
                        out_path = _pdf_result
                    with open(out_path, "rb") as f:
                        result["sha256"] = _hl.sha256(f.read()).hexdigest()
                result["processed"] = True
                result["output_path"] = out_path
                result["logs"] = logs
                result["pii_count"] = sum(
                    1 for l in logs
                    if any(kw in l for kw in ("EMAIL:", "PHONE:", "NAME:", "LOCATION:", "WORKPLACE:"))
                )

            elif ext == ".pdf":
                out_path, logs = process_pdf(
                    filepath, brj_number=brj_number, candidate=candidate,
                    dry=dry, photo_path=photo_path,
                )
                if out_path and Path(out_path).exists():
                    with open(out_path, "rb") as f:
                        result["sha256"] = _hl.sha256(f.read()).hexdigest()
                result["processed"] = True
                result["output_path"] = Path(out_path) if out_path else None
                result["logs"] = logs
                result["pii_count"] = sum(
                    1 for l in logs
                    if any(kw in l for kw in ("EMAIL:", "PHONE:", "NAME:", "LOCATION:", "WORKPLACE:"))
                )

            else:
                # rtf/odt/hwp — PII 처리 모듈 없음 → pass-through
                with open(filepath, "rb") as f:
                    result["sha256"] = _hl.sha256(f.read()).hexdigest()
                result["processed"] = False
                result["output_path"] = filepath
                result["logs"] = [f"[SKIP] {ext} 포맷 PII 처리 미지원 — 원본 보존"]

        # ── IMAGE: 해시 + 크기 검증만 (PII 불필요) ───────────────────────
        elif file_category == "image":
            with open(filepath, "rb") as f:
                data = f.read()
                result["sha256"] = _hl.sha256(data).hexdigest()
            result["processed"] = True   # 처리 완료로 간주 (PII 없음)
            result["output_path"] = filepath
            result["logs"] = [f"[IMAGE] size={len(data)} sha256={result['sha256'][:16]}"]

        # ── VIDEO: 해시 + 메타데이터만 (변환/블러 미지원) ────────────────
        elif file_category == "video":
            with open(filepath, "rb") as f:
                data = f.read()
                result["sha256"] = _hl.sha256(data).hexdigest()
            result["processed"] = True
            result["output_path"] = filepath
            result["logs"] = [f"[VIDEO] size={len(data)} sha256={result['sha256'][:16]}"]

        # ── OTHER: 해시만 ─────────────────────────────────────────────────
        else:
            with open(filepath, "rb") as f:
                data = f.read()
                result["sha256"] = _hl.sha256(data).hexdigest()
            result["processed"] = True
            result["output_path"] = filepath
            result["logs"] = [f"[OTHER] ext={ext} size={len(data)} — 원본 보존"]

    except Exception as e:
        result["error"] = str(e)[:500]
        result["logs"].append(f"[ERROR] {type(e).__name__}: {e}")

    return result


def find_candidate_by_email(email: str):
    """이메일로 후보자 검색 → dict or None (정확 매칭 우선, 없으면 LIKE)"""
    db = _get_db()
    _COLS = ("sheet_number, full_name, nationality, email, "
             "mobile_phone, kakaotalk, current_location, dob, gender")
    # 정확 매칭
    row = db.execute(
        f"SELECT {_COLS} FROM candidates WHERE email = ? LIMIT 1",
        (email,),
    ).fetchone()
    if not row:
        # 암호화된 이메일일 수 있으므로 전체 스캔 (느리지만 정확)
        rows = db.execute(
            f"SELECT {_COLS} FROM candidates WHERE email IS NOT NULL AND email != ''"
        ).fetchall()
        for r in rows:
            decrypted = _try_decrypt(r[3])
            if decrypted and decrypted.lower() == email.lower():
                row = r
                break
    if not row:
        return None
    return {
        "sheet_number": row[0],
        "full_name": _try_decrypt(row[1]),
        "nationality": _try_decrypt(row[2]),
        "email": _try_decrypt(row[3]),
        "mobile_phone": _try_decrypt(row[4]),
        "kakaotalk": _try_decrypt(row[5]),
        "current_location": _try_decrypt(row[6]),
        "dob": _try_decrypt(row[7]) if row[7] else "",
        "gender": _try_decrypt(row[8]) if row[8] else "",
    }


def _extract_emails_from_docx(filepath: Path) -> list[str]:
    """DOCX에서 이메일 주소 추출 (PII 삭제 전 원본에서)"""
    import docx
    doc = docx.Document(str(filepath))
    emails = set()
    all_text = []
    # 본문
    for p in doc.paragraphs:
        all_text.append(p.text)
    # 헤더/푸터
    for sec in doc.sections:
        if sec.header and sec.header.paragraphs:
            for p in sec.header.paragraphs:
                all_text.append(p.text)
        if sec.footer and sec.footer.paragraphs:
            for p in sec.footer.paragraphs:
                all_text.append(p.text)
    # 테이블
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_text.append(p.text)
    full = "\n".join(all_text)
    for m in RE_EMAIL.finditer(full):
        emails.add(m.group().lower())
    return list(emails)


# ══════════════════════════════════════════════════════
#  1b. file_uploads 테이블 (로컬 DB 이력 추적)
# ══════════════════════════════════════════════════════

_FILE_UPLOADS_DDL = """
CREATE TABLE IF NOT EXISTS file_uploads (
    id            INTEGER PRIMARY KEY AUTOINCREMENT,
    entity_type   TEXT    NOT NULL DEFAULT 'candidate',
    entity_id     INTEGER,
    file_type     TEXT,
    file_url      TEXT,
    file_size     INTEGER DEFAULT 0,
    is_deleted    INTEGER NOT NULL DEFAULT 0,
    created_at    TEXT    NOT NULL DEFAULT (datetime('now'))
)
"""


def _ensure_file_uploads_table():
    """file_uploads 테이블이 없으면 생성 (IF NOT EXISTS)."""
    db = _get_db()
    db.execute(_FILE_UPLOADS_DDL)
    db.commit()


def _record_file_upload(entity_id: int, file_type: str, file_url: str, file_size: int = 0):
    """처리 완료 파일을 file_uploads 테이블에 INSERT."""
    _ensure_file_uploads_table()
    db = _get_db()
    db.execute(
        "INSERT INTO file_uploads (entity_type, entity_id, file_type, file_url, file_size) "
        "VALUES (?, ?, ?, ?, ?)",
        ("candidate", entity_id, file_type, file_url, file_size),
    )
    db.commit()


# ══════════════════════════════════════════════════════
#  2. PII 패턴 정의
# ══════════════════════════════════════════════════════

RE_EMAIL = re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b")

# 전화번호: 7자리+ 숫자를 포함하는 시퀀스만 (날짜·연도 충돌 방지)
RE_PHONE = re.compile(
    r"(?:\+?\d{1,3}[\s\-.]?)?\(?\d{2,4}\)?[\s\-.]?\d{3,4}[\s\-.]?\d{3,4}\b"
)

# 서술형 생년월일: "20th of September 1996", "March 5th, 1990", "5 January 1988"
# 생년도 범위: 1940~2009 (2010+ = 문서 날짜 또는 경력 날짜 오탐 방지)
RE_DOB_SPELLED = re.compile(
    r'\b\d{1,2}(?:st|nd|rd|th)?\s+(?:of\s+)?'
    r'(?:January|February|March|April|May|June|July|August|September|October|November|December)'
    r'\s+(?:19[4-9]\d|200[0-9])\b'
    r'|\b(?:January|February|March|April|May|June|July|August|September|October|November|December)'
    r'\s+\d{1,2}(?:st|nd|rd|th)?\s*,?\s*(?:19[4-9]\d|200[0-9])\b',
    re.IGNORECASE,
)

# 한국 전화 프리픽스 잔류 제거: "+82 (0)" 홀로 남는 경우
RE_KR_PHONE_PREFIX = re.compile(r'\+82\s*\(?\s*0\s*\)?(?!\d)', re.IGNORECASE)

# 한국 우편번호 5자리 — 도시명 삭제 후 잔류 방지 (예: "10906", "03000")
# 주의: 연도·번지 오탐 방지 → 줄 시작 또는 쉼표/공백 직후 5자리만
RE_KR_POSTAL = re.compile(r'(?:^|(?<=,\s)|(?<=,))[\s]*\b([0-9]{5})\b(?!\d)', re.MULTILINE)

RE_URL = re.compile(r"https?://[^\s<>\"']+|www\.[^\s<>\"']+", re.I)

RE_LINKEDIN = re.compile(
    r"(?:linkedin\.com/in/[^\s<>\"']+|linkedin\s*:?\s*[^\s,;\n]+)", re.I
)

RE_KAKAO = re.compile(
    r"(?:kakao(?:talk)?|카카오(?:톡)?)\s*:?\s*[A-Za-z0-9._\-]+", re.I
)

RE_SNS = re.compile(
    # (?<!\w): "online" 내 "line", "Ursuline" 내 "line", "anticholinergic" 내 "line" 오탐 방지
    r"(?<!\w)(?:instagram|facebook|twitter|whatsapp|we\s*chat|skype|telegram|signal|viber|line)"
    r"(?:\s+id|\s+no\.?|\s+account|\s+handle)?"
    r"\s*[:\-]?\s*[@]?(?:live:)?[A-Za-z0-9._\-\+]{2,60}",
    re.I,
)

RE_KR_ADDRESS = re.compile(
    r"(?:서울|부산|대구|인천|광주|대전|울산|세종|경기|강원|충북|충남|전북|전남|경북|경남|제주)"
    r"(?:특별시|광역시|도|특별자치시|특별자치도)?"
    r"[\s]?(?:\S+(?:시|군|구|읍|면|동|리|로|길|번지)[\s]?){0,5}",
    re.UNICODE,
)

RE_PASSPORT = re.compile(r"\b[A-Z]{1,2}\d{7,8}\b")

# 한국 아파트/주거 주소 패턴 ("부영 1차 아파트", "래미안 빌라" 등)
RE_KR_RESIDENTIAL = re.compile(
    r"[가-힣]+(?:\s+[가-힣\d]+)*\s+(?:\d+차\s*)?(?:아파트|빌라|오피스텔|주공|맨션|타운하우스|연립)",
    re.UNICODE,
)

# 영문 주소 패턴: "123 Street Name, City, ST 12345"
# 전체 단어 + 약어(St./Ave./Dr. 등) 모두 매칭
RE_EN_ADDRESS = re.compile(
    r"\d{1,5}\s+[\w\s.]{3,40}\b(?:Street|Avenue|Boulevard|Drive|"
    r"Road|Lane|Court|Place|Trail|Circle|Parkway|Highway|"
    r"St\.?|Ave\.?|Blvd\.?|Dr\.?|Rd\.?|Ln\.?|Ct\.?|Pl\.?|Trl\.?|Cir\.?|Pkwy\.?|Hwy\.?)\b"
    r"[.,]?\s*(?:(?:Apt|Suite|Unit|#)\s*\w+[.,]?\s*)?"
    r"(?:\w[\w\s]*,\s*[A-Z]{2}\s*\d{5}(?:-\d{4})?)?",
    re.I,
)

# 한국 로마자 주소 패턴: "80-11, Jangji 1-gil," / "123 Teheran-ro" / "45 Sejong-daero"
RE_KR_ROMANIZED_ADDR = re.compile(
    r'\d+(?:[-–]\d+)?\s*,?\s*\w+(?:\s+\w+)?\s*\d*\s*[-–]?\s*'
    r'(?:gil|ro|daero|dong|gu|si|ri|myeon|eup|gun|do)\b'
    r'(?:\s*,\s*[\w\s]+)?',
    re.I,
)

# ── 학교/기관명 일반화 (전 세계 공통) ──────────────────────────────────────
# "Sheffield University", "Lincoln High School" → 기관 유형만 남김
RE_SCHOOL_NAMED = re.compile(
    # ★ "at" 제거 — "Studied at University" 오매칭 방지
    r'\b[A-Z][a-zA-Z\'\-\.]+(?:\s+(?:of|the|and|&|[A-Z][a-zA-Z\'\-\.]+)){0,5}\s+'
    r'(?:University|College|High\s+School|Secondary\s+School|'
    r'Elementary\s+School|Primary\s+School|Grammar\s+School|'
    r'Boarding\s+School|Prep\s+School|Junior\s+School|'
    r'Sixth\s+Form\s+College|Collegiate|Polytechnic)\b'
)
# "University of Sheffield", "Institute of Technology" → 기관 유형만 남김
RE_UNIV_OF = re.compile(
    r'\b(?:University|Institute|College|School|Academy)\s+of\s+'
    r'[A-Z][a-zA-Z\s\'\-\.]{2,40}\b'
)

# ── 정식 인가 대학교 보호 목록 (졸업학교 삭제 금지) ─────────────────────────
# 미국/영국/캐나다/호주/뉴질랜드/남아공/아일랜드 + 한국 주요 대학
# school_replacer에서 이 목록에 포함된 이름은 삭제하지 않음
_PROTECTED_UNIVERSITIES = frozenset({
    # ── 미국 ──
    "harvard university", "yale university", "princeton university",
    "columbia university", "cornell university", "brown university",
    "dartmouth college", "university of pennsylvania",
    "massachusetts institute of technology",
    "stanford university", "california institute of technology",
    "johns hopkins university", "duke university", "northwestern university",
    "vanderbilt university", "rice university", "notre dame",
    "university of notre dame", "georgetown university", "emory university",
    "carnegie mellon university", "wake forest university", "tufts university",
    "boston university", "boston college", "new york university",
    "fordham university", "villanova university", "loyola university",
    "university of california", "uc berkeley", "ucla",
    "university of michigan", "michigan state university",
    "university of texas", "texas a&m university",
    "university of florida", "florida state university",
    "university of virginia", "virginia tech",
    "university of north carolina", "north carolina state university",
    "ohio state university", "university of ohio",
    "penn state", "pennsylvania state university",
    "university of illinois", "illinois state university",
    "university of wisconsin", "university of washington",
    "university of minnesota", "university of georgia",
    "georgia institute of technology", "georgia tech",
    "purdue university", "indiana university",
    "university of iowa", "iowa state university",
    "university of missouri", "university of arizona",
    "arizona state university", "university of colorado",
    "colorado state university", "university of utah",
    "brigham young university", "university of oregon",
    "oregon state university", "university of southern california",
    "university of pittsburgh", "drexel university",
    "temple university", "university of miami",
    "university of south florida", "university of central florida",
    "university of alabama", "auburn university",
    "louisiana state university", "tulane university",
    "university of tennessee", "university of kentucky",
    "university of arkansas", "oklahoma state university",
    "university of oklahoma", "university of nebraska",
    "university of kansas", "kansas state university",
    "university of nevada", "university of hawaii",
    "san diego state university", "california state university",
    "san francisco state university", "university of new mexico",
    "university of alaska", "university of connecticut",
    "university of massachusetts", "university of vermont",
    "university of new hampshire", "university of rhode island",
    "university of delaware", "university of maryland",
    "university of richmond", "william & mary",
    "college of william and mary",
    "oberlin college", "vassar college", "amherst college",
    "williams college", "swarthmore college", "wellesley college",
    "smith college", "mount holyoke college", "bryn mawr college",
    "barnard college", "colgate university", "hamilton college",
    "colby college", "bates college", "bowdoin college",
    "middlebury college", "trinity college", "connecticut college",
    "lafayette college", "lehigh university", "bucknell university",
    "gettysburg college", "dickinson college", "skidmore college",
    "union college", "ithaca college",
    "eastern michigan university", "western michigan university",
    "central michigan university", "northern michigan university",
    "wayne state university", "university of denver",
    "university of tulsa", "university of dayton",
    "xavier university", "creighton university",
    "gonzaga university", "seattle university",
    "portland state university", "university of portland",
    "university of san francisco", "santa clara university",
    "university of san diego", "pepperdine university",
    "university of redlands", "california baptist university",
    "biola university", "azusa pacific university",
    # ── 영국 ──
    "university of oxford", "university of cambridge",
    "oxford university", "cambridge university",
    "university college london", "ucl",
    "imperial college london", "imperial college",
    "king's college london", "london school of economics",
    "lse", "university of edinburgh",
    "university of manchester", "university of birmingham",
    "university of bristol", "university of warwick",
    "university of exeter", "university of bath",
    "university of leeds", "university of sheffield",
    "university of liverpool", "university of nottingham",
    "university of southampton", "university of leicester",
    "university of reading", "university of surrey",
    "university of sussex", "university of kent",
    "university of east anglia", "university of york",
    "university of durham", "durham university",
    "university of st andrews", "st andrews university",
    "university of glasgow", "university of strathclyde",
    "heriot-watt university", "university of aberdeen",
    "queen's university belfast", "university of ulster",
    "cardiff university", "swansea university",
    "university of wales", "aberystwyth university",
    "queen mary university of london",
    "royal holloway university of london",
    "city university london", "brunel university",
    "university of portsmouth", "university of brighton",
    "university of northampton", "de montfort university",
    "coventry university", "aston university",
    "loughborough university", "keele university",
    "university of lincoln", "university of hull",
    "university of salford", "manchester metropolitan university",
    "sheffield hallam university", "leeds beckett university",
    "northumbria university", "newcastle university",
    "university of newcastle", "teesside university",
    "university of sunderland", "university of the arts london",
    "goldsmiths", "goldsmiths college",
    # ── 캐나다 ──
    "university of toronto", "mcgill university",
    "university of british columbia", "ubc",
    "university of alberta", "university of calgary",
    "western university", "university of western ontario",
    "queen's university", "university of ottawa",
    "university of waterloo", "wilfrid laurier university",
    "mcmaster university", "york university",
    "concordia university", "université de montréal",
    "university of montreal", "université laval",
    "dalhousie university", "university of new brunswick",
    "memorial university of newfoundland",
    "university of victoria", "simon fraser university",
    "university of manitoba", "university of saskatchewan",
    "university of regina", "university of lethbridge",
    "brock university", "ryerson university",
    "toronto metropolitan university",
    "ontario tech university", "trent university",
    "nipissing university", "algoma university",
    # ── 호주 ──
    "university of melbourne", "university of sydney",
    "australian national university", "anu",
    "university of new south wales", "unsw",
    "monash university", "university of queensland",
    "university of western australia", "uwa",
    "university of adelaide", "macquarie university",
    "university of wollongong", "university of newcastle",
    "james cook university", "griffith university",
    "queensland university of technology", "qut",
    "rmit university", "deakin university",
    "la trobe university", "curtin university",
    "murdoch university", "edith cowan university",
    "flinders university", "charles darwin university",
    "university of tasmania", "bond university",
    "swinburne university", "victoria university",
    "university of south australia", "unisa",
    "university of canberra", "charles sturt university",
    # ── 뉴질랜드 ──
    "university of auckland", "auckland university",
    "victoria university of wellington",
    "university of canterbury", "lincoln university",
    "university of otago", "university of waikato",
    "massey university", "auckland university of technology",
    "aut university",
    # ── 남아프리카공화국 ──
    "university of cape town", "uct",
    "university of the witwatersrand", "wits university",
    "wits", "stellenbosch university",
    "university of pretoria", "up",
    "university of kwazulu-natal", "ukzn",
    "rhodes university", "university of the free state",
    "university of johannesburg", "uj",
    "north-west university", "nwu",
    "university of the western cape", "uwc",
    "nelson mandela university", "walter sisulu university",
    "university of limpopo", "university of venda",
    "university of zululand", "university of fort hare",
    "tshwane university of technology", "tut",
    "cape peninsula university of technology", "cput",
    "central university of technology", "cut",
    # ── 아일랜드 ──
    "trinity college dublin", "tcd",
    "university college dublin", "ucd",
    "university college cork", "ucc",
    "national university of ireland galway", "nuig",
    "university of galway",
    "dublin city university", "dcu",
    "maynooth university", "university of limerick",
    "dublin institute of technology", "technological university dublin",
    "tu dublin",
    # ── 한국 ──
    "seoul national university", "snu",
    "yonsei university", "korea university",
    "sungkyunkwan university", "skku",
    "hanyang university", "sogang university",
    "ewha womans university", "ewha woman's university",
    "postech", "pohang university of science and technology",
    "kaist", "korea advanced institute of science and technology",
    "ulsan national institute of science and technology", "unist",
    "dgist", "daegu gyeongbuk institute of science and technology",
    "gist", "gwangju institute of science and technology",
    "university of seoul", "konkuk university",
    "kyung hee university", "inha university",
    "chung-ang university", "cau",
    "ajou university", "sejong university",
    "dongguk university", "hongik university",
    "kookmin university", "myongji university",
    "sookmyung women's university", "duksung women's university",
    "hankuk university of foreign studies", "hufs",
    "korean university of foreign studies",
    "korea national university of education",
    "pusan national university", "kyungpook national university",
    "chonnam national university", "jeonbuk national university",
    "chungnam national university", "chungbuk national university",
    "강남대학교", "경희대학교", "고려대학교", "서울대학교", "연세대학교",
    "성균관대학교", "한양대학교", "이화여자대학교", "부산대학교",
    # ── 스위스/유럽 국제 연구대학 ──
    "eth zurich", "eth zürich", "eth zurich university", "eth zürich university",
    "epfl", "école polytechnique fédérale de lausanne",
    "university of zurich", "university of zürich",
    "university of geneva", "university of bern", "university of basel",
    "university of amsterdam", "university of groningen",
    "leiden university", "delft university of technology",
    "technical university of munich", "tu munich", "tu münchen",
    "ludwig maximilian university", "lmu munich",
    "humboldt university", "humboldt university of berlin",
    "free university of berlin", "freie universität berlin",
    "university of vienna", "vienna university of technology",
    "university of copenhagen", "university of oslo",
    "university of stockholm", "stockholm university",
    "karolinska institute", "chalmers university of technology",
    "university of helsinki",
    "paris sorbonne", "sorbonne university", "paris-sorbonne university",
    "école normale supérieure", "sciences po",
    "university of paris", "paris descartes university",
    "ku leuven", "university of leuven", "ghent university",
    "university of bologna", "università di bologna",
    "sapienza university of rome", "sapienza università di roma",
    "university of milan", "università degli studi di milano",
    "universitat barcelona", "university of barcelona",
    "autonomous university of madrid",
    "university of lisbon", "university of porto",
})

# 미국 도시+주: "Houston, TX", "New Orleans, Louisiana"
_US_STATES = (
    r'AL|AK|AZ|AR|CA|CO|CT|DE|FL|GA|HI|ID|IL|IN|IA|KS|KY|LA|ME|MD|MA|'
    r'MI|MN|MS|MO|MT|NE|NV|NH|NJ|NM|NY|NC|ND|OH|OK|OR|PA|RI|SC|SD|TN|'
    r'TX|UT|VT|VA|WA|WV|WI|WY|'
    r'Alabama|Alaska|Arizona|Arkansas|California|Colorado|Connecticut|'
    r'Delaware|Florida|Georgia|Hawaii|Idaho|Illinois|Indiana|Iowa|Kansas|'
    r'Kentucky|Louisiana|Maine|Maryland|Massachusetts|Michigan|Minnesota|'
    r'Mississippi|Missouri|Montana|Nebraska|Nevada|New\s+Hampshire|'
    r'New\s+Jersey|New\s+Mexico|New\s+York|North\s+Carolina|North\s+Dakota|'
    r'Ohio|Oklahoma|Oregon|Pennsylvania|Rhode\s+Island|South\s+Carolina|'
    r'South\s+Dakota|Tennessee|Texas|Utah|Vermont|Virginia|Washington|'
    r'West\s+Virginia|Wisconsin|Wyoming'
)
RE_US_CITY_STATE = re.compile(
    r'\b[A-Z][a-zA-Z]+(?:[ ]+[A-Z][a-zA-Z]+)?,\s*(?:' + _US_STATES + r')\b',
    re.I
)

# 외국 도시+국가: "Sheffield, United Kingdom", "Beijing, China"
RE_FOREIGN_CITY = re.compile(
    # [a-zA-Z\-]+(?:[ ][a-zA-Z\-]+){0,3}: 단어 사이 단일 공백만 허용 (PDF 다중공백 오탐 방지)
    r'\b[A-Z][a-zA-Z\-]+(?:[ ][a-zA-Z\-]+){0,3},\s*'
    r'(?:United\s+Kingdom|England|Scotland|Wales|Northern\s+Ireland|'
    r'China|Japan|Germany|France|South\s+Africa|Australia|New\s+Zealand|'
    r'Canada|USA|United\s+States(?:\s+of\s+America)?|Brazil|India|'
    r'Philippines|Zimbabwe|Zambia|Nigeria|Kenya|Ghana|Uganda)\b',
    re.I
)

# 한국 도시+국가 주소 패턴: "Gwangmyeong, South Korea", "Suwon, Korea"
# → 통째로 redact (주소줄 전체 삭제)
RE_KR_CITY_COUNTRY = re.compile(
    r'\b[A-Za-z][A-Za-z\s\-]{2,30},\s*(?:South\s+Korea|Korea|Republic\s+of\s+Korea)\b',
    re.I
)

# 나이 언급: "I am a 33 year old", "I'm a 34-year-old"
RE_AGE_MENTION = re.compile(
    r"\bI(?:'m|\s+am)\s+(?:a\s+)?\d{1,2}[\s\-]?year[\s\-]?old\b",
    re.I
)

# 비자 상태 공개: "E-2 VISA ELIGIBLE", "E2 visa holder"
RE_VISA_STATUS = re.compile(
    r'\b(?:E-?\d|F-?\d|H-?\d|G-?\d)\s*VISA\b(?:\s+(?:ELIGIBLE|HOLDER|STATUS|SPONSORSHIP))?\b',
    re.I
)

# 한국 도시/지역 키워드 (소문자)
KR_KEYWORDS = frozenset([
    "korea", "seoul", "busan", "daegu", "incheon", "gwangju",
    "daejeon", "ulsan", "sejong", "gyeonggi", "gangwon",
    "chungbuk", "chungnam", "jeonbuk", "jeonnam",
    "gyeongbuk", "gyeongnam", "jeju",
    # 도(道) 접미사 포함 변형 — "Gangwon-do" 등에서 "-do" 잔류 방지
    "gangwon-do", "gyeonggi-do", "chungbuk-do", "chungnam-do",
    "jeonbuk-do", "jeonnam-do", "gyeongbuk-do", "gyeongnam-do", "jeju-do",
    "south korea", "republic of korea", "rok",
    "한국", "서울", "부산", "대구", "인천", "광주", "대전", "울산",
    "세종", "경기", "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주",
    "suwon", "seongnam", "goyang", "yongin", "changwon",
    "ansan", "anyang", "namyangju", "hwaseong", "pyeongtaek",
    "gimpo", "gwacheon", "uijeongbu", "paju", "yangju",
    "pocheon", "dongducheon", "guri", "hanam", "icheon",
    "osan", "gunpo", "uiwang", "siheung", "bucheon",
    "jeonju", "cheongju", "cheonan", "asan",
    "gimhae", "yangsan", "geoje", "tongyeong",
    "mokpo", "yeosu", "suncheon", "gwangyang",
    "gyeongju", "pohang", "andong", "gumi",
    "chuncheon", "wonju", "gangneung", "sokcho",
    "itaewon", "gangnam", "hongdae", "sinchon", "mapo",
    "jamsil", "songpa", "yongsan", "nowon", "bundang",
    "ilsan", "dongtan", "pangyo",
    # 경기도 주요 도시 (추가)
    "gwangmyeong", "uijeongbu", "gapyeong", "yeoju",
    "anseong", "pyeongtaek", "dongducheon", "gapyeong",
    # 기타 주요 시
    "mokpo", "chungju", "jecheon", "boryeong",
    # 송도/동백/수지 — 인천·용인 자주 등장하는 지구명
    "songdo", "song-do", "dongbaek", "suji", "giheung",
    # 수도권 신도시/지구명
    "bupyeong", "bupyeong-gu",
    # Gyeonggi-do 오탈자/약식 표기
    "gyeongi-do", "gyeongi", "gyeong-do",
    # -si 접미사 (도시명 제거 후 잔류 방지)
    "paju-si", "incheon-si", "suwon-si", "seongnam-si",
])

# PII 라벨 → 해당 줄 전체 삭제
# 주의: "line", "cell" 등 짧은 단어는 ":" 필수 조건으로 오탐 방지
PII_LINE_LABELS = [
    # 연락처 — 구체적 표현 먼저, 범용 나중 (순서 중요)
    ("email address", True), ("email", True), ("e-mail", True),
    ("phone number", True), ("phone no", True), ("phone", True),
    ("telephone number", True), ("telephone", True),
    ("mobile number", True), ("mobile no", True), ("mobile", True),
    ("cell phone", True), ("cell no", True), ("cell", True),
    ("home phone", True), ("work phone", True), ("office phone", True),
    ("korean telephone", True), ("u.k. telephone", True),
    ("uk telephone", True), ("us telephone", True),
    ("contact number", True), ("contact no", True), ("contact", True),
    ("tel", True),
    # 주소
    ("permanent address", True), ("current address", True),
    ("home address", True), ("mailing address", True),
    ("residential address", True), ("address", True),
    # SNS (구체적 표현 먼저)
    ("kakao id", True), ("kakaotalk id", True), ("kakaotalk", True), ("kakao", True),
    ("wechat id", True), ("we chat id", True), ("wechat", True), ("we chat", True),
    ("line id", True), ("line", True),
    ("skype id", True), ("skype", True),
    ("telegram id", True), ("telegram", True),
    ("whatsapp number", True), ("whatsapp no", True), ("whatsapp", True),
    ("signal", True), ("viber", True),
    ("instagram", True), ("facebook", True), ("twitter", True), ("tiktok", True),
    # 링크
    ("linkedin", True), ("personal website", True), ("website", True),
    ("portfolio", True), ("github", True),
    # 신원
    ("passport number", True), ("passport no", True), ("passport", True),
    ("date of birth", True), ("birth date", True), ("dob", True),
    ("national id", True), ("national identification", True), ("ssn", True),
    # nationality/citizenship 은 고용주가 필요한 정보 → 삭제 안 함
    # ("nationality", True), ("citizenship", True),
    ("age", True),                          # "Age: 28"
    ("marital status", True),              # "Marital status: Single"
    ("civil status", True),               # "Civil status: Single"
    ("race", True), ("ethnicity", True), ("religion", True),
    ("birth", True), ("born", True),      # "Birth: UK / 1990"
    ("gender", True), ("sex", True),      # "Gender: Female"
    ("emergency contact", True), ("emergency", True),
    ("next of kin", True),
    # 한국어
    ("이메일", True), ("전화", True), ("연락처", True),
    ("카카오", True), ("여권", True), ("주소", True),
    ("위챗", True), ("텔레그램", True), ("라인", True),
]

# 위치 라벨 → 한국이면 "Korea"로 변환
LOCATION_LABELS = [
    "current location", "current address", "address", "location",
    "residence", "residing in", "living in", "based in",
    "city", "province", "country of residence",
    "현재 위치", "거주지", "주소", "현위치",
]

# 직장명 라벨 → 값 삭제 (institution/university 제외)
WORKPLACE_LABELS = [
    "current employer", "employer", "company", "workplace",
    "school name", "academy name", "hagwon",
    "현 직장", "근무지", "학원명", "학교명",
]

# 한국 학원/학교 키워드 (경력줄에서 한국 근무지 감지용)
KR_WORKPLACE_KEYWORDS = frozenset([
    # ══ 대형 프랜차이즈 체인 (전국구) ══
    "ybm", "ybm ecc", "ecc",
    "pagoda", "pagoda english",
    "chungdahm", "chungdahm learning", "chungdam", "cdl",
    "poly", "poly school", "poly english",
    "avalon", "avalon english",
    "sisa", "sisa english",
    "jungchul", "jeongchul", "jung chul",
    "hansol", "hansol education",
    "daekyo", "daekyo education",
    "kumon",
    "bricks", "bricks english",
    "slp", "slp english",
    "rise english", "rise korea", "rise english academy", "rise english kindergarten",  # rise 단독 제거 (동사 오탐)
    "cedar hill", "cedar hill global prep", "cedar hill global",    # 인천 국제학교
    "korea poly", "korea poly school",                               # "Korea Poly School" 전체 매칭 (poly school 개별 제거 후 Korea 잔류 방지)
    "luxx", "luxx language academy", "luxx english",                 # luxx 단독은 짧지만 단독 단어로만 사용
    "carrot global", "carrot english", "carrot english academy",     # carrot 단독 제거 (일반명사 오탐)
    "4kidsedu", "4kids edu", "4kids english",                        # 숫자 포함 브랜드
    "cdi", "cdi english",
    "jle", "jlec",
    "aclipse",
    "wonderland", "wonderland english",
    "maple bear",
    "little america",
    "little fox",
    "creverse",
    "reading town",
    "english egg",
    "english channel",
    "english town",
    "english village",
    "english plus",
    "english muse",
    "like english",
    "kids college",
    "canlearn",
    "eduplex",
    "lucid english", "lucid academy",       # lucid 단독 제거 (형용사 오탐)
    "topia",
    "snb english",                          # snb 단독 제거
    "jls",
    "gse english", "gse academy",          # gse 단독 제거 (Pearson GSE 점수 오탐)
    "ils english", "ils academy",          # ils 단독 제거
    "fastrackids",
    "ding ding dang",
    "seouldal",
    "top edu",
    "dada", "dadahak",
    "altiora",
    "emg education",
    "emg",
    "hampson",
    "engoo",
    "gnb", "gnb english",           # bcm/ael/tel 제거 — 일반 영어단어 오탐
    "hess english", "hess english school",  # hess 단독 제거 (회사명 오탐)
    "global adventure",
    "chungjae",
    "empire english", "empire academy",     # asset/empire/lucid 단독 제거
    "lucid english",
    # ══ 2026년 5월 실시간 업데이트 — 확인된 신규 브랜드 ══
    # JLS 계열 (CHESS·ACE·V.GROUP — 한국 최대 영어학원 네트워크 125+ 캠퍼스)
    "chess english", "chess english academy", "chess academy",  # chess 단독 제거 (취미)
    "v group", "v.group",
    "canb english", "canb academy",         # canb 단독 제거
    # DYB / 최선어학원 계열
    "dyb", "dyb choisun", "dyb english",
    # 헬렌도론 (글로벌 유아영어 프랜차이즈)
    "helen doron", "helen doron english",
    # iGarten / 아이가르텐 (청담어학원 영어유치원 브랜드)
    "igarten", "i-garten", "i garten", "아이가르텐",
    # 크레버스 계열 (April+청담+iGarten 모회사)
    "april english", "april institute", "chungdahm april",  # april 단독 제거 (월 이름 오탐)
    # 윤선생 계열
    "yoons", "yoon's english", "윤선생",
    "yoons igse", "yoons forest",
    # 삼성영어 (삼성출판사)
    "samsung english", "samsung selena", "삼성영어",
    # 민병철 계열
    "min byungchul", "minbyungchul", "민병철",
    "minbyungchul speaking works",
    # 단비교육
    "danbi", "danbi english", "단비교육",
    # 앙코아어학원
    "ankoa", "앙코아",
    # 링구아포럼
    "lingua forum", "linguaforum", "링구아포럼",
    # 야나두 (온라인 영어)
    "yanadoo", "야나두",
    # 스피킹맥스
    "speaking max", "speakingmax", "스피킹맥스",
    # 케이크 (카카오 계열 영어앱)
    "cake english", "cake app",
    # 튜터링
    "tutoring english", "tutoring",
    # 링글 (1:1 화상영어)
    "ringle",
    # 그 외 korvia 목록 확인 브랜드
    "apex english", "apex",
    "apple english academy",
    "arte english",
    "bis academy",
    "bono english",
    "bibakids", "biba kids",
    "billy english",
    "blooming kinder",
    "brighton english",
    "british education korea",
    "canada beavers",
    "callan institute", "callan english",   # callan 단독 제거 (Callan Method 교수법 오탐)
    "cleverkinder", "clever kinder",
    "dna english",
    "dreamedu",
    "english joy",
    "english republic",
    "englfit",
    "gaon english",
    "gem academy",
    "gemstone english",
    "gilbert language",
    "glovia english",
    "guilford",
    "hans language",
    "heritage english",
    "ile academy",
    "ingeniiis",
    "inspire english",
    "lingua english",
    "dux academy",                          # dux 단독 제거 (수석졸업상 오탐)
    "worwick franklin", "worwick",
    "english vine",
    "eie english", "eie",
    "kics", "kics prep",
    # ══ 영어유치원 / 유아 영어 브랜드 ══
    "wonderkids",
    "kinderhaus",
    "kling", "kling english",
    "smartbear", "smart bear",
    "early birds",
    "little scholars",
    "cozy english",
    "bright english",
    "smart english",
    "kids english",
    "little english",
    "happy english",
    "sunshine english",
    "english garden",
    "english land",
    "english world",
    "english house",
    "english hub",
    "english kids",
    "fun english",
    "magic english",
    "rainbow english",
    "abc english",
    "wondertree",
    "ivykids", "ivy kids",
    "cambridge kids",
    "oxford kids",
    "stanford kids",
    "harvard kids",
    "yale kids",
    "first english",
    "prime english",
    "top english",
    "global kids",
    "global english",
    "i love english",
    "yes english",
    "best english",
    "ace english",
    "star english",
    "sky english",
    "spring english",
    "green english",
    "white english",
    "blue english",
    # ══ 사립 영어전문학원 브랜드 ══
    "gnb english",
    "gnb",
    "april english",                        # april 단독 제거 중복 정리
    "iwell", "i-well",
    "sprintclass",
    "visang", "비상",
    "woongjin", "웅진",
    "longman korea", "longman",
    "ericsson english", "ericsson",
    "gate english", "gate",
    "megastudy english", "megastudy",
    "etoos english", "etoos",
    "ebsi english",
    "class101 english", "class101",
    "engstudy",
    "tmax english",
    "talking english",
    "speak english",
    "native english",
    "real english",
    "pure english",
    "fresh english",
    "open english",
    "new english",
    "one english",
    "all english",
    "cool english",
    "joy english",
    "pro english",
    "max english",
    "go english",
    "fast english",
    "power english",
    "win english",
    "lead english",
    "king english",
    "queen english",
    "champion english",
    "genius english",
    "master english",
    "super english",
    "great english",
    "gold english",
    "silver english",
    "plus english",
    "a+ english", "a plus english",
    "number one english", "no.1 english",
    # ══ 지역 기반 유명 학원 브랜드 ══
    "busan global village", "global village",
    "hillside collegiate", "hillside",
    "berkeley language", "berkeley english",
    "francis parker",
    "daejeon english",
    "incheon english",
    "gyeonggi english",
    "seoul english",
    "gangnam english",
    "apgujeong english",
    "bundang english",
    "ilsan english",
    "suwon english",
    "songdo english",
    # ══ 공공 기관 / 정부 운영 영어 프로그램 ══
    "epik", "gepik", "smoe", "jlec",
    "jeollanamdo epik",
    "english program in korea",
    "global english village",
    "english education center",
    "english learning center",
    # ══ 일반 유형명 (보존 대상) ══
    "language school", "language academy",
    "language institute", "language center", "language centre",
    "english academy", "english institute",
    "private academy", "teaching academy",
    "hagwon", "학원", "어학원", "영어학원", "영어유치원",
    "유치원", "어린이집", "유아원",
    "elementary school", "middle school", "high school",
    "primary school", "secondary school", "independent school",
    "grammar school", "boarding school", "prep school",
    "international school", "kindergarten", "preschool",
    "after-school", "afterschool", "after school",
])

# 이미 일반 유형명인 KR_WORKPLACE_KEYWORDS 항목
# → 앞의 브랜드명만 제거하고 유형명 그대로 보존
# "Broad Language School" → "Language School" (not "English")
_GENERIC_SCHOOL_TYPES: dict = {
    "language school": "Language School",
    "language academy": "Language Academy",
    "language institute": "Language Institute",
    "language center": "Language Center",
    "language centre": "Language Centre",
    "english village": "English Village",
    "english academy": "English Academy",
    "english institute": "English Institute",
    "english center": "English Center",
    "english centre": "English Centre",
    "english school": "English School",
    "teaching academy": "Teaching Academy",
    "private academy": "Private Academy",
    "international school": "International School",
    "elementary school": "Elementary School",
    "primary school": "Primary School",
    "secondary school": "Secondary School",
    "middle school": "Middle School",
    "high school": "High School",
    "kindergarten": "Kindergarten",
    "preschool": "Preschool",
    "after school": "After School",
    "after-school": "After-School",
    "afterschool": "Afterschool",
    # 일반 서비스어 — KR_WORKPLACE_KEYWORDS에 포함되나 비한국 맥락에서도 자주 쓰임
    # 오탐 방지: 삭제 대상에서 제외하고 유형명으로 유지
    "tutoring": "Tutoring",
    "hagwon": "Hagwon",
    "global english": "Global English",
    "global village": "Global Village",
    "native english": "Native English",
}
_GENERIC_TYPES_SET = frozenset(_GENERIC_SCHOOL_TYPES.keys())

# 브랜드 접두사 + 일반 유형명 → 유형명만 보존 (capture group 1 = generic type)
# "Broad Language School" → "Language School"
# !! 줄 경계 넘어 매칭 방지: [ \t]+ 사용 (not \s+) !!
_RE_KR_SCHOOL_PREFIX = re.compile(
    # !! re.IGNORECASE 제거: [A-Z] prefix가 lowercase "in" 등을 매칭하던 버그 수정
    # 예: "in Kindergarten" → IGNORECASE 있으면 매칭됨 → 제거
    r'\b[A-Z][a-zA-Z\'\-\.]+(?:[ \t]+[A-Z][a-zA-Z\'\-\.]+){0,2}[ \t]+'
    r'(Language[ \t]+School|Language[ \t]+Academy|Language[ \t]+Institute|'
    r'Language[ \t]+Cent(?:er|re)|English[ \t]+(?:Village|Academy|Institute|Cent(?:er|re)|School|Kindergarten)|'
    r'Teaching[ \t]+Academy|Private[ \t]+Academy|International[ \t]+School|'
    r'(?:Elementary|Primary|Secondary|Middle|High)[ \t]+School|'
    r'Kindergarten|Preschool)\b',
)

# 이력서 섹션 헤더 (PDF 이름 추출 시 오탐 방지)
_RESUME_SECTION_HEADERS = frozenset({
    "professional summary", "work experience", "education", "references",
    "skills", "qualifications", "experience", "objective",
    "personal statement", "career summary", "career objective",
    "profile summary", "employment history", "work history",
    "professional experience", "relevant experience",
    "academic qualifications", "certifications", "teaching experience",
    "teaching history", "additional skills", "languages", "interests",
    "volunteer experience", "professional development", "awards",
    "curriculum vitae", "personal profile", "about me",
    "contact information", "contact details", "personal information",
    "personal details", "key skills", "core competencies",
    "summary of qualifications", "professional skills",
    "cover letter", "letter of introduction", "summary",
})

# 이름으로 오인하면 안 되는 단어 (파일명 + PDF 텍스트 양쪽 사용)
_NON_NAME_WORDS = frozenset({
    # CV/파일명 관련
    "resume", "cv", "curriculum", "vitae", "cover", "letter",
    "teaching", "teacher", "english", "esl", "tesol", "tefl",
    "application", "updated", "final", "new", "copy", "version",
    "document", "file", "scan", "draft", "official",
    # 주소/장소
    "street", "road", "avenue", "drive", "lane", "court", "place",
    "boulevard", "terrace", "close", "crescent", "way", "park",
    "phase", "block", "floor", "unit", "apt", "suite", "room",
    "north", "south", "east", "west", "city",
    # 국가/지역 (이름 오인 방지)
    "usa", "korea", "canada", "australia", "zealand", "africa",
    "britain", "england", "ireland", "scotland", "wales",
    "america", "states", "republic", "province", "district",
    # 일반 영단어 (이름에는 없는)
    "and", "the", "for", "with", "from", "into", "about",
    "skills", "profile", "summary", "experience", "education",
    # 교육기관 유형명 (이름 오인 방지 — Liberty University, Francis Parker Collegiate 등)
    "university", "college", "collegiate", "polytechnic", "institute",
    "academy", "faculty", "campus", "seminary",
    # 커버레터 인사말 (이름 오인 방지 — "Dear Hiring Manager")
    "dear", "hiring", "manager", "sincerely", "regards", "warmly",
    "faithfully", "truly", "respectfully",
    # 자격증/학력 라벨 (이름 오인 방지 — "Endorsements: Elementary")
    "endorsements", "endorsement", "certifications", "certification",
    "elementary", "bachelor", "master", "doctorate",
    "license", "licensed", "diploma", "degree",
})

# ── EDUCATION 섹션 보호용 헤더 감지 정규식 ──
# Pass 2.7(학교명 일반화)은 EDUCATION 섹션에서 실행 금지
_EDU_SECTION_RE = re.compile(
    r'^(?:education|academic\s+(?:qualifications?|background|history)|'
    r'educational\s+(?:background|history)|qualifications?|schooling)',
    re.IGNORECASE,
)
# 다른 섹션 헤더 감지 → EDUCATION 섹션 종료 신호
_OTHER_SECTION_RE = re.compile(
    r'^(?:(?:related|relevant|additional|other|extra|recent)\s+)?'
    r'(?:experience|work(?:\s+(?:experience|history))?|employment(?:\s+history)?|'
    r'teaching(?:\s+experience)?|professional(?:\s+experience)?|'
    r'career(?:\s+(?:summary|objective))?|volunteer|internship|reference|'
    r'skills?|languages?|interests?|awards?|certifications?|'
    r'summary|objective|profile|about\s+me|personal)',
    re.IGNORECASE,
)

# 경력줄에서 한국 근무 상세 축약 패턴
# "YBM ECC, Uijeongbu, South Korea, March 2021 - Sept 2022" → "South Korea, March 2021 - Sept 2022"
# 날짜 패턴: "Jan 2021", "2021-2023", "March 2021 - Sept 2022", "2021 - Present"
_MONTH = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
)
RE_DATE_RANGE = re.compile(
    r"(?:,\s*)?"
    r"(" + _MONTH + r"\.?\s*)?"
    r"(\d{4})"
    r"(?:\s*[-\u2013~]\s*"
    r"(?:" + _MONTH + r"\.?\s*)?"
    r"(?:\d{4}|[Pp]resent|[Cc]urrent)"
    r")?",
    re.I,
)


# ── 연락 요청 문장 감지 ──────────────────────────────────────────────────────
# "please email me at X", "find me on WhatsApp (+44...)", "reach me at..."
# → 커버레터/이력서에서 해당 문장 전체 삭제
# 조건: 필자(me/us)에게 연락하라는 초대 문구 + 연락처 or 플랫폼 언급
_CONTACT_INVITE_RE = re.compile(
    r'\b(?:'
    r'please\s+(?:email|call|contact|reach|text|message|find)\s+(?:me|us)\b'
    r'|email\s+me\s+(?:at|on|directly)\b'
    r'|reach\s+(?:me|out\s+to\s+me)\s+(?:at|on|via|by|through)\b'
    r'|contact\s+me\s+(?:at|on|via|by|through|directly|using)\b'
    r'|find\s+me\s+(?:on|at|via|through)\b'
    r'|(?:on|via|through|over|using)\s+whatsapp\b'
    r'|(?:on|via|through)\s+wechat\b'
    r'|(?:on|via|through)\s+kakao(?:talk)?\b'
    r'|my\s+(?:email|phone|mobile|number|whatsapp|wechat|contact)\s+(?:is|address\s+is|number\s+is)\b'
    r'|i\s+can\s+be\s+reached\s+(?:at|on|via|by|through)\b'
    r'|i\s+am\s+reachable\s+(?:at|on|via|by)\b'
    r'|feel\s+free\s+to\s+(?:email|call|contact|reach|message)\s+(?:me|us)\b'
    r'|(?:do\s+not|don\'t)\s+hesitate\s+to\s+(?:contact|reach|email|call)\s+(?:me|us)\b'
    r'|at\s+your\s+convenience.*?(?:email|call|contact|reach|message)\s+me\b'
    r')',
    re.IGNORECASE | re.DOTALL,
)
# 연락 플랫폼/수단 단독 언급 (WhatsApp/WeChat 번호 포함 문장)
_CONTACT_PLATFORM_RE = re.compile(
    r'\b(?:whatsapp|wechat|kakao(?:talk)?|viber|telegram|line\s+app)\b',
    re.IGNORECASE,
)

# ══════════════════════════════════════════════════════
#  3. PII 삭제 엔진
# ══════════════════════════════════════════════════════

def _is_korea(text: str) -> bool:
    lower = text.lower().strip()
    return any(kw in lower for kw in KR_KEYWORDS)


def _contains_protected_university(text: str) -> bool:
    """텍스트에 정식 인가 대학교명이 포함되어 있는지 확인 (보호 우선 판별)"""
    tl = text.lower()
    return any(pu in tl for pu in _PROTECTED_UNIVERSITIES)


def _replace_workplace_generic(value: str) -> str:
    """
    근무처 명을 일반명으로 대체.
    - 영어 관련: "English"
    - 학원/아카데미: "Academy"
    - 대학: "University"
    """
    cleaned = value.strip()
    if not cleaned:
        return "English"

    # 기호 제거
    cleaned = re.sub(r"[!@#$%&*\-_~'\".,;():\[\]{}/?\\]+", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    lower = cleaned.lower()

    # 키워드 매칭 (우선순위 순)
    if any(kw in lower for kw in ["english", "esl", "language", "polyglot", "poly", "ybm", "ecc"]):
        return "English"
    if any(kw in lower for kw in ["academy", "hagwon", "cram", "institute", "school"]):
        return "Academy"
    if any(kw in lower for kw in ["university", "college"]):
        return "University"
    if any(kw in lower for kw in ["company", "corporation", "co", "ltd", "inc"]):
        return "Company"

    # 기본값
    return "English"


def _school_to_generic(match_text: str) -> str:
    """학교/기관명을 일반 유형명으로 대체."""
    lower = match_text.lower()
    if "university" in lower or "college" in lower or "polytechnic" in lower:
        return "University"
    return "School"


def _name_variants(full_name: str) -> list[str]:
    """
    이름의 모든 인식 가능 변형 생성.
    - 풀네임, First Last, Last First, 각 단어(3글자 이상만)
    - 세미콜론/괄호 별명 처리
    긴 변형부터 정렬 (greedy 매칭)
    """
    if not full_name or not full_name.strip():
        return []

    raw = full_name.strip()
    variants = set()
    variants.add(raw)

    # 별명 분리: "Jashen Arianne ; Shen" → main + nickname
    main_name = raw
    if ";" in raw:
        parts = raw.split(";", 1)
        main_name = parts[0].strip()
        nick = parts[1].strip()
        variants.add(main_name)
        if nick and len(nick) > 2:
            variants.add(nick)

    if "(" in main_name:
        clean = re.sub(r"\s*\([^)]*\)", "", main_name).strip()
        nick_m = re.search(r"\(([^)]+)\)", main_name)
        variants.add(clean)
        if nick_m and len(nick_m.group(1)) > 2:
            variants.add(nick_m.group(1))
        main_name = clean

    words = main_name.split()

    if len(words) >= 2:
        # First Last
        variants.add(f"{words[0]} {words[-1]}")
        # Last, First
        variants.add(f"{words[-1]}, {words[0]}")
        # Last First
        variants.add(f"{words[-1]} {words[0]}")
        # 모든 인접 2-word 조합 (중간이름 포함)
        for i in range(len(words) - 1):
            variants.add(f"{words[i]} {words[i+1]}")

    # 개별 단어 (3글자 이상만, 너무 짧으면 오탐)
    for w in words:
        if len(w) >= 3:
            variants.add(w)

    # 긴 것부터 매칭 (greedy)
    return sorted(variants, key=len, reverse=True)


def _should_skip_line(stripped_lower: str) -> bool:
    """PII 라벨로 시작하는 줄인지 판별.
    불릿(•·▪▸), 아이콘, 탭 등 앞부분 비알파벳 문자 제거 후 재검사.
    예: '• date of birth: 7 april 1997' → inner='date of birth: ...' → 감지됨
    """
    # 앞 비알파벳 문자 제거 버전 (불릿·아이콘·탭 포함)
    _inner = re.sub(r'^[^\w]+', '', stripped_lower)
    # 원본 + 불릿 제거 버전 둘 다 검사
    _candidates = [stripped_lower]
    if _inner and _inner != stripped_lower:
        _candidates.append(_inner)
    for _t in _candidates:
        for label, needs_colon in PII_LINE_LABELS:
            if not _t.startswith(label):
                continue
            rest = _t[len(label):]
            if not rest:
                return True  # 라벨만 있는 줄
            first_char = rest.lstrip()[0] if rest.lstrip() else ""
            if needs_colon and first_char == ":":
                return True
            if not needs_colon:
                if first_char in (":", " "):
                    return True
    return False


def remove_pii(text: str, candidate: dict = None, skip_school_names: bool = False) -> tuple[str, list[str]]:
    """
    텍스트에서 PII 삭제.
    Returns: (cleaned_text, list[삭제 로그])
    """
    log = []

    # ── Pass 1: 줄 단위 PII 라벨 삭제 ──
    lines = text.split("\n")
    kept_lines = []
    for line in lines:
        stripped = line.strip().lower()
        if stripped and _should_skip_line(stripped):
            log.append(f"LINE_DEL: {line.strip()[:80]}")
            continue
        kept_lines.append(line)
    result = "\n".join(kept_lines)

    # ── Pass 1.5: 연락 초대 문장 삭제 ──
    # "please email me at ...", "find me on WhatsApp (+44 ...)", "I can be reached at ..."
    # → 해당 문장 전체 삭제 (커버레터/이력서 공통 적용)
    # 전략: 문단 분리 → 문장 분리 → 초대 문구 포함 문장 삭제 → 재조합
    _ci_paragraphs = result.split("\n\n")
    _ci_out_paragraphs = []
    for _para in _ci_paragraphs:
        # 줄바꿈 없는 단락은 문장 단위로 분리
        _sentences = re.split(r'(?<=[.!?])\s+', _para)
        _kept_sentences = []
        for _sent in _sentences:
            if _CONTACT_INVITE_RE.search(_sent) or _CONTACT_PLATFORM_RE.search(_sent):
                log.append(f"CONTACT_INVITE_DEL: {_sent.strip()[:80]}")
            else:
                _kept_sentences.append(_sent)
        # 문장 삭제로 단락이 비면 단락 자체 드롭
        _rebuilt = " ".join(_kept_sentences).strip()
        if _rebuilt:
            _ci_out_paragraphs.append(_rebuilt)
        # 빈 단락(원래 빈 줄 구분자)도 유지해 형태 보존
        elif _para.strip() == "":
            _ci_out_paragraphs.append(_para)
    result = "\n\n".join(_ci_out_paragraphs)

    # ── Pass 2: regex 기반 인라인 삭제 ──
    def _sub(pattern, repl, tag):
        nonlocal result
        def replacer(m):
            log.append(f"{tag}: {m.group()[:60]}")
            return repl
        result = pattern.sub(replacer, result)

    # 여권 먼저 (전화번호 regex와 숫자 충돌 방지)
    _sub(RE_PASSPORT, "", "PASSPORT")
    _sub(RE_EMAIL, "", "EMAIL")
    _sub(RE_LINKEDIN, "", "LINKEDIN")
    _sub(RE_KAKAO, "", "KAKAO")
    _sub(RE_SNS, "", "SNS")
    _sub(RE_URL, "", "URL")
    # 서술형 생년월일 ("20th of September 1996" 등)
    _sub(RE_DOB_SPELLED, "", "DOB_SPELLED")
    # 한국 전화 프리픽스 잔류 (+82 (0))
    _sub(RE_KR_PHONE_PREFIX, "", "KR_PHONE_PREFIX")

    # 전화번호 (7자리+ 숫자만)
    def phone_replacer(m):
        digits = re.sub(r"\D", "", m.group())
        if len(digits) < 7:
            return m.group()
        # 연도 범위 오탐 방지: 모든 숫자 그룹이 연도(1900-2099)면 건너뜀
        # "2022-2023", "2017\n2021-2022", "2015-2020\n2017" 등 커버
        digit_groups = re.findall(r'\d+', m.group())
        if digit_groups and all(re.match(r'^(?:19|20)\d{2}$', g) for g in digit_groups):
            return m.group()
        log.append(f"PHONE: {m.group()[:40]}")
        return ""
    result = RE_PHONE.sub(phone_replacer, result)

    # 한국 주소 → "Korea"
    def addr_replacer(m):
        log.append(f"KR_ADDR→Korea: {m.group()[:40]}")
        return "Korea"
    result = RE_KR_ADDRESS.sub(addr_replacer, result)

    # 한국 아파트/주거 주소 삭제
    _sub(RE_KR_RESIDENTIAL, "", "KR_RESID")

    # 영문 주소 삭제 (거리+도시+주/zip)
    _sub(RE_EN_ADDRESS, "", "EN_ADDR")

    # 한국 로마자 주소 삭제 (80-11, Jangji 1-gil 등)
    _sub(RE_KR_ROMANIZED_ADDR, "", "KR_ROMAN_ADDR")

    # ── Pass 2.4b: 한국 고용주 위치줄 통합 교체 ──
    # "[School Name], [Korean City]" → "English school, South Korea"
    # 예: "Berkeley Language School, Busan" → "English school, South Korea"
    #     "Hillside Collegiate, Daegu" → "English school, South Korea"
    #     "Busan Global Village, South Korea" → "English school, South Korea"
    # "korea", "south korea", "republic of korea" 포함 — 도시명 없이 국가명만 있는 줄도 처리
    _kr_city_pat = re.compile(
        r'\b(?:' + '|'.join(re.escape(c) for c in KR_KEYWORDS if len(c) > 3) + r')\b',
        re.IGNORECASE,
    )
    _has_kr_workplace = any(kw in result.lower() for kw in KR_WORKPLACE_KEYWORDS
                            if kw not in _GENERIC_TYPES_SET)
    _has_kr_school_named = bool(RE_SCHOOL_NAMED.search(result))
    # 줄 단위로 검사: 한국 도시 포함 + 학원/학교명 있는 짧은 줄 → 전체 교체
    if (_has_kr_workplace or _has_kr_school_named) and not skip_school_names:
        _new_lines = []
        for _rl in result.split("\n"):
            _rls = _rl.strip()
            if (len(_rls) > 0 and len(_rls) <= 80
                    and _kr_city_pat.search(_rls)
                    and not _contains_protected_university(_rls)  # ★ 정식 대학교 보호
                    and (any(kw in _rls.lower() for kw in KR_WORKPLACE_KEYWORDS
                             if kw not in _GENERIC_TYPES_SET)
                         or RE_SCHOOL_NAMED.search(_rls))):
                log.append(f"KR_EMPLOYER_LINE→English school, South Korea: {_rls[:60]}")
                # 원본 줄의 indentation 유지
                _indent = _rl[:len(_rl) - len(_rl.lstrip())]
                _new_lines.append(_indent + "English school, South Korea")
            else:
                _new_lines.append(_rl)
        result = "\n".join(_new_lines)

    # ── Pass 2.45: 인라인 진행형 + KR 업체 문장 통째 제거 (도시 제거 BEFORE) ──
    # "I am currently teaching at Little Fox in Gyeonggi-do," → 통째 삭제
    # ★ Pass 2.5 (city removal) 보다 먼저 실행 — 도시 부분까지 함께 캡처
    # alternation: 가장 구체적인 패턴(접미사 -do/-si/-gu 등) 먼저
    # Python re는 left-to-right first match — "Gyeonggi-do" 가 "Gyeonggi" 보다 먼저 매칭되도록
    _kr_loc_alt_e = (
        r'(?:[A-Z][a-z]+(?:-[a-z]+)*-(?:do|si|gu|dong|gun|myeon|eup|ri)|'
        r'South\s+Korea|Korea|'
        r'Seoul|Busan|Incheon|Daegu|Daejeon|Gwangju|Ulsan|Sejong|'
        r'Gyeonggi|Gangwon|Chungcheong|Jeolla|Gyeongsang|Jeju)'
    )
    for _wkp_e in KR_WORKPLACE_KEYWORDS:
        if _wkp_e in _GENERIC_TYPES_SET:
            continue
        _ph_e = re.compile(
            r'(?:I[\s\'’]?\s*(?:am|m)\s+)?'
            r'(?:currently\s+|now\s+|recently\s+|previously\s+|formerly\s+)?'
            r'(?:teaching|working|employed|teach|work|am\s+teaching|am\s+working)\s+'
            r'(?:at|for|in|with)\s+'
            r'\b' + re.escape(_wkp_e) + r'\b'
            r'(?:\s+(?:in|at)\s+' + _kr_loc_alt_e
            + r'(?:\s*,\s*' + _kr_loc_alt_e + r')?)?'
            r'\s*[,.;]?',
            re.IGNORECASE,
        )
        for _ph_m_e in _ph_e.finditer(result):
            log.append(f"KR_INLINE_PHRASE→DELETE: {_ph_m_e.group()[:80]}")
        result = _ph_e.sub('', result)

    # ── Pass 2.5: 한국 도시명 처리 ──
    # "Daegu, South Korea" → "South Korea"
    # "Seoul" → ""
    # ★ 정식 대학교명에 포함된 도시는 삭제 금지 (예: Seoul National University)
    for city in KR_KEYWORDS:
        if len(city) > 2 and city.lower() not in ("korea", "south korea", "rok"):
            # 단어 경계: 하이픈 포함 키워드("gangwon-do")는 \b 대신 앞뒤 공백/구두점 경계
            if "-" in city:
                pat = re.compile(r"(?<![A-Za-z])" + re.escape(city) + r"(?![A-Za-z])", re.IGNORECASE)
            else:
                pat = re.compile(r"\b" + re.escape(city) + r"\b", re.IGNORECASE)
            if not pat.search(result):
                continue
            # 도시명이 정식 대학교명 내에 포함된 경우 해당 줄 전체 보호
            # ★ _snap: sub() 호출 전 스냅샷 명시적 캡처 → 콜백 내부에서 인덱스 일관성 보장
            _snap = result
            def _kr_city_replacer(m, _city=city, _s=_snap):
                matched = m.group()
                start = m.start()
                end = m.end()
                line_start = _s.rfind('\n', 0, start) + 1
                line_end_idx = _s.find('\n', end)
                full_line = _s[line_start:] if line_end_idx == -1 else _s[line_start:line_end_idx]
                if _contains_protected_university(full_line):
                    return matched  # 대학교명 포함 줄 → 도시명 유지
                log.append(f"KR_CITY→REMOVE: {_city}")
                return ""
            result = pat.sub(_kr_city_replacer, result)
            result = re.sub(r",\s*,", ",", result)
            result = re.sub(r"^\s*,\s*", "", result, flags=re.MULTILINE)
            result = re.sub(r"\s*,\s*$", "", result, flags=re.MULTILINE)
    # 도시 제거 후 남은 "— , Text" 패턴 정리: "— , South Korea" → "— South Korea"
    result = re.sub(r'([—–])\s*,+\s*', r'\1 ', result)
    result = re.sub(r',+\s*([—–])', r' \1', result)

    # ── Pass 2.5b: 한국 우편번호 잔류 제거 ──
    # 도시명 삭제 후 남은 "10906", "03000" 등 5자리 우편번호 제거
    # 줄 전체가 숫자만 남거나, 쉼표/공백 뒤 5자리 숫자인 경우
    def _postal_replacer(m):
        log.append(f"KR_POSTAL→REMOVE: {m.group().strip()}")
        return ""
    result = RE_KR_POSTAL.sub(_postal_replacer, result)
    # 줄 단독 숫자(4~6자리) 잔류 처리: "  10906" 같이 줄이 숫자만 남은 경우
    result = re.sub(r'^\s*\d{4,6}\s*$', '', result, flags=re.MULTILINE)

    # ── Pass 2.6: 한국 업체명 처리 ──
    # Step 1: 브랜드 접두사 제거, 유형명 보존
    # "Broad Language School" → "Language School"
    def _school_prefix_replacer(m):
        canonical = " ".join(w.capitalize() for w in m.group(1).split())
        log.append(f"KR_SCHOOL_PREFIX→{canonical}: {m.group()[:60]}")
        return canonical
    result = _RE_KR_SCHOOL_PREFIX.sub(_school_prefix_replacer, result)

    # 잔류 정리: 인라인 제거 후 갭 정리
    result = re.sub(r'\s{2,}', ' ', result)
    result = re.sub(r'^\s*,\s*', '', result, flags=re.MULTILINE)

    # Step 3: 나머지 단독 KR 업체명 → 일반명 대체
    # "YBM ECC" → "English", "POLY" → "English"
    # 이미 유형명인 항목(language school 등)은 Step 1에서 처리됐으므로 건너뜀
    for workplace in KR_WORKPLACE_KEYWORDS:
        if workplace in _GENERIC_TYPES_SET:
            continue  # 유형명 그대로 보존
        pat = re.compile(r"\b" + re.escape(workplace) + r"\b", re.IGNORECASE)
        if pat.search(result):
            generic = _replace_workplace_generic(workplace)
            log.append(f"KR_WORKPLACE→{generic}: {workplace}")
            result = pat.sub(generic, result)

    # ── Pass 2.7: 학교/기관명 일반화 ──
    # skip_school_names=True이면 건너뜀 (EDUCATION 섹션 보호용)
    # "Sheffield University" → "University", "Lincoln High School" → "School"
    def school_replacer(m):
        matched = m.group()
        matched_lower = matched.lower()
        # 정식 인가 대학교 화이트리스트 → 삭제하지 않음
        for _pu in _PROTECTED_UNIVERSITIES:
            if _pu in matched_lower or matched_lower in _pu:
                return matched  # 보호: 원본 그대로 반환
        generic = _school_to_generic(matched)
        log.append(f"SCHOOL→{generic}: {matched[:60]}")
        return generic
    if not skip_school_names:
        result = RE_SCHOOL_NAMED.sub(school_replacer, result)
        result = RE_UNIV_OF.sub(school_replacer, result)

    # ── Pass 2.8: 외국 도시/주 + 나이 + 비자 삭제 ──
    # ★ FOREIGN_CITY / US_CITY_STATE: 정식 대학교명 포함 시 보호 (삭제 금지)
    def _city_replacer_safe(pattern, tag):
        nonlocal result
        def replacer(m):
            matched = m.group()
            # 도시명은 반드시 대문자로 시작해야 함 — re.I 오탐("cooperating teacher, Ms") 방지
            if not matched[0].isupper():
                return matched
            start = m.start()
            end = m.end()
            # 현재 줄 전체 추출 (짧은 매칭이라도 줄 전체로 대학교명 판별)
            line_start = result.rfind('\n', 0, start) + 1
            line_end_i = result.find('\n', end)
            full_line = result[line_start:] if line_end_i == -1 else result[line_start:line_end_i]
            # 매칭 자체 또는 해당 줄에 정식 대학교명 포함 → 보호
            if _contains_protected_university(matched) or _contains_protected_university(full_line):
                return matched
            log.append(f"{tag}: {matched[:60]}")
            return ""
        result = pattern.sub(replacer, result)

    _city_replacer_safe(RE_US_CITY_STATE, "US_CITY_STATE")
    _city_replacer_safe(RE_FOREIGN_CITY, "FOREIGN_CITY")
    _sub(RE_AGE_MENTION, "", "AGE_MENTION")
    _sub(RE_VISA_STATUS, "", "VISA_STATUS")

    # ── Pass 3: 후보자별 PII (DB에서 가져온 값) ──
    if candidate:
        name = candidate.get("full_name") or candidate.get("name", "")
        if name:
            # 이름 전체 삭제 (성+이름 모두 제거)
            words = name.strip().split()
            if len(words) >= 2:
                last_name = words[-1]
                first_name = " ".join(words[:-1])  # 성을 제외한 모든 단어

                # 방법 1: 전체 이름(풀네임) 삭제
                full_name_pat = re.compile(r"\b" + re.escape(name) + r"\b", re.IGNORECASE)
                matches = full_name_pat.findall(result)
                if matches:
                    log.append(f"FULLNAME→DELETE: {name}")
                    result = full_name_pat.sub("", result)

                # 방법 2: 성(Last name) 단독 삭제
                if len(last_name) >= 2:
                    last_name_pat = re.compile(r"\b" + re.escape(last_name) + r"\b", re.IGNORECASE)
                    matches = last_name_pat.findall(result)
                    if matches:
                        log.append(f"LASTNAME→DELETE: {last_name}")
                        result = last_name_pat.sub("", result)

                # 방법 3: 이름(First name) 단독 삭제 — 서명/커버레터 잔류 이름 제거
                if len(first_name) >= 2:
                    first_name_pat = re.compile(r"\b" + re.escape(first_name) + r"\b", re.IGNORECASE)
                    matches = first_name_pat.findall(result)
                    if matches:
                        log.append(f"FIRSTNAME→DELETE: {first_name}")
                        result = first_name_pat.sub("", result)

                    # 개별 이름 단어도 각각 삭제 (multi-word first name 대응)
                    _skip_words_lower = frozenset({"van", "von", "de", "del", "la", "le", "the", "and"})
                    for _fw in first_name.split():
                        if len(_fw) >= 4 and _fw.lower() not in _skip_words_lower:
                            fw_pat = re.compile(r"\b" + re.escape(_fw) + r"\b", re.IGNORECASE)
                            if fw_pat.search(result):
                                log.append(f"FIRSTWORD→DELETE: {_fw}")
                                result = fw_pat.sub("", result)

                # 앞뒤 공백 정리
                result = re.sub(r"\s+", " ", result)

        # DB에서 가져온 이메일/전화/카카오도 직접 삭제
        for field in ("email", "mobile_phone", "kakaotalk"):
            val = candidate.get(field, "")
            if val and len(val) > 2:
                escaped = re.escape(val)
                if re.search(escaped, result, re.I):
                    log.append(f"DB_{field.upper()}: {val[:30]}")
                    result = re.sub(escaped, "", result, flags=re.I)

    # ── Pass 4: 위치/직장 라벨 값 처리 ──
    # 모든 줄을 검사하여 라벨을 포함하는 줄을 처리
    lines_final = result.split("\n")
    final_lines = []

    for line in lines_final:
        stripped = line.strip()
        modified = line
        found = False

        # LOCATION_LABELS 검사
        for label in LOCATION_LABELS:
            # "Location: Seoul" 또는 "location : seoul" 등 다양한 형식 대응
            label_pat = re.compile(
                r"(.*?)\b" + re.escape(label) + r"\s*:?\s*(.+)",
                re.IGNORECASE
            )
            m = label_pat.match(stripped)
            if m:
                prefix, value = m.group(1), m.group(2).strip()
                if _is_korea(value):
                    # 도시명 제거: "Daegu, South Korea" → "South Korea"
                    # KR_KEYWORDS의 도시명들을 모두 제거
                    cleaned_value = value
                    for city in KR_KEYWORDS:
                        if len(city) > 2:
                            city_pat = re.compile(r"\b" + re.escape(city) + r"\s*,?\s*", re.IGNORECASE)
                            cleaned_value = city_pat.sub("", cleaned_value).strip()

                    # 결과가 비면 "South Korea"로 설정
                    if not cleaned_value or cleaned_value.lower() == "south korea":
                        cleaned_value = "South Korea"

                    log.append(f"LOCATION→{cleaned_value}: {value[:40]}")
                    # 원래 라인 구조 유지
                    indent = len(line) - len(line.lstrip())
                    modified = " " * indent + stripped.replace(value, cleaned_value)
                    found = True
                    break

        if not found:
            # WORKPLACE_LABELS 검사
            for label in WORKPLACE_LABELS:
                label_pat = re.compile(
                    r"(.*?)\b" + re.escape(label) + r"\s*:?\s*(.+)",
                    re.IGNORECASE
                )
                m = label_pat.match(stripped)
                if m:
                    prefix, value = m.group(1), m.group(2).strip()
                    generic_value = _replace_workplace_generic(value)
                    log.append(f"WORKPLACE→{generic_value}: {value[:40]}")
                    indent = len(line) - len(line.lstrip())
                    modified = " " * indent + stripped.replace(value, generic_value)
                    found = True
                    break

        final_lines.append(modified)

    result = "\n".join(final_lines)

    # 빈줄 정리
    result = re.sub(r"\n{3,}", "\n\n", result)
    # 빈 라벨줄 정리 (라벨: 만 남은 줄)
    result = re.sub(r"^[ \t]*\w[\w\s]*:\s*$", "", result, flags=re.MULTILINE)
    result = re.sub(r"\n{3,}", "\n\n", result)

    return result, log


# ══════════════════════════════════════════════════════
#  4. DOCX 처리 (서식 보존)
# ══════════════════════════════════════════════════════

def _replace_in_runs(paragraph, cleaned_text):
    """
    단락의 run들에서 텍스트를 교체하되 서식을 최대한 보존.
    전략: 각 run의 텍스트 내에서 PII 부분만 교체.
    """
    runs = paragraph.runs
    if not runs:
        return

    original_full = "".join(r.text for r in runs)
    if not original_full.strip():
        return

    # 삭제된 텍스트가 없으면 스킵
    if original_full == cleaned_text:
        return

    # run이 1개면 단순 교체 (서식 유지)
    if len(runs) == 1:
        runs[0].text = cleaned_text
        return

    # run이 여러개: cleaned_text를 run 경계에 맞춰 분배
    # 전략: 각 run의 원래 길이 비율로 배분
    # 단, PII가 run 경계를 걸칠 수 있으므로 최선의 근사치 사용
    # cleaned가 비면 모든 run을 비움
    if not cleaned_text.strip():
        for r in runs:
            r.text = ""
        return

    # 원문과 cleaned의 공통 접두사/접미사 기반 매칭
    # 실용적 접근: 첫 run에 전체 cleaned, 나머지 비움
    # (완벽한 run-level 매칭은 diff 알고리즘이 필요하여 실무 대비 과도)
    #
    # 개선: 각 run의 텍스트가 cleaned에 그대로 존재하면 유지,
    #       아니면 cleaned에서 해당 위치 텍스트 추출
    pos = 0
    remaining = cleaned_text
    for i, run in enumerate(runs):
        orig_len = len(run.text)
        if i == len(runs) - 1:
            # 마지막 run에 나머지 전부
            run.text = remaining
        else:
            # 원래 run의 텍스트가 remaining 시작부에 있으면 유지
            if remaining.startswith(run.text):
                remaining = remaining[orig_len:]
            else:
                # PII가 이 run에 있었음 → 이 run 비우고 다음으로
                run.text = ""


def _extract_names_from_filename(filepath: Path) -> list[str]:
    """파일명에서 사람 이름 추출. 'Resume-DoniaBland2026 (1) - Donia Bland.docx' → ['Donia Bland', ...]"""
    stem = filepath.stem
    # " - Name" 패턴 (가장 흔함)
    parts = re.split(r"\s*-\s*", stem)
    names = []
    for part in parts:
        # 숫자만/키워드만 제외
        clean = re.sub(r"\(\d+\)", "", part).strip()
        clean = re.sub(r"\d{4,}", "", clean).strip()
        if clean.lower() in ("resume", "cv", "cover letter", "coverletter", ""):
            continue
        # 이름이 아닌 단어 제거 (Teaching, CV, English 등)
        words = [w for w in clean.split()
                 if w and w.lower() not in _NON_NAME_WORDS and len(w) >= 2]
        if len(words) >= 2 and all(w[0].isupper() for w in words if w):
            name_clean = " ".join(words)
            names.append(name_clean)
            # 개별 단어도 추가 (3글자 이상, non-name 제외)
            for w in words:
                if len(w) >= 3 and w.lower() not in _NON_NAME_WORDS:
                    names.append(w)
    return names


def _find_photo_for_candidate(filepath: Path) -> Path | None:
    """incoming 폴더에서 같이 넣은 사진 파일 찾기 (JPG/PNG)"""
    incoming = filepath.parent
    photo_exts = {".jpg", ".jpeg", ".png"}
    # 1) 같은 번호로 시작하는 사진
    stem_numbers = re.findall(r"(\d{3,5})", filepath.stem)
    for f in incoming.iterdir():
        if f.suffix.lower() in photo_exts:
            for num in stem_numbers:
                if num in f.stem:
                    return f
    # 2) incoming에 사진이 1개만 있으면 그거
    photos = [f for f in incoming.iterdir() if f.suffix.lower() in photo_exts]
    if len(photos) == 1:
        return photos[0]
    # 3) processed_docs 폴더에서도 찾기
    proc_parent = incoming.parent
    for subdir in [proc_parent / "incoming", proc_parent]:
        if not subdir.exists():
            continue
        for f in subdir.iterdir():
            if f.suffix.lower() in photo_exts:
                for num in stem_numbers:
                    if num in f.stem:
                        return f
    return None


def process_docx(filepath: Path, brj_number: int, candidate: dict = None,
                 dry: bool = False, photo_path: Path = None):
    """DOCX 처리: PII 삭제 + 번호/사진 삽입. Returns (doc, log_entries)"""
    import docx
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document(str(filepath))
    all_logs = []

    # 파일명에서 이름 추출 (DB 불일치 대비)
    filename_names = _extract_names_from_filename(filepath)

    # Reference 섹션 헤더 감지용
    _REF_HEADERS = frozenset({
        "references", "reference", "professional references",
        "references & recommendations", "references and recommendations",
        "referee", "referees", "character references",
    })
    _ref_active = [False]  # mutable flag (closure 공유)

    def _process_paragraphs(paragraphs, section_name="body"):
        _in_edu = [False]  # EDUCATION 섹션 플래그 (클로저용 리스트)

        for para in paragraphs:
            original = para.text
            if not original.strip():
                continue

            stripped_lower = original.strip().lower()

            # EDUCATION 섹션 추적 — 본인 학력은 학교명 보호
            if _EDU_SECTION_RE.match(stripped_lower):
                _in_edu[0] = True
            elif _in_edu[0] and _OTHER_SECTION_RE.match(stripped_lower):
                _in_edu[0] = False

            # Reference 섹션 헤더 감지 → 이후 전체 삭제
            if stripped_lower in _REF_HEADERS:
                _ref_active[0] = True
                all_logs.append(f"[{section_name}] REF_SECTION_START")
                if not dry:
                    for run in para.runs:
                        run.text = ""
                continue

            # 다음 주요 섹션 헤더가 나오면 Reference 종료
            if _ref_active[0] and stripped_lower in _RESUME_SECTION_HEADERS:
                _ref_active[0] = False

            # Reference 섹션 내 내용 전체 삭제
            if _ref_active[0]:
                all_logs.append(f"[{section_name}] REF_DEL: {original.strip()[:60]}")
                if not dry:
                    for run in para.runs:
                        run.text = ""
                continue

            # ── 하이퍼링크 내 이메일/전화번호 제거 (para.text에 포함 안 됨) ──
            for _child in list(para._element):
                _ctag = _child.tag.split("}")[-1] if "}" in _child.tag else _child.tag
                if _ctag == "hyperlink":
                    _link_parts = []
                    for _r in _child:
                        _rtag = _r.tag.split("}")[-1] if "}" in _r.tag else _r.tag
                        if _rtag == "r":
                            for _t in _r:
                                _ttag = _t.tag.split("}")[-1] if "}" in _t.tag else _t.tag
                                if _ttag == "t" and _t.text:
                                    _link_parts.append(_t.text)
                    _link_text = "".join(_link_parts)
                    if RE_EMAIL.search(_link_text) or RE_KR_PHONE_PREFIX.search(_link_text):
                        all_logs.append(f"[{section_name}] HYPERLINK_EMAIL: {_link_text[:60]}")
                        if not dry:
                            para._element.remove(_child)

            cleaned, log = remove_pii(original, candidate, skip_school_names=_in_edu[0])
            # 파일명에서 추출한 이름도 삭제
            for fn_name in filename_names:
                pat = re.compile(re.escape(fn_name), re.I)
                if pat.search(cleaned):
                    log.append(f"FILENAME_NAME: {fn_name}")
                    cleaned = pat.sub("", cleaned)
            if cleaned != original:
                all_logs.extend([f"[{section_name}] {l}" for l in log])
                if not dry:
                    # South Korea → 볼드 처리
                    if "South Korea" in cleaned:
                        _replace_with_bold_korea(para, cleaned)
                    else:
                        _replace_in_runs(para, cleaned)
                    # 하이퍼링크/필드코드 등 runs에 안 잡히는 요소 강제 클리어
                    if not cleaned.strip() and para.text.strip():
                        from docx.oxml.ns import qn
                        for child in list(para._element):
                            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                            if tag in ("hyperlink", "r", "fldSimple", "smartTag"):
                                para._element.remove(child)

    # 헤더/푸터 — 전체 삭제 (이름/주소/연락처 모두 비움)
    def _clean_header_footer(paragraphs, section_name="header"):
        for para in paragraphs:
            original = para.text
            if not original.strip():
                continue
            all_logs.append(f"[{section_name}] HDR_CLEAR: {original.strip()[:60]}")
            if not dry:
                for run in para.runs:
                    run.text = ""

    # ── 사전 처리: PII 삭제 전 국적 추출 → DB 업데이트 ──
    if not dry and candidate:
        _full_text = "\n".join(p.text for p in doc.paragraphs)
        _nat_extracted = _extract_nationality_from_text(_full_text)
        if _nat_extracted:
            _current_nat = (candidate.get("nationality") or "").strip()
            if _current_nat in ("", "미상", "unknown", "Unknown"):
                candidate["nationality"] = _nat_extracted
                _update_candidate_nationality(
                    candidate.get("sheet_number", 0), _nat_extracted
                )
                all_logs.append(f"[pre] NAT_EXTRACTED: {_nat_extracted}")

    # 본문
    _process_paragraphs(doc.paragraphs, "body")

    # 테이블
    for ti, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs, f"table{ti}")

    # 헤더/푸터 — 전체 삭제
    for si, section in enumerate(doc.sections):
        for hf_name, hf in [("header", section.header), ("footer", section.footer)]:
            if hf and hf.paragraphs:
                _clean_header_footer(hf.paragraphs, f"{hf_name}{si}")
        try:
            fph = section.first_page_header
            if fph and fph.paragraphs:
                _clean_header_footer(fph.paragraphs, f"first_header{si}")
        except Exception:
            pass
        try:
            fpf = section.first_page_footer
            if fpf and fpf.paragraphs:
                _clean_header_footer(fpf.paragraphs, f"first_footer{si}")
        except Exception:
            pass

    # 문서 메타데이터 클리어 (이름/제목 제거)
    try:
        cp = doc.core_properties
        cp.title = ""
        cp.author = ""
        cp.subject = ""
        cp.keywords = ""
        cp.comments = ""
        cp.last_modified_by = ""
    except Exception:
        pass

    # 번호 + 사진 삽입 (최상단)
    if not dry:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        _is_cl = _is_cover_letter(filepath)
        body = doc.element.body

        if not _is_cl:
            # 기존 본문 이미지 모두 제거 (프로필 사진 중복 방지)
            for drawing in body.findall('.//' + qn('w:drawing')):
                parent = drawing.getparent()
                if parent is not None:
                    parent.remove(drawing)
                    all_logs.append("[photo] 기존 이미지 제거")
            for pict in body.findall('.//' + qn('w:pict')):
                parent = pict.getparent()
                if parent is not None:
                    parent.remove(pict)

            photo = photo_path or _find_photo_for_candidate(filepath)

            # 1행 2열 테이블: [번호 | 사진]
            tbl = doc.add_table(rows=1, cols=2)
            tbl.autofit = True
            # 테두리 없음
            tbl_pr = tbl._element.tblPr
            borders = tbl_pr.find(qn("w:tblBorders"))
            if borders is not None:
                tbl_pr.remove(borders)

            # 왼쪽: 번호 (크고 굵게)
            left_cell = tbl.cell(0, 0)
            left_cell.text = ""
            p_num = left_cell.paragraphs[0]
            run_num = p_num.add_run(str(brj_number))
            run_num.bold = True
            run_num.font.size = Pt(28)
            run_num.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            # 오른쪽: 사진
            right_cell = tbl.cell(0, 1)
            right_cell.text = ""
            p_photo = right_cell.paragraphs[0]
            p_photo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if photo and photo.exists():
                run_photo = p_photo.add_run()
                run_photo.add_picture(str(photo), height=Cm(3.5))
                all_logs.append(f"[photo] 삽입: {photo.name}")
            else:
                all_logs.append("[photo] 사진 없음 (스킵)")

            # body의 맨 앞 자식 앞에 삽입 (테이블/단락 무관하게 최상단)
            first_child = body[0] if len(body) > 0 else None
            if first_child is not None:
                first_child.addprevious(tbl._element)
            else:
                body.append(tbl._element)
        else:
            all_logs.append("[cover_letter] 번호/사진 삽입 스킵")

    return doc, all_logs


def _replace_with_bold_korea(para, cleaned_text: str):
    """South Korea를 볼드로 처리하여 단락에 삽입"""
    from docx.shared import Pt, RGBColor
    # 기존 runs의 서식 복사용 (첫 run의 font 정보 보존)
    base_font_size = None
    if para.runs:
        base_font_size = para.runs[0].font.size
    # 기존 runs 텍스트 비우기
    for run in para.runs:
        run.text = ""
    # South Korea 기준으로 분할하여 새 run 추가
    parts = cleaned_text.split("South Korea")
    for i, part in enumerate(parts):
        if part:
            run = para.add_run(part)
            if base_font_size:
                run.font.size = base_font_size
        if i < len(parts) - 1:
            kr_run = para.add_run("South Korea")
            kr_run.bold = True
            kr_run.font.size = base_font_size or Pt(11)
            kr_run.font.color.rgb = RGBColor(0, 0, 0)


# ══════════════════════════════════════════════════════
#  5. PDF 처리 (인-플레이스 redaction)
# ══════════════════════════════════════════════════════

def process_pdf(filepath: Path, brj_number: int, candidate: dict = None,
                dry: bool = False, photo_path: Path = None):
    """
    PDF 처리: PyMuPDF redaction으로 PII를 흰색 사각형으로 덮기.
    첫 페이지 상단 영역은 화이트아웃 후 번호+사진 삽입.
    Returns (doc, log_entries)
    """
    import fitz

    doc = fitz.open(str(filepath))
    all_logs = []

    # ── 스캔 PDF 감지 + OCR 기반 redaction (텍스트 레이어 없는 이미지 PDF) ──
    _scanned_face_covers = {}  # page_idx -> list of fitz.Rect (face 영역)
    if not dry:
        _scanned_pages = [pn for pn, pg in enumerate(doc) if _is_scanned_page(pg)]
        if _scanned_pages:
            all_logs.append(f"[ocr] 스캔 페이지 감지: {_scanned_pages}")
            _ocr_patterns = _build_search_patterns(candidate)
            for _pn in _scanned_pages:
                # 1) 얼굴 탐색 → 원본 사진 영역 흰박스 (번호/새사진과 겹침 방지)
                _covers = _detect_and_cover_existing_photo(doc[_pn], _pn, all_logs)
                if _covers:
                    _scanned_face_covers[_pn] = _covers
                # 2) OCR 기반 텍스트 PII redaction
                _ocr_redact_scanned_page(
                    doc[_pn], _pn, candidate, _ocr_patterns, all_logs
                )

    # ── 사전 처리: PII 삭제 전 국적 추출 → DB 업데이트 ──
    if not dry and candidate:
        _pdf_full_text = "".join(page.get_text() for page in doc)
        _nat_extracted = _extract_nationality_from_text(_pdf_full_text)
        if _nat_extracted:
            _current_nat = (candidate.get("nationality") or "").strip()
            if _current_nat in ("", "미상", "unknown", "Unknown"):
                candidate["nationality"] = _nat_extracted
                _update_candidate_nationality(
                    candidate.get("sheet_number", 0), _nat_extracted
                )
                all_logs.append(f"[pre] NAT_EXTRACTED: {_nat_extracted}")

    # 모든 PII 패턴을 수집
    pii_patterns = _build_search_patterns(candidate)

    # 파일명에서 이름 추출 (DB 불일치 대비)
    filename_names = _extract_names_from_filename(filepath)

    # PDF 첫 페이지 텍스트에서 이름 추출 (상단 영역)
    _pdf_extracted_names = []
    # 직책/타이틀 줄: 풀텍스트만 blank, 단어 분리 금지 (educator/school 등 오탐 방지)
    _header_title_texts: set = set()
    if len(doc) > 0:
        first_text = doc[0].get_text()
        first_lines = [l.strip() for l in first_text.split("\n") if l.strip()]

        # ── 스페이스 장식체 이름 감지: "A RY A M OH AN KA" → "ARYA MOHANKA" ──
        # 패턴: 공백으로 분리된 1~3자 대문자 토큰이 4개 이상 연속
        for _fl in first_lines[:6]:
            _toks = _fl.split()
            if (len(_toks) >= 4
                    and all(1 <= len(t) <= 3 and t.isupper() for t in _toks)):
                _collapsed = re.sub(r'\s+', '', _fl)  # "ARYAMOHANKA"
                # 2~4단어로 분리 시도 (단어 경계를 추정하기 어려우므로 전체 collapsed 삭제)
                _pdf_extracted_names.append(_fl)        # 원문 그대로 ("A RY A M OH AN KA")
                _pdf_extracted_names.append(_collapsed)  # 합친 것 ("ARYAMOHANKA")
                all_logs.append(f"[pre] SPACED_NAME: {_fl!r} → {_collapsed!r}")

        for line in first_lines[:6]:  # 상위 6줄 (이전 5줄에서 확장)
            # ── suffix 제거: "MiCaila Rae Holland, M.A.T." → "MiCaila Rae Holland" ──
            _line_clean = re.sub(
                r',?\s+(?:[A-Z]\.){1,4}[A-Z]?\.?\s*$'   # ", M.A.T." / "Ph.D."
                r'|,?\s+(?:Jr|Sr|II|III|IV|Esq)\.?\s*$',  # ", Jr." etc.
                '', line).strip()
            # 이름 패턴: 2~5단어, 각 첫글자 대문자, 짧은 줄
            if len(_line_clean) < 60 and not _should_skip_line(_line_clean.lower()):
                # 섹션 헤더 제외 (PROFESSIONAL SUMMARY 등)
                if _line_clean.strip().lower() in _RESUME_SECTION_HEADERS:
                    continue
                # ALL-CAPS 2+단어 → 섹션 제목으로 간주 (이름은 보통 Mixed Case)
                # 단, 2단어이고 각 단어가 짧으면(≤8자) 성씨일 가능성 → 개별 추가
                if _line_clean.isupper() and len(_line_clean.split()) >= 2:
                    _caps_words = _line_clean.split()
                    if len(_caps_words) == 2 and all(len(w) <= 8 for w in _caps_words):
                        _pdf_extracted_names.append(_line_clean)
                        for _cw in _caps_words:
                            if len(_cw) >= 3:
                                _pdf_extracted_names.append(_cw)
                    continue
                words = _line_clean.split()
                # 주소/일반 단어 포함 줄은 이름이 아님
                if any(w.lower() in _NON_NAME_WORDS for w in words):
                    continue
                # 라벨 단어("Endorsements:", "Phone:", "Date:") 포함 줄은 이름이 아님
                if any(w.endswith(':') for w in words):
                    continue
                # 커버레터 인사말("Dear Hiring Manager,") 패턴 → 이름이 아님
                if words and words[0].lower() in ("dear", "to", "from", "re", "subject"):
                    continue
                if 2 <= len(words) <= 5 and all(
                    (w[0].isupper() if w else False) or w in ("de", "van", "von", "del", "K.", "K")
                    for w in words if w
                ):
                    _pdf_extracted_names.append(_line_clean)
                    # suffix 제거 전 원문도 추가 (검색 일치 위해)
                    if _line_clean != line:
                        _pdf_extracted_names.append(line)
                    for w in words:
                        # 쉼표 등 후처리
                        w_clean = w.rstrip(",.;:")
                        if len(w_clean) >= 3 and w_clean[0].isupper() and w_clean.lower() not in _NON_NAME_WORDS:
                            _pdf_extracted_names.append(w_clean)

        # ── 인접 두 줄이 각각 단일 단어 이름인 경우 합쳐서 처리 ──
        # 예: "DéyJah" / "Burrows" → "DéyJah Burrows"
        # 이름 탐지 후 바로 다음 줄이 직책/타이틀 설명이면 함께 blank
        # 예: "DéyJah" / "Burrows" / "American Kindergarten & Lower Primary Educator" → 타이틀줄도 삭제
        for _pi in range(min(12, len(first_lines) - 1)):
            _l1, _l2 = first_lines[_pi], first_lines[_pi + 1]
            _w1, _w2 = _l1.split(), _l2.split()
            if (len(_w1) == 1 and len(_w2) == 1
                    and 2 <= len(_l1) <= 20 and 2 <= len(_l2) <= 20
                    and _l1[0].isupper() and _l2[0].isupper()
                    and not _should_skip_line(_l1.lower())
                    and not _should_skip_line(_l2.lower())
                    and not any(w in _l1.lower() for w in _NON_NAME_WORDS)
                    and not any(w in _l2.lower() for w in _NON_NAME_WORDS)):
                _combined = f"{_l1} {_l2}"
                _pdf_extracted_names.append(_combined)
                _pdf_extracted_names.append(_l1)
                _pdf_extracted_names.append(_l2)
                all_logs.append(f"[pre] SPLIT_NAME: {_l1!r} + {_l2!r}")
                # ── 이름 직후 직책/타이틀 설명줄 blank ──
                # 이름 바로 뒤에 오는 짧은 설명줄(≤8단어, 섹션헤더 아님)도 삭제
                # "American Kindergarten & Lower Primary Educator" 같은 직책줄 제거
                for _adj in range(_pi + 2, min(_pi + 4, len(first_lines))):
                    _lt = first_lines[_adj].strip()
                    if not _lt:
                        continue
                    _lt_lower = _lt.lower()
                    if _lt_lower in _RESUME_SECTION_HEADERS:
                        break  # 섹션 헤더 → 직책줄 종료
                    _lt_words = _lt.split()
                    if (2 <= len(_lt_words) <= 8
                            and not _should_skip_line(_lt_lower)):
                        # 직책 타이틀: 단어 분리 없이 풀텍스트만 blank
                        # (_pdf_extracted_names에 추가하면 "Educator", "Primary" 등 단독 단어가 본문 오탐)
                        _header_title_texts.add(_lt)
                        all_logs.append(f"[pre] NAMETITLE_BLANK: {_lt!r}")
                    else:
                        break

    # ── 헤더 페이지 자동 감지 (More About Me 인트로 페이지 스킵) ──
    # "More About Me" / "About Me" 같은 자기소개 페이지는 Page 0이지만
    # 실제 이력서 헤더(이름+연락처)는 다른 페이지일 수 있음
    _header_page_num = 0
    _MORE_ABOUT_RE = re.compile(
        r'^\s*(more\s+about|about\s+me|personal\s+profile|personal\s+statement'
        r'|my\s+profile|profile\s+summary)',
        re.I,
    )
    for _hpi, _hpg in enumerate(doc):
        _top_txt = _hpg.get_text("text", clip=fitz.Rect(0, 0, _hpg.rect.width, 280))
        _top_lower = _top_txt.lower().strip()
        if _MORE_ABOUT_RE.match(_top_lower):
            continue  # About Me 페이지 스킵
        # 이메일/전화 패턴이 상단에 있으면 이력서 헤더 페이지
        if re.search(r'@[a-zA-Z]|(?:\+?\d{1,3}[\s\-]?\(?\d{2,4}\)?[\s\-]?\d{3})', _top_txt):
            _header_page_num = _hpi
            break
        # 컬러 헤더 사각형(베이지 등)이 상단에 있으면 이력서 헤더 페이지
        for _drw in _hpg.get_drawings():
            _dfill = _drw.get("fill")
            _dr = _drw.get("rect", fitz.Rect())
            if (_dfill and _dfill not in ((1, 1, 1), (0, 0, 0))
                    and _dr.y0 < 30 and _dr.width > _hpg.rect.width * 0.4):
                _header_page_num = _hpi
                break
        else:
            continue
        break

    for page_num, page in enumerate(doc):
        text = page.get_text()

        # 줄 단위 PII 라벨 검사 → redact 대상 텍스트 수집
        redact_texts = set()
        # 헤더 직책 타이틀 전체 blank (단어 분리 없음 — "educator" 등 본문 오탐 방지)
        redact_texts.update(_header_title_texts)
        replacement_redacts = []  # (원문, 대체텍스트) 쌍
        job_title_replaces: list = []  # 직업 타이틀줄 전체교체: (원문, 클린텍스트) — 갭 없이 제거
        _processed_jt_bboxes: list = []  # JOB_LINE_REINSERT_PRE로 처리된 bbox 추적

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            # PII 라벨 줄 → 전체 줄 redact
            if _should_skip_line(stripped.lower()):
                redact_texts.add(stripped)
                all_logs.append(f"[page{page_num}] LINE_DEL: {stripped[:60]}")
                continue

        # regex 패턴 매칭
        for pattern, tag in pii_patterns:
            for m in pattern.finditer(text):
                matched = m.group().strip()
                if not matched or len(matched) <= 1:
                    continue
                # 전화번호 오탐 방지: 7자리 미만 또는 연도 범위 건너뜀
                if tag == "PHONE":
                    if len(re.sub(r'\D', '', matched)) < 7:
                        continue
                    _dg = re.findall(r'\d+', matched)
                    if _dg and all(re.match(r'^(?:19|20)\d{2}$', g) for g in _dg):
                        continue
                redact_texts.add(matched)
                all_logs.append(f"[page{page_num}] {tag}: {matched[:60]}")

        # ── 학교명 redact (비보호 학교만) ──
        # 고등학교/중학교/초등학교는 항상 redact (개인 식별 가능)
        # 대학교는 _PROTECTED_UNIVERSITIES 체크 후 비보호 시 redact
        _SECONDARY_SUFFIXES = frozenset({
            "high school", "secondary school", "grammar school",
            "boarding school", "prep school", "junior school",
            "elementary school", "primary school", "sixth form college",
        })
        for _sline in text.split("\n"):
            _sstripped = _sline.strip()
            if not _sstripped:
                continue
            for _sm in RE_SCHOOL_NAMED.finditer(_sstripped):
                _school = _sm.group().strip()
                _school_lower = _school.lower()
                _is_secondary = any(_school_lower.endswith(sfx) for sfx in _SECONDARY_SUFFIXES)
                _is_protected = _school_lower in _PROTECTED_UNIVERSITIES
                if _is_secondary or not _is_protected:
                    redact_texts.add(_school)
                    all_logs.append(f"[page{page_num}] SCHOOL: {_school[:60]}")
            # RE_UNIV_OF: "University of Sheffield" style (short lines only — 오탐 방지)
            if len(_sstripped) <= 80:
                for _um in RE_UNIV_OF.finditer(_sstripped):
                    _univ = _um.group().strip()
                    if _univ.lower() not in _PROTECTED_UNIVERSITIES:
                        redact_texts.add(_univ)
                        all_logs.append(f"[page{page_num}] UNIV_UNPROTECTED: {_univ[:60]}")

        # ── 후보자 이름: PDF는 전체 blank (다중 렌더링 레이어 완전 제거) ──
        if candidate and (candidate.get("full_name") or candidate.get("name")):
            name = (candidate.get("full_name") or candidate.get("name", "")).strip()
            words = name.split()
            if len(words) >= 2:
                last_name = words[-1]
                first_name = " ".join(words[:-1])
                text_lower = text.lower()
                # 풀네임 전체 blank (replacement 사용 금지 — 다중 레이어 잔류 방지)
                if name.lower() in text_lower:
                    redact_texts.add(name)
                    all_logs.append(f"[page{page_num}] FULLNAME_BLANK: {name}")
                # 이름 부분(first_name)도 blank — "MiCaila Rae" 단독 블록 제거
                if first_name and first_name.lower() in text_lower:
                    redact_texts.add(first_name)
                    all_logs.append(f"[page{page_num}] FIRSTNAME_BLANK: {first_name}")
                # 성 단독 삭제 (풀네임 미검출 시에도 성은 반드시 제거)
                if len(last_name) >= 2 and last_name.lower() in text_lower:
                    redact_texts.add(last_name)
                    all_logs.append(f"[page{page_num}] LASTNAME: {last_name}")
                # ── 개별 이름 단어 blank — 부분 처리된 파일에서 잔류 이름 제거 ──
                # "My name is Katerina" 같은 본문 참조도 제거
                _skip_words = frozenset({"van", "von", "de", "del", "la", "le", "the", "and"})
                for _nw in words:
                    _nwc = _nw.rstrip(",.;:")
                    if (len(_nwc) >= 4
                            and _nwc.lower() not in _skip_words
                            and _nwc.lower() in text_lower):
                        redact_texts.add(_nwc)
                        all_logs.append(f"[page{page_num}] NAMEWORD: {_nwc}")
                # ── CamelCase 성 분리: "VanZyl" → "Van Zyl" / "VAN ZYL" 대응 ──
                _cc_parts = re.findall(r'[A-Z][a-z]+|[A-Z]{2,}', last_name)
                if len(_cc_parts) > 1:
                    _spaced_ln = " ".join(_cc_parts)
                    if _spaced_ln.lower() in text_lower:
                        redact_texts.add(_spaced_ln)
                        all_logs.append(f"[page{page_num}] LASTNAME_SPACE: {_spaced_ln}")
                    for _cp in _cc_parts:
                        if len(_cp) >= 3 and _cp.lower() in text_lower:
                            redact_texts.add(_cp)
                            all_logs.append(f"[page{page_num}] LASTNAME_PART: {_cp}")
                # PDF가 이름을 \n으로 분리 저장하는 경우 대비
                if name.lower() not in text_lower:
                    for _line in text.split("\n"):
                        _ls = _line.strip()
                        if name.lower() in _ls.lower():
                            replacement_redacts.append((_ls if len(_ls) < len(name) + 10 else name, first_name))
                            all_logs.append(f"[page{page_num}] FULLNAME_LINE→{first_name}: {_ls[:50]}")

        # 파일명 기반 이름: 성만 redact (이름 유지)
        # 단일 단어 이름 중 다중 단어 이름의 첫 번째 단어(이름)는 redact 제외
        _fn_first_names = {
            fn_n.split()[0].lower()
            for fn_n in filename_names
            if len(fn_n.split()) >= 2
        }
        for fn_name in filename_names:
            fn_words = fn_name.split()
            if len(fn_words) >= 2:
                fn_last = fn_words[-1]
                fn_first = " ".join(fn_words[:-1])
                if fn_name.lower() in text.lower():
                    replacement_redacts.append((fn_name, fn_first))
                    all_logs.append(f"[page{page_num}] FILENAME_NAME→{fn_first}: {fn_name}")
                # 성(last name): 최소 3자 이상이어야 redact — "T", "K" 등 이니셜 제외
                if len(fn_last) >= 3 and fn_last.lower() in text.lower():
                    redact_texts.add(fn_last)
            elif (fn_name.lower() not in _fn_first_names
                  and len(fn_name) >= 3  # 이니셜 단독 방지
                  and fn_name.lower() in text.lower()):
                # 단일 단어가 이름(first name)이면 redact 금지 — 성(last name)만 허용
                redact_texts.add(fn_name)
                all_logs.append(f"[page{page_num}] FILENAME_NAME: {fn_name}")

        # PDF 텍스트에서 추출한 이름: 성만 redact
        for pdf_name in _pdf_extracted_names:
            pdf_words = pdf_name.split()
            if len(pdf_words) >= 2:
                pdf_last = pdf_words[-1]
                pdf_first = " ".join(pdf_words[:-1])
                # 풀네임·이름부분 모두 blank (replacement 금지 — 다중 레이어 잔류 방지)
                if pdf_name.lower() in text.lower():
                    redact_texts.add(pdf_name)
                    all_logs.append(f"[page{page_num}] PDF_NAME_BLANK: {pdf_name}")
                if pdf_first and pdf_first.lower() in text.lower():
                    redact_texts.add(pdf_first)
                    all_logs.append(f"[page{page_num}] PDF_FIRSTNAME_BLANK: {pdf_first}")
                if pdf_last.lower() in text.lower():
                    redact_texts.add(pdf_last)
            elif len(pdf_name) >= 3 and pdf_name.lower() in text.lower():
                redact_texts.add(pdf_name)
                all_logs.append(f"[page{page_num}] PDF_NAME: {pdf_name}")

        # ── 한국 학원명/도시명 redact ──
        page_lower = text.lower()
        # page_has_korea: KR_KEYWORDS 중 하나라도 있으면 True
        # (기존 좁은 regex 대신 전체 키워드 집합으로 확장 → Gyeonggi 등 오탐 방지)
        page_has_korea = any(kw in page_lower for kw in KR_KEYWORDS)

        # 인라인 진행형 + KR 업체 문장 통째 redact (커버레터/본문 공통)
        # 예: "I am currently teaching at Little Fox in Gyeonggi-do," → 통째 삭제
        _kr_loc_alt_pdf = (
            r'(?:[A-Z][a-z]+(?:-[a-z]+)*-(?:do|si|gu|dong|gun|myeon|eup|ri)|'
            r'South\s+Korea|Korea|'
            r'Seoul|Busan|Incheon|Daegu|Daejeon|Gwangju|Ulsan|Sejong|'
            r'Gyeonggi|Gangwon|Chungcheong|Jeolla|Gyeongsang|Jeju)'
        )
        for _wkp in KR_WORKPLACE_KEYWORDS:
            if _wkp in _GENERIC_TYPES_SET:
                continue
            if _wkp.lower() not in page_lower:
                continue
            _inline_pat = re.compile(
                r'(?:I[\s\'’]?\s*(?:am|m)\s+)?'
                r'(?:currently\s+|now\s+|recently\s+|previously\s+|formerly\s+)?'
                r'(?:teaching|working|employed|teach|work|am\s+teaching|am\s+working)\s+'
                r'(?:at|for|in|with)\s+'
                r'\b' + re.escape(_wkp) + r'\b'
                r'(?:\s+(?:in|at)\s+' + _kr_loc_alt_pdf
                + r'(?:\s*,\s*' + _kr_loc_alt_pdf + r')?)?'
                r'\s*[,.;]?',
                re.IGNORECASE,
            )
            for _im in _inline_pat.finditer(text):
                _phrase = _im.group().strip()
                if _phrase:
                    redact_texts.add(_phrase)
                    all_logs.append(f"[page{page_num}] KR_INLINE_PHRASE: {_phrase[:80]}")

        keep_keywords = {"korea", "south korea", "rok", "republic of korea", "한국"}

        # ── 사전 계산: EDUCATION 섹션 줄 인덱스 ──
        # 규칙: Education 섹션 내 전 세계 대학/학교명 삭제 절대 금지
        #       Work/Experience 섹션 업체명만 삭제
        # 2컬럼 PDF 대응: 텍스트 추출 순서가 레이아웃과 다를 수 있음
        # → 정방향(헤더→다음섹션) + degree 인접 검사 두 가지 병용
        _page_lines = text.split("\n")
        _edu_lines: set = set()
        _edu_flag = False
        for _idx, _ll in enumerate(_page_lines):
            _ll_s = _ll.strip().lower()
            if _ll_s and _EDU_SECTION_RE.match(_ll_s):
                _edu_flag = True
            elif _edu_flag and _ll_s and _OTHER_SECTION_RE.match(_ll_s):
                _edu_flag = False
            if _edu_flag:
                _edu_lines.add(_idx)
        # ── 보조: 학위 키워드 인접 줄도 교육 섹션으로 간주 ──
        # 2컬럼 PDF에서 Education 헤더 전에 교육 내용(Liberty University 등)이 추출되는 경우 보호
        _degree_kws_edu = frozenset({
            "degree", "bachelor", "master", "associate", "diploma",
            "doctorate", "phd", "m.a.", "b.a.", "b.s.", "m.s.", "mba",
            "b.ed.", "m.ed.", "b.sc.", "m.sc.", "honors", "honours",
        })
        for _idx, _ll in enumerate(_page_lines):
            _ll_lower = _ll.lower()
            if any(dk in _ll_lower for dk in _degree_kws_edu):
                # 이 줄 ±3줄을 교육 섹션으로 추가 표시
                for _adj in range(max(0, _idx - 3), min(len(_page_lines), _idx + 4)):
                    _edu_lines.add(_adj)

        # ── 한국 학원명/도시명 redact ──
        # KR_WORKPLACE_KEYWORDS: EDUCATION 섹션 제외, 유형명(_GENERIC_TYPES_SET) 건너뜀
        for _idx, line in enumerate(_page_lines):
            stripped = line.strip()
            if not stripped:
                continue
            s_lower = stripped.lower()
            has_kr = any(kw in s_lower for kw in KR_KEYWORDS)
            has_kr_work = any(kw in s_lower for kw in KR_WORKPLACE_KEYWORDS
                              if kw not in _GENERIC_TYPES_SET)

            # 학원 유형명 + 한국 지표(Branch/도시) → "{Type}, South Korea"
            # 단, "X — JobTitle" 같이 ↔직책 구분자 있는 줄은 기존 로직(아래)에
            # 위임 — 직책 부분("— English Teacher") 보존 필요
            # _edu_lines 검사 전 — 2컬럼 PDF에서 워크경험 줄이 EDU 옆에 위치하는 케이스 대응
            _has_jt_sep_pre = (" — " in stripped) or (" – " in stripped)
            # " - " 는 "Branch -" 같은 suffix marker일 수도 있으므로
            # 직책 키워드까지 함께 있을 때만 job title sep 로 간주
            if not _has_jt_sep_pre and (" - " in stripped):
                _after_dash = stripped.split(" - ", 1)[1].lower() if " - " in stripped else ""
                if any(jt in _after_dash for jt in ("teacher", "tutor", "instructor",
                                                    "lecturer", "professor", "trainer",
                                                    "manager", "coach")):
                    _has_jt_sep_pre = True
            if not _has_jt_sep_pre and len(stripped) <= 100:
                _matched_generic = None
                for _gt in sorted(_GENERIC_TYPES_SET, key=len, reverse=True):
                    if _gt in s_lower:
                        _matched_generic = _gt
                        break
                if _matched_generic:
                    _has_br_marker = "branch" in s_lower
                    _has_kr_city = any(
                        kw.lower() in s_lower for kw in KR_KEYWORDS
                        if len(kw) > 2 and kw.lower() not in ("korea", "south korea")
                    )
                    if _has_br_marker or _has_kr_city or has_kr_work:
                        _gen_cap = " ".join(w.capitalize() for w in _matched_generic.split())
                        _new_line = f"{_gen_cap}, South Korea"
                        _existing_idx = next(
                            (i for i, (o, _) in enumerate(job_title_replaces) if o == stripped),
                            -1
                        )
                        if _existing_idx >= 0:
                            job_title_replaces[_existing_idx] = (stripped, _new_line)
                        else:
                            job_title_replaces.append((stripped, _new_line))
                        all_logs.append(
                            f"[page{page_num}] KR_INST_REPLACE: {stripped[:60]} → {_new_line}"
                        )

            if _idx in _edu_lines:
                continue

            # 브랜드 접두사 + 일반 유형명 → 유형명 보존 (replacement redact)
            # "Broad Language School" → "Language School"
            # !! 80자 초과 줄(서술문)은 제외 — "management for Kindergarten" 오탐 방지
            # !! has_kr 또는 has_kr_work 조건 추가 — "Homeroom Kindergarten Teacher"(한국 키워드 없음) 오탐 방지
            if len(stripped) <= 80 and (has_kr or has_kr_work):
                for m in _RE_KR_SCHOOL_PREFIX.finditer(stripped):
                    canonical = " ".join(w.capitalize() for w in m.group(1).split())
                    replacement_redacts.append((m.group(), canonical))
                    all_logs.append(f"[page{page_num}] KR_SCHOOL_PREFIX→{canonical}: {m.group()[:60]}")

            # 학원명 redact: 같은 줄에 한국 키워드 있거나, 짧은 줄(<50자)
            # 유형명(language school 등)은 건너뜀 — 위에서 처리
            _is_job_line_kw = " — " in stripped or " - " in stripped  # 직업 타이틀줄
            if has_kr_work:
                if _is_job_line_kw:
                    # 직업 타이틀줄: 갭 없이 제거 → 전체 줄 교체 방식 (job_title_replaces)
                    # (redact_texts에 추가하지 않음 — 아래 full-line clean 섹션에서 처리)
                    _cleaned_jt = stripped
                    # 1순위: Branch 전체 패턴 (Gwanak Branch 등) — KW 개별 제거 전에 먼저 처리
                    _cleaned_jt = re.sub(r'\b[A-Za-z]+(?:\s+[A-Za-z]+)?\s+Branch\b', '', _cleaned_jt, flags=re.I)
                    # 2순위: 괄호 위치 태그: (Yongsan-Gu SMU Camp)
                    _cleaned_jt = re.sub(
                        r'\s*\([^)]{3,50}(?:Gu|Branch|Camp|District|Center|Centre|City)\b[^)]*\)',
                        '', _cleaned_jt, flags=re.I)
                    # 3순위: KW 브랜드명 (긴 것 먼저)
                    _sorted_kws = sorted(
                        (kw for kw in KR_WORKPLACE_KEYWORDS if kw not in _GENERIC_TYPES_SET and kw in s_lower),
                        key=len, reverse=True
                    )
                    for _kw in _sorted_kws:
                        _cleaned_jt = re.sub(r'\b' + re.escape(_kw) + r'\b', '', _cleaned_jt, flags=re.I)
                        # Fallback: redact_texts에도 직접 추가
                        # (job_title_replaces는 apply_redactions() 이후 실행 → 선행 redact로 _orig_jt 불일치 가능)
                        _fb_pat = re.compile(r'\b' + re.escape(_kw) + r'\b', re.I)
                        for _fb_m in _fb_pat.finditer(stripped):
                            redact_texts.add(_fb_m.group())
                            all_logs.append(f"[page{page_num}] JOB_KW_FALLBACK: {_fb_m.group()}")
                    # 연속 공백 정리
                    _cleaned_jt = re.sub(r'[ \t]{2,}', ' ', _cleaned_jt).strip()
                    # "Junior -  — Teacher" → "Junior - Teacher" (branch 제거 후 dangling — 정리)
                    _cleaned_jt = re.sub(r'\s+-\s+—\s+', ' - ', _cleaned_jt)
                    # 앞뒤 고립 " - " 또는 " — " 정리
                    _cleaned_jt = re.sub(r'^[\s\-—]+', '', _cleaned_jt).strip()
                    # 파이프 뒤 고아 "Korea" 제거: "Teacher | Korea — South Korea" → "Teacher — South Korea"
                    # (Korea Poly School → Korea 잔류 후 "|" 구분자 사이에 홀로 남은 경우)
                    _cleaned_jt = re.sub(
                        r'\|\s*Korea\s*(?=[—\-])', '| ', _cleaned_jt, flags=re.I)
                    _cleaned_jt = re.sub(r'\|\s*—', '—', _cleaned_jt)
                    _cleaned_jt = re.sub(r'\|\s*-\s', '- ', _cleaned_jt)
                    # 브랜드 제거 후 "— 한국도시, South Korea" 잔류 제거
                    # "Head Teacher — Song-do, Incheon, South Korea" → "Head Teacher — South Korea"
                    # ★ 사용자 요청: 광역 행정구역(-do/-si 접미사) 은 보존
                    #   "Gyeonggi-do" 안의 "Gyeonggi"도 보존 (negative lookahead)
                    _keep_kws_jt = {"korea", "south korea"}
                    for _city_jt in KR_KEYWORDS:
                        if len(_city_jt) <= 2 or _city_jt.lower() in _keep_kws_jt:
                            continue
                        # -do (도, 광역) / -si (시) 보존 (Gyeonggi-do, Seongnam-si 등)
                        _cl = _city_jt.lower()
                        if _cl.endswith("-do") or _cl.endswith("-si"):
                            continue
                        # "Gyeonggi" 가 "Gyeonggi-do" 안에 있으면 매칭 제외
                        _city_jt_pat = re.compile(
                            r'\b' + re.escape(_city_jt) + r'\b(?!-(?:do|si))\s*,?\s*', re.I)
                        # ① 직업타이틀 clean text에서 제거
                        _cleaned_jt = _city_jt_pat.sub('', _cleaned_jt)
                        # ② PDF 텍스트 분리 대응: redact_texts에도 직접 추가
                        for _cm in _city_jt_pat.finditer(stripped):
                            redact_texts.add(_cm.group().rstrip(', '))
                            all_logs.append(f"[page{page_num}] JOB_CITY_FALLBACK: {_cm.group().strip()}")
                    # 대시 뒤 쉼표/공백 정리
                    _cleaned_jt = re.sub(r'([—\-])\s*,+\s*', r'\1 ', _cleaned_jt)
                    _cleaned_jt = re.sub(r'[ \t]{2,}', ' ', _cleaned_jt).strip()
                    if _cleaned_jt != stripped:
                        job_title_replaces.append((stripped, _cleaned_jt))
                        all_logs.append(f"[page{page_num}] JOB_TITLE_CLEAN: {stripped[:50]} → {_cleaned_jt[:50]}")
            # 직업 타이틀줄 중 KR_WORKPLACE 없어도 괄호/브랜치 있는 경우 처리
            # (Sookmyung Women's University — Teacher (Yongsan-Gu Camp) 같은 경우)
            if _is_job_line_kw and not has_kr_work:
                _paren_re = re.compile(
                    r'\s*\([^)]{3,50}(?:Gu|Branch|Camp|District|Center|Centre|City)\b[^)]*\)', re.I)
                _branch_re = re.compile(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\s+Branch\b')
                _cleaned_np = stripped
                _cleaned_np = _paren_re.sub('', _cleaned_np)
                _cleaned_np = _branch_re.sub('', _cleaned_np)
                _cleaned_np = re.sub(r'[ \t]{2,}', ' ', _cleaned_np).strip()
                if _cleaned_np != stripped:
                    job_title_replaces.append((stripped, _cleaned_np))
                    all_logs.append(f"[page{page_num}] JOB_PAREN_CLEAN: {stripped[:50]} → {_cleaned_np[:50]}")
            # ── 직업 타이틀줄 후처리: 대학교 고용주 + 한국인 이름형 브랜드 ──────
            # 위 두 블록 이후 실행 — has_kr_work 무관
            # 대상: "Sookmyung Women's University — Teacher" / "Min-Byung-Chul Junior - Teacher"
            if _is_job_line_kw:
                # 현재 job_title_replaces에서 이 줄의 최신 클린 텍스트 가져오기
                _current_clean_jt = next(
                    (_c for _o, _c in job_title_replaces if _o == stripped), stripped)
                _cleaned_jt2 = _current_clean_jt
                # 4순위: 고용주 대학교명 제거
                # "Sookmyung Women's University — Native English Teacher" → "Native English Teacher"
                _cleaned_jt2 = re.sub(
                    r'^.+?University\s*(?:—|-)\s*', '', _cleaned_jt2, flags=re.I).strip()
                # 5순위: 한국인 이름형 학원 브랜드 (Min-Byung-Chul Junior)
                # 패턴: 대문자+소문자 대시 연결 3단어 + 선택적 Junior/Senior
                _cleaned_jt2 = re.sub(
                    r'\b[A-Z][a-z]+-[A-Z][a-z]+-[A-Z][a-z]+(?:\s+(?:Junior|Senior|Academy|Institute))?\b',
                    '', _cleaned_jt2).strip()
                # 6순위: 한국 세부 행정구역만 제거 (X-gu/dong/gun/myeon/eup/ri)
                # ★ -do(도, 광역) / -si(시) 는 보존 — 사용자 요청 (Gyeonggi-do 등)
                _cleaned_jt2 = re.sub(
                    r'\b[A-Z][a-z]+(?:-[a-z]+)*-(?:gu|dong|gun|myeon|eup|ri)\b\s*,?\s*',
                    '', _cleaned_jt2, flags=re.I).strip()
                # 앞뒤 고립 구분자 정리
                _cleaned_jt2 = re.sub(r'^[\s\-—]+', '', _cleaned_jt2).strip()
                _cleaned_jt2 = re.sub(r'[ \t]{2,}', ' ', _cleaned_jt2).strip()
                # 중복 콤마 정리: "Institute, , South Korea" → "Institute, South Korea"
                _cleaned_jt2 = re.sub(r',\s*,', ',', _cleaned_jt2)
                _cleaned_jt2 = re.sub(r',\s*(?=[—–-])', ' ', _cleaned_jt2)  # ", —" → " —"
                if _cleaned_jt2 != stripped:
                    _existing_idx = next(
                        (i for i, (o, _) in enumerate(job_title_replaces) if o == stripped), -1)
                    if _existing_idx >= 0:
                        job_title_replaces[_existing_idx] = (stripped, _cleaned_jt2)
                    else:
                        job_title_replaces.append((stripped, _cleaned_jt2))
                    all_logs.append(
                        f"[page{page_num}] JOB_EXTRA_CLEAN: {stripped[:50]} → {_cleaned_jt2[:50]}")

            # 비직업타이틀 학원명: 한국 키워드 있거나 짧은 줄일 때만 redact
            # len <= 55: 경계값(50자) 라인도 포함 ("at Poly. New teachers..." 등)
            # not has_kr 조건 제거: "Korea Poly School" 처럼 국가명+학원명 조합도 전체 삭제
            if has_kr_work and not _is_job_line_kw and (has_kr or len(stripped) <= 55):
                if len(stripped) <= 55:
                    redact_texts.add(stripped)
                    all_logs.append(f"[page{page_num}] KR_INST_LINE: {stripped[:60]}")
                for kw in KR_WORKPLACE_KEYWORDS:
                    if kw in _GENERIC_TYPES_SET:
                        continue  # 유형명 유지
                    kw_pat = re.compile(r"\b" + re.escape(kw) + r"\b", re.I)
                    for m in kw_pat.finditer(stripped):
                        redact_texts.add(m.group())
                        all_logs.append(f"[page{page_num}] KR_ACADEMY: {m.group()}")

            # 한국 도시명 처리 (korea 자체는 유지)
            # 사용자 요청: "특정지역 언급시 South Korea로 변경"
            # - 줄에 이미 "South Korea"/"Korea" 있으면: 도시 삭제만
            # - 없으면: 도시 → "South Korea" 로 대체 (지역 정보 보존)
            if page_has_korea and has_kr and not has_kr_work:
                _line_has_country = bool(re.search(
                    r'\b(?:south\s+korea|korea|rok|republic\s+of\s+korea|한국)\b',
                    stripped, re.I))
                for city in KR_KEYWORDS:
                    if len(city) <= 2 or city.lower() in keep_keywords:
                        continue
                    city_pat = re.compile(r"\b" + re.escape(city) + r"\b", re.I)
                    for m in city_pat.finditer(stripped):
                        if _line_has_country:
                            redact_texts.add(m.group())
                            all_logs.append(f"[page{page_num}] KR_CITY_DEL: {m.group()}")
                        else:
                            replacement_redacts.append((m.group(), "South Korea"))
                            all_logs.append(f"[page{page_num}] KR_CITY→SouthKorea: {m.group()}")

        # ── 학교/기관명 일반화 + 외국 도시/나이/비자 — 줄 단위 스캔 ──
        # EDUCATION 섹션 줄 + 직업 타이틀줄(" — " 포함)은 RE_SCHOOL_NAMED 미적용
        # → 고용주 대학명(Sookmyung Women's University — Teacher) 보호
        for _idx, _line in enumerate(_page_lines):
            _ls = _line.strip()
            if not _ls:
                continue
            _is_job_title_line = " — " in _ls or " - " in _ls and any(
                w in _ls.lower() for w in ("teacher", "tutor", "instructor", "manager", "scholar")
            )
            # RE_SCHOOL_NAMED: PDF 경로에서 비활성화
            # ★ 규칙: Education 섹션 전 세계 대학/학교명 삭제 절대 금지
            # PDF 섹션 감지는 2컬럼 레이아웃에서 불안정 → RE_SCHOOL_NAMED 사용 금지
            # Work 섹션 고용주 학교명은 KR_WORKPLACE_KEYWORDS + KR_INST_LINE으로 처리
            # (francis parker, hillside collegiate 등 KR_WORKPLACE_KEYWORDS에 직접 등록)
            # 직업 타이틀줄 위치/괄호 태그는 job_title_replaces에서 일괄 처리 (위)
            # (JOB_PAREN_LOC / JOB_BRANCH → job_title_replaces 방식으로 통합, 갭 없이 제거)
            # ★ 학습 폴더 호환: 학교/경력 줄("Univ X | Canterbury, UK | 2020-2022")은
            #   separator(|·•) 있어 그대로 유지 — 도시명 단어 redact 무력화.
            #   단독 주소줄(개인 주소)만 PRE-PASS가 line-level로 처리함.
            _line_has_sep = ("|" in _ls) or ("•" in _ls) or ("·" in _ls)

            # 해외 도시/주는 redact 하지 않음 (사용자 요청: 해외지명 보존)
            # 단독 개인 주소줄(짧고 separator 없는 줄)은 PRE-PASS LINE_PURGE/
            # US_ADDR_LINE / FOREIGN_ADDR_LINE에서 줄 단위로 처리됨
            # 나이/비자 → blank redact (학교 줄에서도 삭제 — 개인정보)
            for m in RE_AGE_MENTION.finditer(_ls):
                redact_texts.add(m.group())
                all_logs.append(f"[page{page_num}] AGE: {m.group()[:40]}")
            for m in RE_VISA_STATUS.finditer(_ls):
                redact_texts.add(m.group())
                all_logs.append(f"[page{page_num}] VISA: {m.group()[:40]}")
            # 한국 도시+국가 주소줄: separator 없는 단독 주소만 삭제
            if not _line_has_sep:
                for m in RE_KR_CITY_COUNTRY.finditer(_ls):
                    redact_texts.add(m.group())
                    all_logs.append(f"[page{page_num}] KR_CITY_COUNTRY: {m.group()[:40]}")
            # 연락 초대 줄: "please email me at X", "find me on WhatsApp" → 줄 전체 redact
            if _CONTACT_INVITE_RE.search(_ls) or _CONTACT_PLATFORM_RE.search(_ls):
                redact_texts.add(_ls)
                all_logs.append(f"[page{page_num}] CONTACT_INVITE_LINE: {_ls[:80]}")

        # 위치 라벨 → 값만 redact (Korea 대체 텍스트 삽입)
        # 직장명 라벨 → 값만 redact

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            s_lower = stripped.lower()

            # 위치: "Current Location: Seoul, South Korea" → redact → "Korea"
            # 비한국 주소 ("location: Greater London, SW16 5DE") → 값 전체 redact
            for label in LOCATION_LABELS:
                pat = re.match(rf"({re.escape(label)}\s*:?\s*)(.+)", s_lower)
                if not pat:
                    continue
                label_end = len(pat.group(1))
                value_text = stripped[label_end:].strip()
                if not value_text:
                    continue
                if _is_korea(pat.group(2)):
                    replacement_redacts.append((value_text, "Korea"))
                    all_logs.append(f"[page{page_num}] LOC→Korea: {value_text[:40]}")
                else:
                    # 비한국 주소: 라벨 포함 줄 전체 삭제 (개인 거주지 정보)
                    redact_texts.add(value_text)
                    redact_texts.add(stripped)  # 라벨("location:") 포함 전체 줄도 삭제
                    all_logs.append(f"[page{page_num}] LOC_NONKR_DEL: {stripped[:60]}")

            # 직장명: "Current Employer: YBM Academy" → redact (빈 대체)
            for label in WORKPLACE_LABELS:
                pat = re.match(rf"({re.escape(label)}\s*:?\s*)(.+)", s_lower)
                if pat:
                    label_end = len(pat.group(1))
                    value_text = stripped[label_end:].strip()
                    if value_text:
                        redact_texts.add(value_text)
                        all_logs.append(f"[page{page_num}] WORK_REDACT: {value_text[:40]}")

        if dry:
            continue

        # ── [PRE-PASS] PII 줄 line-level 깔끔 삭제 (학습폴더 스타일 호환) ─
        # 학습 폴더 5739: "Canterbury, United Kingdom" 유지 → 외국 도시 단독 줄은 보존
        # 학습 폴더 6155: 이름·연락처·생일·정확한 한국 주소 깔끔 제거
        # 규칙:
        #   1) PII 라벨 줄 (이름/생일/국적/연락처 라벨)
        #   2) 연락초대 / 카톡·텔레그램 핸들
        #   3) 한국 도시+국가 (Gwangmyeong, South Korea) — 정확한 한국 주소
        #   4) 짧은 외국 도시+국가만 있는 줄 (개인 주소) — "|" 없고 30자 이하
        #   5) 짧은 US 도시+주만 있는 줄 — "|" 없고 30자 이하
        # ※ 학교/경력 줄 ("University X | Canterbury, UK | 2020-2022") 은 유지
        _purge_pdict = page.get_text("dict", flags=0)
        _purge_count = 0
        for _pblk in _purge_pdict.get("blocks", []):
            if _pblk.get("type") != 0:
                continue
            for _pln in _pblk.get("lines", []):
                _pspans = _pln.get("spans", [])
                _pltxt = "".join(s.get("text", "") for s in _pspans).strip()
                if not _pltxt or len(_pltxt) < 3:
                    continue
                _pltxt_lower = _pltxt.lower()
                _has_separator = ("|" in _pltxt) or ("•" in _pltxt) or ("·" in _pltxt)
                _is_short = len(_pltxt) <= 35

                # job_title_replaces 대상 라인은 PRE-PASS LINE_PURGE 스킵
                # (전체 줄 reinsert가 적절한 클린 텍스트로 처리)
                if any(_orig == _pltxt for _orig, _ in job_title_replaces):
                    continue

                # 줄 단위 PII 감지
                _hit = False
                _reason = ""
                if _should_skip_line(_pltxt_lower):
                    _hit, _reason = True, "PII_LABEL"
                elif _CONTACT_INVITE_RE.search(_pltxt_lower) or _CONTACT_PLATFORM_RE.search(_pltxt_lower):
                    _hit, _reason = True, "CONTACT_INVITE"
                elif RE_KR_CITY_COUNTRY.search(_pltxt):
                    # 한국 주소 — 학교 줄("English Academy, Seoul" 등)도 통째 삭제하지 말고
                    # "|" 없는 짧은 단독 주소줄만
                    if not _has_separator:
                        _hit, _reason = True, "KR_ADDR_LINE"
                elif (RE_FOREIGN_CITY.search(_pltxt) and _is_short and not _has_separator):
                    _hit, _reason = True, "FOREIGN_ADDR_LINE"
                elif (RE_US_CITY_STATE.search(_pltxt) and _is_short and not _has_separator):
                    _hit, _reason = True, "US_ADDR_LINE"

                if not _hit:
                    continue
                _pbbox = fitz.Rect(_pln["bbox"])
                _pbbox.x0 -= 1.0; _pbbox.x1 += 1.0
                _pbbox.y0 -= 0.5; _pbbox.y1 += 0.5
                page.add_redact_annot(_pbbox, fill=(1, 1, 1))
                _purge_count += 1
                all_logs.append(f"[page{page_num}] LINE_PURGE/{_reason}: {_pltxt[:80]}")

        if _purge_count > 0:
            page.apply_redactions()

        # Redact 적용 — 일반 (흰색 덮기) — 옆 글자 클리핑 방지용 안전 패딩
        def _safe_rect(r):
            """글자 박스 ±0.4px 패딩 — adjacent glyph 보호."""
            nr = fitz.Rect(r)
            nr.x0 -= 0.4; nr.x1 += 0.4
            nr.y0 -= 0.2; nr.y1 += 0.2
            return nr

        # ── JOB_LINE_REINSERT를 main redact 전에 처리 ─────────────────────
        # 이유: 개별 redact(KW_FALLBACK 등)가 원본 줄을 부분 변경하면
        #       이후 _orig_jt 매칭 실패. 줄 단위 처리를 우선 적용.
        if not dry and job_title_replaces:
            import fitz as _fitz_pre
            _pg_fonts_pre = page.get_fonts()
            _jt_bold_ref_pre = None
            _jt_reg_ref_pre = None
            for _pgf in _pg_fonts_pre:
                _pfbase = _pgf[3].lower()
                _pfref = _pgf[4]
                if "bold" in _pfbase:
                    _jt_bold_ref_pre = _pfref
                elif "italic" not in _pfbase and "oblique" not in _pfbase:
                    _jt_reg_ref_pre = _pfref
            _jt_dict_pre = page.get_text("dict", flags=0)
            _done_origs = set()
            _processed_jt_bboxes: list = []  # 처리완료된 bbox 추적
            for _orig_jt, _clean_jt in job_title_replaces:
                if _orig_jt in _done_origs:
                    continue
                for _jblk in _jt_dict_pre.get("blocks", []):
                    if _jblk.get("type") != 0:
                        continue
                    _matched = False
                    for _jln in _jblk.get("lines", []):
                        _jlt = "".join(s.get("text", "") for s in _jln.get("spans", [])).strip()
                        if _orig_jt not in _jlt:
                            continue
                        _jrect = _fitz_pre.Rect(_jln["bbox"])
                        _jspans = _jln.get("spans", [])
                        _jfst = _jspans[0] if _jspans else {}
                        _jfsize = _jfst.get("size", 8)
                        _jfflags = _jfst.get("flags", 0)
                        _jcolor_i = _jfst.get("color", 0)
                        _jrgb = (
                            ((_jcolor_i >> 16) & 0xFF) / 255.0,
                            ((_jcolor_i >> 8) & 0xFF) / 255.0,
                            (_jcolor_i & 0xFF) / 255.0,
                        )
                        _is_bold_pre = bool(_jfflags & 16)
                        _jx0, _jy0, _jx1, _jy1 = _jrect
                        # ASCII normalize (em-dash -> "-", NBSP -> " ")
                        _safe_clean = _clean_jt.replace("—", "-").replace("–", "-")
                        _safe_clean = _safe_clean.replace("​", "").replace(" ", " ")
                        _safe_clean = re.sub(r"\s+", " ", _safe_clean).strip()
                        # add_redact_annot text param: whiteout + replacement in one step
                        _exp_rect = _fitz_pre.Rect(_jx0, _jy0, max(_jx1, _jx0 + 350), _jy1)
                        _processed_jt_bboxes.append(_jrect)
                        try:
                            page.add_redact_annot(
                                _exp_rect, text=_safe_clean, fill=(1, 1, 1),
                                text_color=_jrgb, fontsize=max(7, _jfsize - 1),
                                align=0,
                            )
                            page.apply_redactions()
                            all_logs.append(f"[page{page_num}] JOB_LINE_REINSERT_PRE: {_safe_clean[:60]}")
                        except Exception as _e:
                            all_logs.append(f"[page{page_num}] JOB_INSERT_FAIL_PRE: {_e}")

                        _done_origs.add(_orig_jt)
                        _matched = True
                        break
                    if _matched:
                        break
            # 처리 완료된 항목 제거 (후속 JOB_LINE_REINSERT 루프에서 중복 처리 방지)
            job_title_replaces = [(_o, _c) for _o, _c in job_title_replaces if _o not in _done_origs]

        for target in redact_texts:
            instances = page.search_for(target)
            for rect in instances:
                page.add_redact_annot(_safe_rect(rect), fill=(1, 1, 1))

        # Redact 적용 — 대체 텍스트 삽입 (위치→Korea)
        for target, replacement in replacement_redacts:
            instances = page.search_for(target)
            for rect in instances:
                page.add_redact_annot(
                    _safe_rect(rect), text=replacement, fill=(1, 1, 1),
                    text_color=(0, 0, 0), fontsize=10,
                )

        if redact_texts or replacement_redacts:
            page.apply_redactions()

        # ── Supplementary: get_text("dict") 기반 정밀 redact ──
        # 이유: search_for()는 다중 스팬 줄(Bold+Normal 혼합), 사이드바 레이아웃에서 실패 가능
        # 해결: 각 라인/스팬의 실제 bbox를 직접 사용해 추가 whiteout
        _pdict = page.get_text("dict", flags=0)
        _has_extra = False
        for _blk in _pdict.get("blocks", []):
            if _blk.get("type") != 0:
                continue
            for _lnobj in _blk.get("lines", []):
                _lspans = _lnobj.get("spans", [])
                _ltxt = "".join(s.get("text", "") for s in _lspans)
                _lstripped = _ltxt.strip()
                if not _lstripped:
                    continue
                # PII 라벨 줄 → 전체 라인 bbox whiteout (Bold: "Citizenship:" + Normal: " English" 같은 경우)
                if _should_skip_line(_lstripped.lower()):
                    page.add_redact_annot(fitz.Rect(_lnobj["bbox"]), fill=(1, 1, 1))
                    _has_extra = True
                    all_logs.append(f"[page{page_num}] LINE_BBOX: {_lstripped[:60]}")
                    continue
                # 스팬 단위 이메일/전화 (사이드바/컬럼 레이아웃에서 search_for 실패 보완)
                _ltxt_lower = _lstripped.lower()
                _line_has_kr = any(kw in _ltxt_lower for kw in KR_KEYWORDS)
                _line_has_kr_work = any(kw in _ltxt_lower for kw in KR_WORKPLACE_KEYWORDS
                                        if kw not in _GENERIC_TYPES_SET)
                for _span in _lspans:
                    _st = _span.get("text", "")
                    if not _st or len(_st) < 4:
                        continue
                    _sr = fitz.Rect(_span["bbox"])
                    if RE_EMAIL.search(_st):
                        page.add_redact_annot(_sr, fill=(1, 1, 1))
                        _has_extra = True
                        all_logs.append(f"[page{page_num}] EMAIL_BBOX: {_st[:60]}")
                    elif (RE_PHONE.search(_st) and len(re.sub(r"\D", "", _st)) >= 7
                          and not all(re.match(r'^(?:19|20)\d{2}$', g) for g in re.findall(r'\d+', _st))):
                        page.add_redact_annot(_sr, fill=(1, 1, 1))
                        _has_extra = True
                        all_logs.append(f"[page{page_num}] PHONE_BBOX: {_st[:60]}")
                    # 아이콘 폰트 글리프 삭제 (FontAwesome/icomoon — U+E000~U+F8FF Private Use Area)
                    # 순수 아이콘 스팬만 (공백 제외 모든 문자가 PUA 범위) — 헤더 상단 400px 이내
                    _is_pure_glyph = (
                        bool(_st.strip())
                        and all(0xE000 <= ord(c) <= 0xF8FF or c.isspace() for c in _st)
                    )
                    if _is_pure_glyph:
                        _span_y0 = _span["bbox"][1]
                        if _span_y0 < 400:
                            page.add_redact_annot(_sr, fill=(1, 1, 1))
                            _has_extra = True
                            all_logs.append(f"[page{page_num}] ICON_GLYPH y={_span_y0:.0f}: {repr(_st[:8])}")
                # 한국 학원명 라인 보충 bbox 처리 (search_for Bold/Normal 혼합 실패 보완)
                # 전체 라인에 학원 키워드 + 한국 키워드 있을 때 → 브랜드 접두어 스팬 whiteout
                # ★ KR_INST_REPLACE 대상 라인은 스킵 — job_title_replaces가 전체 교체 처리
                _is_inst_replace = any(_orig == _lstripped for _orig, _ in job_title_replaces)
                if not _is_inst_replace and _line_has_kr_work and (_line_has_kr or len(_lstripped) < 60):
                    for _m in _RE_KR_SCHOOL_PREFIX.finditer(_lstripped):
                        # prefix 부분(match 에서 capture group 1 앞부분) 을 span 단위로 삭제
                        _prefix_txt = _lstripped[:_m.start(1)].strip()  # "JLS " 부분
                        if _prefix_txt:
                            for _sp2 in _lspans:
                                _sp2_t = _sp2.get("text", "").strip()
                                if not _sp2_t:
                                    continue
                                # 멀티스팬: "JLS" 스팬이 prefix에 포함되는 경우 (Bold/Normal 분리)
                                if _sp2_t.lower() in _prefix_txt.lower():
                                    page.add_redact_annot(fitz.Rect(_sp2["bbox"]), fill=(1, 1, 1))
                                    _has_extra = True
                                    all_logs.append(f"[page{page_num}] KR_BRAND_BBOX: {_sp2_t[:30]}")
                                # 싱글스팬: 스팬 안에 prefix가 포함된 경우 → search_for로 정밀 위치
                                elif (len(_lspans) == 1 and _prefix_txt.lower() in _sp2_t.lower()):
                                    for _prf_rect in page.search_for(
                                            _prefix_txt, clip=fitz.Rect(_sp2["bbox"])):
                                        page.add_redact_annot(_prf_rect, fill=(1, 1, 1))
                                        _has_extra = True
                                        all_logs.append(
                                            f"[page{page_num}] KR_BRAND_SINGLE: {_prefix_txt[:30]}"
                                        )
                    # 한국 도시명 스팬 삭제 — 학원명 포함 줄은 도시명 유지
                    # "English Academy, Seoul" → Seoul 살림 (사용자 요구)
                    # (도시 삭제는 독립 주소줄에만 적용)
                # 주소줄 bbox: "Gwangmyeong, South Korea" → 라인 전체 삭제
                # 단, JOB_LINE_REINSERT_PRE로 처리된 bbox는 스킵 (재redact 방지)
                _ln_bbox_y = _lnobj["bbox"][1]
                try:
                    _is_jt_processed = any(abs(_pjr[1] - _ln_bbox_y) < 2 for _pjr in _processed_jt_bboxes)
                except NameError:
                    _is_jt_processed = False
                if not _is_jt_processed and RE_KR_CITY_COUNTRY.search(_lstripped):
                    page.add_redact_annot(fitz.Rect(_lnobj["bbox"]), fill=(1, 1, 1))
                    _has_extra = True
                    all_logs.append(f"[page{page_num}] KR_ADDR_BBOX: {_lstripped[:50]}")
        if _has_extra:
            page.apply_redactions()

        # ── 직업 타이틀줄 전체 교체 (갭 없이) ──────────────────────────────────
        # 브랜드명/지점명 제거: 개별 단어 whitebox 대신 줄 전체 교체 → 공백 갭 없음
        if not dry and job_title_replaces:
            import fitz as _fitz
            # 페이지에 내장된 폰트 참조명 감지 (F5=Bold, F4=Regular 등)
            # em-dash(U+2014) 지원: 내장 폰트 사용 시 정상 렌더링
            _pg_fonts = page.get_fonts()
            _jt_bold_ref = None
            _jt_reg_ref = None
            for _pgf in _pg_fonts:
                _pfbase = _pgf[3].lower()  # basefont name
                _pfref = _pgf[4]           # /F-name in resource dict
                if "bold" in _pfbase:
                    _jt_bold_ref = _pfref
                elif "italic" not in _pfbase and "oblique" not in _pfbase:
                    _jt_reg_ref = _pfref

            _jt_dict = page.get_text("dict", flags=0)
            for _orig_jt, _clean_jt in job_title_replaces:
                # 원문 줄을 get_text("dict")에서 찾기
                for _jblk in _jt_dict.get("blocks", []):
                    if _jblk.get("type") != 0:
                        continue
                    for _jln in _jblk.get("lines", []):
                        _jlt = "".join(s.get("text", "") for s in _jln.get("spans", [])).strip()
                        # 원문이 이 줄에 포함되어 있으면 교체
                        if _orig_jt not in _jlt:
                            continue
                        _jrect = _fitz.Rect(_jln["bbox"])
                        _jspans = _jln.get("spans", [])
                        _jfst = _jspans[0] if _jspans else {}
                        _jfsize = _jfst.get("size", 8)
                        _jfflags = _jfst.get("flags", 0)
                        _jcolor_i = _jfst.get("color", 0)
                        _jrgb = (
                            ((_jcolor_i >> 16) & 0xFF) / 255.0,
                            ((_jcolor_i >> 8) & 0xFF) / 255.0,
                            (_jcolor_i & 0xFF) / 255.0,
                        )
                        # 폰트: 내장 PDF 폰트 참조 우선 (em-dash 지원)
                        _is_bold = bool(_jfflags & 16)
                        _jfontref = (_jt_bold_ref if _is_bold else _jt_reg_ref)
                        # 1. 전체 줄 whiteout
                        page.add_redact_annot(_jrect, fill=(1, 1, 1))
                        page.apply_redactions()
                        # 2. 클린 텍스트 재삽입 (baseline = bbox 하단)
                        _jx0, _jy0, _jx1, _jy1 = _jrect
                        _ins_ok = False
                        # 시도 1: 내장 폰트
                        if _jfontref:
                            try:
                                page.insert_text(
                                    _fitz.Point(_jx0, _jy1), _clean_jt,
                                    fontsize=_jfsize, color=_jrgb,
                                    fontname=_jfontref,
                                )
                                _ins_ok = True
                            except Exception:
                                pass
                        # 시도 2: 표준 폰트 폴백
                        if not _ins_ok:
                            try:
                                page.insert_text(
                                    _fitz.Point(_jx0, _jy1), _clean_jt,
                                    fontsize=_jfsize, color=_jrgb,
                                    fontname=("tibo" if _is_bold else "tiro"),
                                )
                                _ins_ok = True
                            except Exception as _ins_err:
                                all_logs.append(f"[page{page_num}] JOB_INSERT_FAIL: {_ins_err}")
                        # 3. 밑줄 (원본이 bold+underline인 경우)
                        if _jfflags & 4:  # bit 2 = underline
                            _ul_y = _jy1 + 0.5
                            page.draw_line(
                                _fitz.Point(_jx0, _ul_y),
                                _fitz.Point(_jx1, _ul_y),
                                color=_jrgb, width=0.5
                            )
                        all_logs.append(f"[page{page_num}] JOB_LINE_REINSERT: {_clean_jt[:60]}")
                        break  # 한 줄만 처리 후 다음 교체 항목으로

    # ── 헤더 페이지 연락처 스트립 정리 + 번호/사진 삽입 ──
    # _header_page_num: 위에서 자동 감지된 실제 이력서 헤더 페이지 번호
    # 전략: 베이지 헤더 배경색 감지 → 연락처 스트립을 같은 색으로 덮어 아이콘/PII 흔적 제거
    #       → 번호(왼쪽) + 사진(오른쪽) 삽입
    if not dry and len(doc) > 0:
        hdr_page = doc[_header_page_num]
        pw = hdr_page.rect.width

        # ── 헤더 배경색 + 높이 감지 ──────────────────────────────────────────
        # 실제 colored 사각형이 페이지 상단에 있을 때만 헤더로 인정
        # fill=(0,0,0)/(1,1,1) 제외, y0<50, width>40% 조건
        _hdr_fill_color = None   # None = 헤더 없음
        _hdr_bottom = 0
        for _drw in hdr_page.get_drawings():
            _dfill = _drw.get("fill")
            _dr = _drw.get("rect", fitz.Rect())
            if (_dfill and _dfill not in ((1, 1, 1), (0, 0, 0))
                    and _dr.y0 < 50 and _dr.width > pw * 0.4
                    and _dr.y1 > _hdr_bottom):
                _hdr_fill_color = _dfill
                _hdr_bottom = _dr.y1

        _has_real_header = _hdr_fill_color is not None and _hdr_bottom > 40

        if _has_real_header:
            # ── 실제 colored 헤더가 있는 경우 ─────────────────────────────
            # 구분선 Y 감지 (헤더 아래 얇은 수평선 = 연락처 스트립 하단)
            _divider_y = _hdr_bottom + 110  # fallback
            for _drw in hdr_page.get_drawings():
                _dr = _drw.get("rect", fitz.Rect())
                if (_dr.y0 > _hdr_bottom and (_dr.y1 - _dr.y0) < 3
                        and _dr.width > pw * 0.5 and _dr.y0 < _divider_y):
                    _divider_y = _dr.y0

            # ── 헤더 밝기 기반 커버 방식 결정 ────────────────────────────
            # 밝은 헤더(beige/tan, 0.5이상): 왼쪽 사진 플레이스홀더만 덮기 → 중앙 이름 보존 (6165형)
            # 어두운 헤더(navy/dark, 0.5미만): 전체 이름 영역 덮기 → 자간벌림/벡터 이름 제거 (6189형)
            _r, _g, _b = _hdr_fill_color
            _brightness = 0.299 * _r + 0.587 * _g + 0.114 * _b

            if _brightness >= 0.5:
                # 밝은 헤더 (6165형): 왼쪽 사진 플레이스홀더만 덮기
                _name_cover_right = pw * 0.34
                _cover_mode = "LEFT"
            else:
                # 어두운 헤더 (6189형): 이름 영역 전체 덮기
                photo_reserve = _hdr_bottom - 10
                _name_cover_right = pw - photo_reserve - 5
                _cover_mode = "FULL"

            hdr_page.draw_rect(
                fitz.Rect(0, 0, _name_cover_right, _hdr_bottom),
                color=None, fill=_hdr_fill_color,
            )
            all_logs.append(
                f"[page{_header_page_num}] HDR_COVER({_cover_mode}): 0-{_name_cover_right:.0f} / 0-{_hdr_bottom:.0f}"
            )

            # ② 연락처 스트립 화이트아웃 (헤더 하단 ~ 구분선)
            # 아이콘 박스(벡터) + PII 텍스트 완전 제거
            hdr_page.draw_rect(
                fitz.Rect(0, _hdr_bottom - 2, pw, _divider_y + 1),
                color=None, fill=(1, 1, 1),
            )
            all_logs.append(
                f"[page{_header_page_num}] STRIP_WHITE: y={_hdr_bottom:.0f}~{_divider_y:.0f}"
            )

            # ③ 강사 번호 삽입 (헤더 왼쪽)
            _num_y = _hdr_bottom * 0.70
            # 밝은 헤더 → 어두운 텍스트, 어두운 헤더 → 밝은 텍스트
            _r, _g, _b = _hdr_fill_color
            _brightness = 0.299 * _r + 0.587 * _g + 0.114 * _b
            _num_color = (0.95, 0.95, 0.95) if _brightness < 0.5 else (0.20, 0.14, 0.06)
            hdr_page.insert_text(
                fitz.Point(12, _num_y),
                str(brj_number),
                fontsize=24,
                fontname="helv",
                color=_num_color,
            )
            all_logs.append(f"[page{_header_page_num}] BRJNUM: {brj_number}")

            # ④ 사진 삽입 (헤더 오른쪽) — 커버레터 제외
            photo = None if _is_cover_letter(filepath) else (photo_path or _find_photo_for_candidate(filepath))
            if photo and photo.exists():
                _img_margin = 5
                _img_h = int(_hdr_bottom - _img_margin * 2)
                photo_rect = fitz.Rect(
                    pw - _img_h - _img_margin, _img_margin,
                    pw - _img_margin, _hdr_bottom - _img_margin,
                )
                try:
                    hdr_page.insert_image(photo_rect, filename=str(photo))
                    all_logs.append(f"[photo] {_header_page_num}p {_img_h}px: {photo.name}")
                except Exception as e:
                    all_logs.append(f"[photo] 삽입 실패: {e}")
            else:
                all_logs.append("[photo] 없음 (스킵)")

        elif _is_cover_letter(filepath):
            # 커버레터: 번호/사진 일체 삽입 안 함 — PII redact만 적용됨
            all_logs.append(f"[page{_header_page_num}] COVER_LETTER_SKIP_HEADER")
        else:
            # ── 헤더 없는 일반 CV ─────────────────────────────────────────
            # 이름이 redact된 자리(이름 영역 ~ 첫 섹션 헤더) 에 번호+사진 배치

            # 1) 첫 섹션 헤더(PROFILE/CAREER 등) y 위치 탐색
            _SECTION_HDRS = {
                "PROFILE", "ABOUT", "ABOUT ME", "SUMMARY",
                "CAREER", "EXPERIENCE", "EDUCATION", "SKILLS",
            }
            _first_sec_y = 200  # fallback
            _pdict_nh = hdr_page.get_text("dict", flags=0)
            for _blk_nh in _pdict_nh.get("blocks", []):
                if _blk_nh.get("type") != 0:
                    continue
                for _ln_nh in _blk_nh.get("lines", []):
                    _lt_nh = "".join(
                        s.get("text", "") for s in _ln_nh.get("spans", [])
                    ).strip().upper()
                    _ly_nh = _ln_nh["bbox"][1]
                    if _lt_nh in _SECTION_HDRS and 30 < _ly_nh < 300:
                        _first_sec_y = min(_first_sec_y, int(_ly_nh) - 4)

            # 2) 삽입 존 결정 — 사용자 정답 기준 좌표
            #    정답: 6300 at x=39 y=66-115 (fs~48) / photo at (435,20)-(543,138)
            _face_covers = _scanned_face_covers.get(_header_page_num, [])
            if _face_covers:
                # 스캔 PDF — 사용자 정답 좌표 고정
                _num_left_x = 39
                _num_baseline_y = 115
                _num_fs = 48
                # photo rect: 정답 기준
                _photo_rect_target = fitz.Rect(435, 20, 543, 138)
                _is_scanned_layout = True
            else:
                # 텍스트 PDF — 최상단 텍스트 감지하여 충돌 회피
                _topmost_y = None
                for _blk_t in _pdict_nh.get("blocks", []):
                    if _blk_t.get("type") != 0:
                        continue
                    for _ln_t in _blk_t.get("lines", []):
                        _txt = "".join(s.get("text","") for s in _ln_t.get("spans",[])).strip()
                        if not _txt:
                            continue
                        _ly = _ln_t["bbox"][1]
                        if _ly > 5 and (_topmost_y is None or _ly < _topmost_y):
                            _topmost_y = _ly
                if _topmost_y is None:
                    _topmost_y = max(_first_sec_y, 100)

                # 정답 기준: 6321 at y=2-69 x=35-142 (fs~60) — 페이지 최상단 큰 글씨
                # 텍스트 PDF에서 이름/연락처 redact 후 빈 공간 활용
                if _topmost_y >= 55:
                    # 충분한 위쪽 공간 — 큰 번호 배치
                    _zone_top = 2
                    _zone_bot = max(int(_topmost_y), 60)
                    _num_fs = 60  # 정답 기준 고정
                    _num_left_x = 35
                else:
                    # 위 공간 부족 — 작게
                    _zone_top = max(8, int(_topmost_y) - 4)
                    _zone_bot = _zone_top + 40
                    _num_fs = min(36, max(20, int((_zone_bot - _zone_top) * 0.75)))
                    _num_left_x = 15
                _zone_h = _zone_bot - _zone_top
                _num_baseline_y = _zone_top + int(_num_fs * 0.95)
                _is_scanned_layout = False

            # 3) 번호 삽입 (왼쪽, 크게)
            hdr_page.insert_text(
                fitz.Point(_num_left_x, _num_baseline_y),
                str(brj_number),
                fontsize=_num_fs,
                fontname="helv",
                color=(0.15, 0.15, 0.15),
            )
            all_logs.append(f"[page{_header_page_num}] NO_HDR BRJNUM: {brj_number}  x={_num_left_x} y={_num_baseline_y} fs={_num_fs}")

            # 4) 사진 삽입 (오른쪽, 적정 크기) — 커버레터 제외
            photo = photo_path or _find_photo_for_candidate(filepath)
            if photo and photo.exists():
                if _is_scanned_layout:
                    photo_rect = _photo_rect_target
                else:
                    # 텍스트 PDF: 정답 기준 (464,58)-(552,161) — W=88 H=103, 우상단
                    _pw2 = 88
                    _ph2 = 103
                    photo_rect = fitz.Rect(
                        pw - _pw2 - 60, 58,
                        pw - 60,        58 + _ph2,
                    )
                try:
                    hdr_page.insert_image(photo_rect, filename=str(photo))
                    all_logs.append(f"[photo] NO_HDR {_header_page_num}p: {photo.name} rect=({photo_rect.x0:.0f},{photo_rect.y0:.0f})-({photo_rect.x1:.0f},{photo_rect.y1:.0f})")
                except Exception as e:
                    all_logs.append(f"[photo] 삽입 실패: {e}")
            else:
                all_logs.append("[photo] NO_HDR 없음 (스킵)")

    # ── 링크 어노테이션 제거 (이메일/URI 하이퍼링크 → redaction 이후에도 잔류) ──
    # 이유: add_redact_annot/apply_redactions은 텍스트 스트림만 덮음.
    #       link annotation은 별도 오브젝트 → PDF 뷰어에서 클릭 가능하게 남음.
    #       이메일 파란 밑줄, CamScanner URI 등 완전 제거.
    if not dry:
        for _pg in doc:
            _links = _pg.get_links()
            for _lnk in _links:
                # URI (이메일 mailto:, http: 등) 및 내부 점프 링크 모두 제거
                _lkind = _lnk.get("kind", -1)
                if _lkind in (
                    fitz.LINK_URI,     # 0 — http/mailto
                    fitz.LINK_GOTO,    # 1 — 내부 페이지 이동
                    fitz.LINK_NAMED,   # 4 — named dest
                    2,                 # LINK_GOTOR — remote goto
                ):
                    try:
                        _pg.delete_link(_lnk)
                    except Exception:
                        pass
        all_logs.append(f"[post] LINK_ANNOTS_CLEARED: {sum(len(p.get_links()) for p in doc)} remaining")

    # PDF 메타데이터 클리어 (이름/제목 제거)
    if not dry:
        doc.set_metadata({
            "title": "", "author": "", "subject": "",
            "keywords": "", "creator": "", "producer": "",
        })

    return doc, all_logs


def _build_search_patterns(candidate: dict = None):
    """PDF redaction용 검색 패턴 목록 생성"""
    patterns = [
        (RE_EMAIL, "EMAIL"),
        (RE_LINKEDIN, "LINKEDIN"),
        (RE_KAKAO, "KAKAO"),
        (RE_SNS, "SNS"),
        (RE_URL, "URL"),
        (RE_PASSPORT, "PASSPORT"),
        (RE_KR_ADDRESS, "KR_ADDR"),
        (RE_KR_RESIDENTIAL, "KR_RESID"),   # 부영 1차 아파트 등
        (RE_EN_ADDRESS, "EN_ADDR"),
        (RE_KR_ROMANIZED_ADDR, "KR_ROMAN_ADDR"),
        (RE_DOB_SPELLED, "DOB_SPELLED"),   # "20th of September 1996" 서술형 생년월일
        (RE_KR_PHONE_PREFIX, "KR_PHONE_PREFIX"),  # "+82 (0)" 전화 프리픽스 잔류
        (RE_KR_POSTAL, "KR_POSTAL"),       # "10906" 우편번호 잔류
    ]

    # 전화번호는 별도 처리 (7자리+ 필터)
    re_phone_strict = re.compile(
        r"(?:\+?\d{1,3}[\s\-.]?)?\(?\d{2,4}\)?[\s\-.]?\d{3,4}[\s\-.]?\d{3,4}"
    )
    patterns.append((re_phone_strict, "PHONE"))

    # 후보자 DB 정보
    if candidate:
        for field in ("email", "mobile_phone", "kakaotalk"):
            val = candidate.get(field, "")
            if val and len(val) > 3:
                patterns.append((re.compile(re.escape(val), re.I), f"DB_{field.upper()}"))

    return patterns


# ══════════════════════════════════════════════════════
#  6. 파일 저장 + 백업 + 로그
# ══════════════════════════════════════════════════════

def _convert_docx_to_pdf(docx_path: Path, pdf_path: Path, timeout: int = 60) -> bool:
    """DOCX → PDF 변환 (Word COM 자동화). 성공 시 True.
    timeout: Word COM 최대 대기 초 (기본 60초, 초과 시 강제 종료 후 False)
    """
    import threading as _thr

    def _word_com() -> bool:
        word = None
        try:
            import comtypes.client
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0   # wdAlertsNone — 보안/업데이트 팝업 차단
            doc = word.Documents.Open(
                str(docx_path.resolve()),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
            )
            doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)  # 17 = wdFormatPDF
            doc.Close(SaveChanges=False)
            return True
        except ImportError:
            return False
        except Exception as e:
            print(f"  [PDF WARN] COM 변환 실패: {e}")
            return False
        finally:
            if word is not None:
                try:
                    word.Quit(SaveChanges=False)
                except Exception:
                    pass

    result = [False]
    t = _thr.Thread(target=lambda: result.__setitem__(0, _word_com()), daemon=True)
    t.start()
    t.join(timeout)
    if t.is_alive():
        print(f"  [PDF WARN] Word COM {timeout}초 타임아웃 — 강제 중단")
        # Word 프로세스 종료
        try:
            import subprocess
            subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE"],
                           capture_output=True)
        except Exception:
            pass
        return False
    if result[0]:
        return True
    # fallback: LibreOffice
    try:
        import subprocess
        lo_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        lo = next((p for p in lo_paths if Path(p).exists()), None)
        if lo:
            subprocess.run(
                [lo, "--headless", "--convert-to", "pdf",
                 "--outdir", str(pdf_path.parent), str(docx_path)],
                capture_output=True, timeout=60,
            )
            # LibreOffice는 원본 이름으로 저장하므로 rename
            lo_out = pdf_path.parent / (docx_path.stem + ".pdf")
            if lo_out.exists() and lo_out != pdf_path:
                lo_out.rename(pdf_path)
            return pdf_path.exists()
    except Exception as e:
        print(f"  [PDF WARN] LibreOffice 실패: {e}")
    print("  [PDF FAIL] PDF 변환 불가 — comtypes/docx2pdf/LibreOffice 모두 없음")
    print("             pip install comtypes 또는 pip install docx2pdf 로 설치")
    return False


def backup_original(filepath: Path):
    """원본 → originals/ 복사"""
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    dest = BACKUP_DIR / f"{ts}_{filepath.name}"
    shutil.copy2(str(filepath), str(dest))
    return dest


def _docx_to_pdf(docx_path: Path) -> Path | None:
    """DOCX → 소용량 PDF 변환 래퍼. 성공 시 PDF Path, 실패 시 None."""
    pdf_path = docx_path.with_suffix(".pdf")
    if _convert_docx_to_pdf(docx_path, pdf_path, timeout=90):
        return pdf_path
    return None


def save_log(filepath, brj_number: int, logs: list[str]):
    """삭제 로그 저장"""
    LOG_DIR.mkdir(parents=True, exist_ok=True)
    stem = filepath.stem if isinstance(filepath, Path) else str(filepath).replace("#", "")
    log_path = LOG_DIR / f"BRJ{brj_number}_{stem}.json"
    data = {
        "source": str(filepath),
        "brj_number": brj_number,
        "processed_at": datetime.now().isoformat(),
        "removals": logs,
        "total_removals": len(logs),
    }
    log_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return log_path


def auto_detect_candidate(filepath: Path):
    """파일명에서 강사번호 자동 감지 (이름/위치 교차 검증)"""
    stem = filepath.stem
    numbers = re.findall(r"(\d{3,5})", stem)
    for num_str in numbers:
        candidate = get_candidate(int(num_str))
        if not candidate:
            continue
        # 번호가 파일명 맨 앞이면 확실한 매칭 (예: "3060_resume.pdf")
        if stem.startswith(num_str):
            return candidate
        # _# 뒤에 오면 확실한 매칭 (예: "resume_3060.pdf", "#3060.pdf")
        if re.search(rf"[_#]{num_str}\b", stem):
            return candidate
        # 그 외: 후보자 이름이 파일명에 포함되는지 검증
        cand_name = (candidate.get("full_name") or "").lower()
        stem_lower = stem.lower()
        if cand_name:
            name_words = [w for w in cand_name.split() if len(w) >= 3]
            if any(w in stem_lower for w in name_words):
                return candidate
        # 이름 매칭 실패 → 우연한 숫자 (연도 2026 등) → 스킵
    return None


# ══════════════════════════════════════════════════════
#  7. CLI 명령어
# ══════════════════════════════════════════════════════

def cmd_process(args):
    """process 명령"""
    input_path = Path(args.input)
    output_dir = Path(args.output) if args.output else DEFAULT_OUTPUT
    brj_number = args.number
    dry = args.dry
    photo_path = Path(args.photo) if getattr(args, "photo", None) else None
    if photo_path and not photo_path.exists():
        print(f"  [WARN] photo not found: {photo_path}")
        photo_path = None

    if not input_path.exists():
        print(f"[ERROR] Path not found: {input_path}")
        sys.exit(1)

    # 파일 목록
    if input_path.is_file():
        files = [input_path]
    else:
        files = sorted(
            f for f in input_path.iterdir()
            if f.suffix.lower() in (".docx", ".pdf")
            and not f.name.startswith("~")
            and not f.name.startswith(".")
        )

    if not files:
        print(f"[INFO] No .docx/.pdf files in {input_path}")
        return

    mode = "DRY RUN" if dry else "PROCESS"
    print(f"\n{'=' * 60}")
    print(f"  BRIDGE Document Processor v2.0  [{mode}]")
    print(f"  Input:  {input_path}")
    print(f"  Output: {output_dir}")
    print(f"  Files:  {len(files)}")
    print(f"{'=' * 60}\n")

    for filepath in files:
        print(f"[{mode}] {filepath.name}")

        # 번호 결정
        current_number = brj_number
        candidate = None

        if current_number:
            candidate = get_candidate(current_number)
            if candidate:
                print(f"  #{current_number} {candidate['full_name']}")
            else:
                print(f"  #{current_number} (DB에 없음)")
        else:
            candidate = auto_detect_candidate(filepath)
            if candidate:
                current_number = candidate["sheet_number"]
                print(f"  Auto: #{current_number} {candidate['full_name']}")
            else:
                nums = re.findall(r"(\d{3,5})", filepath.stem)
                if nums:
                    current_number = int(nums[0])
                    print(f"  Filename: #{current_number}")
                else:
                    print(f"  [SKIP] 번호 없음. --number N 사용")
                    continue

        try:
            ext = filepath.suffix.lower()

            if ext == ".docx":
                doc, logs = process_docx(filepath, current_number, candidate, dry,
                                         photo_path=photo_path)
            elif ext == ".pdf":
                doc, logs = process_pdf(filepath, current_number, candidate, dry,
                                        photo_path=photo_path)
            else:
                print(f"  [SKIP] 미지원: {ext}")
                continue

            # 로그 출력
            if logs:
                print(f"  PII {len(logs)}건 {'발견' if dry else '삭제'}:")
                for l in logs[:10]:
                    print(f"    - {l}")
                if len(logs) > 10:
                    print(f"    ... +{len(logs)-10}건")

            if dry:
                if ext == ".pdf":
                    doc.close()
                print(f"  [DRY] 미리보기 완료\n")
                continue

            # 백업
            backup = backup_original(filepath)
            print(f"  Backup: {backup.name}")

            # 저장
            output_dir.mkdir(parents=True, exist_ok=True)
            out_name = _build_output_filename(current_number, candidate, ext)
            out_path = output_dir / out_name

            if ext == ".docx":
                # DOCX 임시 저장 후 Word COM으로 PDF 변환
                doc.save(str(out_path))
                pdf_result = _docx_to_pdf(out_path)
                if pdf_result and pdf_result.exists():
                    out_path.unlink(missing_ok=True)  # DOCX 삭제
                    out_path = pdf_result
                    print(f"  Output: {out_path.name} (PDF 변환됨)")
                else:
                    print(f"  Output: {out_path.name} (DOCX — PDF 변환 실패)")
            elif ext == ".pdf":
                # 최대 압축: garbage 컬렉션 + deflate + 이미지/폰트 압축 + 정리
                doc.save(
                    str(out_path),
                    garbage=4, deflate=True, deflate_images=True,
                    deflate_fonts=True, clean=True,
                )
                doc.close()
                _sz_kb = out_path.stat().st_size // 1024
                print(f"  Output: {out_path.name} ({_sz_kb}KB)")
                # 2차 압축: 큰 이미지 다운샘플링 (>500KB만)
                if _sz_kb > 500:
                    try:
                        _compress_pdf_images(out_path)
                        _new_sz = out_path.stat().st_size // 1024
                        if _new_sz < _sz_kb:
                            print(f"  Compressed: {_sz_kb}KB → {_new_sz}KB")
                    except Exception as _ce:
                        print(f"  [WARN] 2차 압축 실패: {_ce}")

            # 삭제 로그 저장
            if logs:
                log_path = save_log(filepath, current_number, logs)
                print(f"  Log: {log_path.name}")

            print(f"  [OK]\n")

        except Exception as e:
            print(f"  [ERROR] {e}\n")
            import traceback
            traceback.print_exc()

    print(f"{'=' * 60}")
    print(f"  Done! {output_dir}")
    print(f"{'=' * 60}")


def cmd_lookup(args):
    """lookup 명령"""
    results = lookup_candidate(args.query)
    if not results:
        print(f"[INFO] No results for: {args.query}")
        return

    print(f"\n  {'#':>5}  {'Name':<35}  {'Nationality':<12}  Email")
    print(f"  {'─'*5}  {'─'*35}  {'─'*12}  {'─'*25}")
    for sn, name, nat, email in results:
        print(f"  {sn or '?':>5}  {(name or '')[:35]:<35}  {(nat or '')[:12]:<12}  {email or ''}")
    print()


def _resolve_candidate(filepath: Path) -> tuple[int | None, dict | None]:
    """파일에서 후보자 번호+정보 자동 탐지. Returns (number, candidate_dict)."""
    candidate = auto_detect_candidate(filepath)
    if candidate and candidate.get("sheet_number"):
        return candidate["sheet_number"], candidate

    # auto_detect_candidate가 이미 파일명 숫자를 이름 교차검증했으므로
    # 여기서 재시도 불필요 → 이메일 매칭으로 직행

    # 이메일 기반 매칭 (DOCX + PDF)
    emails = []
    ext = filepath.suffix.lower()
    try:
        if ext == ".docx":
            emails = _extract_emails_from_docx(filepath)
        elif ext == ".pdf":
            import fitz
            doc = fitz.open(str(filepath))
            for page in doc:
                for m in RE_EMAIL.finditer(page.get_text()):
                    emails.append(m.group().lower())
            doc.close()
    except Exception:
        pass
    for em in emails:
        cand = find_candidate_by_email(em)
        if cand:
            if cand.get("sheet_number"):
                return cand["sheet_number"], cand
            return None, cand
    return None, None


def _quick_extract_names(filepath: Path) -> set[str]:
    """문서에서 이름 빠르게 추출 (사진 매칭용). Returns lowercase name parts (3글자+)."""
    names = set()
    # 파일명 기반
    for fn_name in _extract_names_from_filename(filepath):
        names.update(w.lower().rstrip(".") for w in fn_name.split() if len(w) >= 3)
    # PDF 첫 페이지 텍스트 기반
    if filepath.suffix.lower() == ".pdf":
        try:
            import fitz
            doc = fitz.open(str(filepath))
            if len(doc) > 0:
                first_text = doc[0].get_text()
                for line in (l.strip() for l in first_text.split("\n") if l.strip()):
                    if len(line) >= 60:
                        continue
                    if line.strip().lower() in _RESUME_SECTION_HEADERS:
                        continue
                    if line.isupper() and len(line.split()) >= 2:
                        continue
                    if _should_skip_line(line.lower()):
                        continue
                    words = line.split()
                    if 2 <= len(words) <= 5 and all(
                        w[0].isupper() or w in ("de", "van", "von", "del", "K.", "K")
                        for w in words if w
                    ):
                        names.update(w.lower().rstrip(".") for w in words if len(w) >= 3)
                        break  # 첫 이름 줄만
            doc.close()
        except Exception:
            pass
    return names


def _extract_nationality_gender(filepath: Path) -> dict:
    """문서 텍스트에서 국적/성별 추출 (DB 없는 후보자용)."""
    text = ""
    ext = filepath.suffix.lower()
    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(str(filepath))
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
        elif ext == ".docx":
            import docx
            doc = docx.Document(str(filepath))
            text = "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return {}

    info = {}
    text_lower = text.lower()

    # 국적 — 명시적 라벨 우선
    nat_line = re.search(
        r"(?:nationality|citizenship|country of (?:origin|birth))\s*:?\s*(.+)",
        text, re.I,
    )
    if nat_line:
        val = nat_line.group(1).strip().lower()
        for key in _NAT_KR:
            if key in val:
                info["nationality"] = key
                break
    # 라벨 없으면 본문에서 국가명 빈도 추론
    if "nationality" not in info:
        nat_hints = {
            "american": ["united states", "u.s.a", "u.s. citizen"],
            "british": ["united kingdom", "u.k.", "british citizen", "british national"],
            "irish": ["ireland", "irish citizen"],
            "canadian": ["canada", "canadian citizen"],
            "australian": ["australia", "australian citizen"],
            "south african": ["south africa"],
        }
        for nat, patterns in nat_hints.items():
            if any(p in text_lower for p in patterns):
                info["nationality"] = nat
                break

    # 성별 — 명시적 라벨 우선
    gender_line = re.search(
        r"(?:gender|sex)\s*:?\s*(male|female|m|f)\b", text, re.I,
    )
    if gender_line:
        val = gender_line.group(1).strip().lower()
        if val in ("male", "m"):
            info["gender"] = "male"
        elif val in ("female", "f"):
            info["gender"] = "female"

    # DOB — 라벨 패턴
    dob_match = re.search(
        r"(?:date\s+of\s+birth|d\.?o\.?b\.?|born|birth\s*date)\s*:?\s*"
        r"(\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}|\d{4}[/\-\.]\d{1,2}[/\-\.]\d{1,2}|"
        r"\w+\s+\d{1,2},?\s+\d{4}|\d{4})",
        text, re.I,
    )
    if dob_match:
        info["dob"] = dob_match.group(1).strip()
    # DOB 없으면 고등학교 졸업년도로 출생년 추정 (졸업년 - 18)
    if "dob" not in info:
        hs_match = re.search(r"high\s+school.*?(\d{4})", text, re.I)
        if hs_match:
            grad_year = int(hs_match.group(1))
            if 2000 <= grad_year <= 2030:
                info["dob"] = str(grad_year - 18)

    return info


def cmd_batch(args):
    """batch 명령: incoming/ 폴더 전체 일괄 처리 (커버레터+이력서 자동 병합)"""
    dry = args.dry
    incoming = INCOMING_DIR
    if not incoming.exists() or not any(incoming.iterdir()):
        print(f"[INFO] incoming/ 폴더가 비어 있습니다: {incoming}")
        print(f"  파일을 여기에 넣으세요: {incoming}")
        return

    files = sorted(
        f for f in incoming.iterdir()
        if f.suffix.lower() in (".docx", ".pdf")
        and not f.name.startswith("~")
        and not f.name.startswith(".")
    )

    if not files:
        print(f"[INFO] .docx/.pdf 파일 없음: {incoming}")
        return

    mode = "DRY RUN" if dry else "BATCH"

    # ── 파일 분류: 커버레터 vs 이력서 ──
    cover_letters = []
    resumes = []
    for f in files:
        if _is_cover_letter(f):
            cover_letters.append(f)
        else:
            resumes.append(f)

    # ── 사진 맵 수집 ──
    photo_map = {}
    for f in incoming.iterdir():
        if f.suffix.lower() in (".jpg", ".jpeg", ".png") and not f.name.startswith(("~", ".")):
            photo_map[f.stem.lower()] = f

    # ── 자동 번호 카운터 (충돌 방지) ──
    db = _get_db()
    _next_auto = (db.execute("SELECT MAX(sheet_number) FROM candidates").fetchone()[0] or 3000) + 1

    # ── 후보자별 그룹핑 ──
    groups = {}  # number → {resume, cover, candidate}

    for filepath in resumes:
        num, cand = _resolve_candidate(filepath)
        if not num:
            num = _next_auto
            _next_auto += 1
        groups[num] = {"resume": filepath, "cover": None, "candidate": cand}

    # 커버레터 매칭: 번호 → 파일명 이름 유사도 → 신규 그룹
    _stop_words = {"cover", "letter", "resume", "copy", "of", "the", "for"}
    for filepath in cover_letters:
        num, cand = _resolve_candidate(filepath)

        # 1) 번호 매칭
        if num and num in groups:
            groups[num]["cover"] = filepath
            if cand and not groups[num]["candidate"]:
                groups[num]["candidate"] = cand
            continue

        # 2) 파일명 이름 유사도 매칭 ("Cover letter - grace peers" ↔ "Copy of Resume - grace peers")
        cl_words = set(re.findall(r"[a-z]{3,}", filepath.stem.lower())) - _stop_words
        matched = False
        for gnum, grp in groups.items():
            if grp["resume"] and not grp["cover"]:
                res_words = set(re.findall(r"[a-z]{3,}", grp["resume"].stem.lower())) - _stop_words
                common = cl_words & res_words
                if len(common) >= 1:
                    groups[gnum]["cover"] = filepath
                    if cand and not grp["candidate"]:
                        grp["candidate"] = cand
                    print(f"  [MATCH] Cover ↔ Resume by name: {common}")
                    matched = True
                    break
        if matched:
            continue

        # 3) 신규 그룹
        if not num:
            num = _next_auto
            _next_auto += 1
        groups[num] = {"resume": None, "cover": filepath, "candidate": cand}

    total = len(groups)
    print(f"\n{'=' * 60}")
    print(f"  BRIDGE Document Processor v2.4  [{mode}]")
    print(f"  Incoming:  {incoming}")
    print(f"  Output:    {DEFAULT_OUTPUT}")
    print(f"  Resumes:   {len(resumes)} | Cover Letters: {len(cover_letters)}")
    print(f"  Groups:    {total}")
    print(f"{'=' * 60}\n")

    import tempfile

    for current_number, grp in groups.items():
        resume_path = grp["resume"]
        cover_path = grp["cover"]
        candidate = grp["candidate"]

        # DB 없는 후보자: 문서에서 국적/성별 추출 (파일명 생성용)
        if not candidate:
            doc_path = resume_path or cover_path
            if doc_path and doc_path.exists():
                extracted_info = _extract_nationality_gender(doc_path)
                if extracted_info:
                    candidate = extracted_info
                    print(f"  [INFO] 문서에서 추출: {extracted_info}")

        label = f"#{current_number}"
        if candidate:
            label += f" {candidate.get('full_name', '')}"
        print(f"[{mode}] {label}")
        if resume_path:
            print(f"  Resume: {resume_path.name}")
        if cover_path:
            print(f"  Cover:  {cover_path.name}")

        # 사진 매칭: 번호 → 이력서 파일명 유사도 → 이름 유사도
        photo = None
        for pk, pf in list(photo_map.items()):
            if str(current_number) in pk:
                photo = pf
                break
        if not photo and resume_path:
            res_stem = resume_path.stem.lower()
            for pk, pf in list(photo_map.items()):
                # 이력서 이니셜/이름과 사진 파일명 비교
                res_words = set(re.findall(r"[a-z]{2,}", res_stem))
                pk_words = set(re.findall(r"[a-z]{2,}", pk))
                if res_words & pk_words:
                    photo = pf
                    break
        if not photo and candidate and candidate.get("full_name"):
            name_parts = candidate["full_name"].lower().split()
            for pk, pf in list(photo_map.items()):
                if any(part in pk for part in name_parts if len(part) >= 3):
                    photo = pf
                    break
        # NEW: 문서 텍스트에서 추출한 이름으로 사진 매칭
        if not photo:
            doc_names = set()
            for doc_path in [resume_path, cover_path]:
                if doc_path and doc_path.exists():
                    doc_names |= _quick_extract_names(doc_path)
            if doc_names:
                for pk, pf in list(photo_map.items()):
                    pk_words = set(re.findall(r"[a-z]{3,}", pk))
                    common = pk_words & doc_names
                    if common:
                        photo = pf
                        print(f"  [MATCH] Photo ↔ Doc name: {common}")
                        break
        # 그룹 1개 + 사진 1개면 자동 매칭
        if not photo and len(groups) == 1 and len(photo_map) == 1:
            photo = next(iter(photo_map.values()))
        all_logs = []
        temp_pdfs = []  # (순서, pdf_path) — 최종 병합용

        try:
            # ── 커버레터 처리 ──
            if cover_path:
                ext = cover_path.suffix.lower()
                if ext == ".docx":
                    # 커버레터 DOCX: PII 제거 (번호/사진 삽입 안 함)
                    cl_doc, cl_logs = process_docx(
                        cover_path, current_number, candidate, dry, photo_path=None
                    )
                    all_logs.extend(cl_logs)
                    if not dry:
                        backup_original(cover_path)
                        # 임시 DOCX 저장 → PDF 변환
                        tmp_docx = Path(tempfile.mktemp(suffix=".docx"))
                        cl_doc.save(str(tmp_docx))
                        tmp_pdf = Path(tempfile.mktemp(suffix=".pdf"))
                        if _convert_docx_to_pdf(tmp_docx, tmp_pdf):
                            temp_pdfs.append(("cover", tmp_pdf))
                            tmp_docx.unlink()
                        else:
                            temp_pdfs.append(("cover_docx", tmp_docx))
                        cover_path.unlink()
                elif ext == ".pdf":
                    cl_doc, cl_logs = process_pdf(
                        cover_path, current_number, candidate, dry
                    )
                    all_logs.extend(cl_logs)
                    if not dry:
                        backup_original(cover_path)
                        tmp_pdf = Path(tempfile.mktemp(suffix=".pdf"))
                        cl_doc.save(str(tmp_pdf), garbage=4, deflate=True)
                        cl_doc.close()
                        temp_pdfs.append(("cover", tmp_pdf))
                        cover_path.unlink()

            # ── 이력서 처리 ──
            if resume_path:
                ext = resume_path.suffix.lower()
                if ext == ".docx":
                    doc, res_logs = process_docx(
                        resume_path, current_number, candidate, dry, photo_path=photo
                    )
                    all_logs.extend(res_logs)
                    if not dry:
                        backup_original(resume_path)
                        tmp_docx = Path(tempfile.mktemp(suffix=".docx"))
                        doc.save(str(tmp_docx))
                        tmp_pdf = Path(tempfile.mktemp(suffix=".pdf"))
                        if _convert_docx_to_pdf(tmp_docx, tmp_pdf):
                            temp_pdfs.append(("resume", tmp_pdf))
                            tmp_docx.unlink()
                        else:
                            temp_pdfs.append(("resume_docx", tmp_docx))
                        resume_path.unlink()
                elif ext == ".pdf":
                    doc, res_logs = process_pdf(
                        resume_path, current_number, candidate, dry,
                        photo_path=photo,
                    )
                    all_logs.extend(res_logs)
                    if not dry:
                        backup_original(resume_path)
                        tmp_pdf = Path(tempfile.mktemp(suffix=".pdf"))
                        doc.save(str(tmp_pdf), garbage=4, deflate=True)
                        doc.close()
                        temp_pdfs.append(("resume", tmp_pdf))
                        resume_path.unlink()

            # ── 로그 출력 ──
            if all_logs:
                print(f"  PII {len(all_logs)}건 {'발견' if dry else '삭제'}:")
                for l in all_logs[:10]:
                    print(f"    - {l}")
                if len(all_logs) > 10:
                    print(f"    ... +{len(all_logs)-10}건")

            if dry:
                print(f"  [DRY] 미리보기 완료\n")
                continue

            # ── 최종 PDF 병합/출력 ──
            DEFAULT_OUTPUT.mkdir(parents=True, exist_ok=True)
            final_name = _build_output_filename(current_number, candidate, ".pdf")
            final_path = DEFAULT_OUTPUT / final_name

            pdf_parts = [p for _, p in temp_pdfs if p.suffix == ".pdf"]
            if len(pdf_parts) >= 2:
                # 커버레터 + 이력서 PDF 병합
                cover_p = next((p for t, p in temp_pdfs if "cover" in t), None)
                resume_p = next((p for t, p in temp_pdfs if "resume" in t), None)
                _merge_pdfs(cover_p, resume_p, final_path)
                print(f"  Output: {final_name} (커버레터+이력서 병합 PDF)")
            elif len(pdf_parts) == 1:
                shutil.move(str(pdf_parts[0]), str(final_path))
                print(f"  Output: {final_name} (PDF)")
            else:
                # PDF 변환 실패 — DOCX 그대로 출력
                docx_parts = [p for _, p in temp_pdfs if p.suffix == ".docx"]
                if docx_parts:
                    final_path = DEFAULT_OUTPUT / _build_output_filename(
                        current_number, candidate, ".docx")
                    shutil.move(str(docx_parts[0]), str(final_path))
                    print(f"  Output: {final_path.name} (DOCX — PDF 변환 실패)")

            # 임시 파일 정리
            for _, tmp in temp_pdfs:
                if tmp.exists():
                    tmp.unlink()

            if all_logs:
                log_path = save_log(
                    resume_path or cover_path or Path(f"#{current_number}"),
                    current_number, all_logs,
                )
                print(f"  Log: {log_path.name}")

            # DB 기록
            try:
                file_size = final_path.stat().st_size if final_path.exists() else 0
                _record_file_upload(
                    entity_id=current_number,
                    file_type=final_path.suffix.lstrip("."),
                    file_url=str(final_path),
                    file_size=file_size,
                )
                print(f"  [DB] file_uploads 기록 완료")
            except Exception as db_err:
                print(f"  [DB WARN] 기록 실패 (무시): {db_err}")

            # 사진 사용 후 incoming에서 제거 (다음 batch 오염 방지)
            if photo and photo.exists() and photo.parent == incoming:
                backup_original(photo)
                photo.unlink()
                photo_map.pop(photo.stem.lower(), None)
                print(f"  [PHOTO] {photo.name} → originals/ 이동")

            print(f"  [OK]\n")

        except Exception as e:
            print(f"  [ERROR] {e}\n")
            import traceback
            traceback.print_exc()

    print(f"{'=' * 60}")
    print(f"  Done! 결과: {DEFAULT_OUTPUT}")
    print(f"{'=' * 60}")


def cmd_download(args):
    """download 명령: S3에서 후보자 이력서+사진 다운로드 → PII 삭제 → PDF 출력"""
    number = args.number
    candidate = get_candidate(number)
    if not candidate:
        print(f"[ERROR] #{number} 후보자를 찾을 수 없습니다")
        return

    cand_id = candidate.get("candidate_id", "")
    print(f"  #{number} {candidate['full_name']} (ID: {cand_id or 'N/A'})")

    # bx.py로 AWS 자격증명 로드
    try:
        sys.path.insert(0, str(BASE_DIR / "tools"))
        from bx import load_to_env
        load_to_env()
    except Exception as e:
        print(f"[ERROR] AWS 자격증명 로드 실패: {e}")
        print(f"  bx.py set AWS_ACCESS_KEY_ID / AWS_SECRET_ACCESS_KEY / AWS_S3_BUCKET 필요")
        return

    import boto3

    bucket = os.environ.get("AWS_S3_BUCKET")
    region = os.environ.get("AWS_REGION", "ap-northeast-2")
    if not bucket:
        print("[ERROR] AWS_S3_BUCKET 환경변수 미설정")
        return

    s3 = boto3.client(
        "s3",
        region_name=region,
        aws_access_key_id=os.environ.get("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.environ.get("AWS_SECRET_ACCESS_KEY"),
    )

    # S3 경로: candidate_id 기반 (cnd_xxxx) + sheet_number 폴백
    prefixes_to_try = []
    if cand_id:
        prefixes_to_try.append(f"candidates/{cand_id}/")
    prefixes_to_try.extend([
        f"candidates/cnd_{number}/",
        f"candidates/{number}/",
        f"resumes/{number}/",
    ])

    _DOC_EXTS = {".pdf", ".docx", ".doc", ".hwp"}
    _IMG_EXTS = {".jpg", ".jpeg", ".png", ".webp"}
    downloaded_docs = []
    downloaded_photo = None
    dest_dir = args.output or str(INCOMING_DIR)
    Path(dest_dir).mkdir(parents=True, exist_ok=True)

    for prefix in prefixes_to_try:
        try:
            response = s3.list_objects_v2(Bucket=bucket, Prefix=prefix)
            for obj in response.get("Contents", []):
                key = obj["Key"]
                ext = Path(key).suffix.lower()
                s3_name = Path(key).name

                if ext in _DOC_EXTS:
                    filename = f"{number}_{s3_name}"
                    dest = Path(dest_dir) / filename
                    print(f"  [DOC] {key} → {filename}")
                    s3.download_file(bucket, key, str(dest))
                    downloaded_docs.append(dest)
                elif ext in _IMG_EXTS and not downloaded_photo:
                    # 사진: thumb 제외, 원본 우선
                    if "thumb" in s3_name.lower():
                        continue
                    filename = f"{number}_photo{ext}"
                    dest = Path(dest_dir) / filename
                    print(f"  [PHOTO] {key} → {filename}")
                    s3.download_file(bucket, key, str(dest))
                    downloaded_photo = dest
        except Exception:
            pass

    total = len(downloaded_docs) + (1 if downloaded_photo else 0)
    if total == 0:
        print(f"  S3에서 파일을 찾을 수 없습니다")
        print(f"  수동으로 파일을 incoming/에 넣고 batch 명령 사용:")
        print(f"    python doc_processor.py batch")
        return

    print(f"\n  {total}개 다운로드 완료 (문서 {len(downloaded_docs)}, 사진 {1 if downloaded_photo else 0})")

    if args.no_process:
        return

    # ── 자동 처리 ──
    import tempfile
    print(f"\n  자동 처리 시작...")
    temp_pdfs = []

    for f in downloaded_docs:
        try:
            ext = f.suffix.lower()
            is_cover = _is_cover_letter(f)

            if ext == ".docx":
                doc, logs = process_docx(
                    f, number, candidate, photo_path=None if is_cover else downloaded_photo)
            elif ext == ".pdf":
                doc, logs = process_pdf(
                    f, number, candidate, photo_path=None if is_cover else downloaded_photo)
            else:
                continue

            backup_original(f)

            if ext == ".docx":
                tmp_docx = Path(tempfile.mktemp(suffix=".docx"))
                doc.save(str(tmp_docx))
                tmp_pdf = Path(tempfile.mktemp(suffix=".pdf"))
                label = "cover" if is_cover else "resume"
                if _convert_docx_to_pdf(tmp_docx, tmp_pdf):
                    temp_pdfs.append((label, tmp_pdf))
                    tmp_docx.unlink()
                else:
                    temp_pdfs.append((f"{label}_docx", tmp_docx))
            elif ext == ".pdf":
                tmp_pdf = Path(tempfile.mktemp(suffix=".pdf"))
                doc.save(str(tmp_pdf), garbage=4, deflate=True)
                doc.close()
                label = "cover" if is_cover else "resume"
                temp_pdfs.append((label, tmp_pdf))

            f.unlink()
            if logs:
                save_log(f, number, logs)
            print(f"  PII {len(logs)}건 삭제: {f.name}")

        except Exception as e:
            print(f"  [ERROR] {f.name}: {e}")
            import traceback
            traceback.print_exc()

    # ── 최종 PDF 병합/출력 ──
    DEFAULT_OUTPUT.mkdir(parents=True, exist_ok=True)
    final_name = _build_output_filename(number, candidate, ".pdf")
    final_path = DEFAULT_OUTPUT / final_name

    pdf_parts = [p for _, p in temp_pdfs if p.suffix == ".pdf"]
    if len(pdf_parts) >= 2:
        cover_p = next((p for t, p in temp_pdfs if "cover" in t and p.suffix == ".pdf"), None)
        resume_p = next((p for t, p in temp_pdfs if "resume" in t and p.suffix == ".pdf"), None)
        _merge_pdfs(cover_p, resume_p, final_path)
        print(f"  Output: {final_name} (커버레터+이력서 병합)")
    elif len(pdf_parts) == 1:
        shutil.move(str(pdf_parts[0]), str(final_path))
        print(f"  Output: {final_name}")
    else:
        docx_parts = [p for _, p in temp_pdfs if p.suffix == ".docx"]
        if docx_parts:
            final_path = DEFAULT_OUTPUT / _build_output_filename(number, candidate, ".docx")
            shutil.move(str(docx_parts[0]), str(final_path))
            print(f"  Output: {final_path.name} (DOCX — PDF 변환 실패)")

    for _, tmp in temp_pdfs:
        if tmp.exists():
            tmp.unlink()

    # 사진 정리
    if downloaded_photo and downloaded_photo.exists():
        backup_original(downloaded_photo)
        downloaded_photo.unlink()

    print(f"  [OK] {final_path.name}")


def cmd_setup(_args=None):
    """setup 명령: 폴더 구조 생성 + 상태 확인"""
    dirs = [INCOMING_DIR, DEFAULT_OUTPUT, BACKUP_DIR, LOG_DIR]
    for d in dirs:
        d.mkdir(parents=True, exist_ok=True)

    print(f"\n  BRIDGE Document Processor v2.6 — 폴더 구조")
    print(f"  {'─' * 50}")
    print(f"  incoming/   : {INCOMING_DIR}")
    print(f"  processed/  : {DEFAULT_OUTPUT}")
    print(f"  originals/  : {BACKUP_DIR}")
    print(f"  logs/       : {LOG_DIR}")
    print()

    # incoming 파일 수
    incoming_files = [f for f in INCOMING_DIR.iterdir()
                      if f.suffix.lower() in (".docx", ".pdf")] if INCOMING_DIR.exists() else []
    processed_files = list(DEFAULT_OUTPUT.iterdir()) if DEFAULT_OUTPUT.exists() else []

    print(f"  대기 중: {len(incoming_files)}개")
    print(f"  처리됨: {len(processed_files)}개")

    if incoming_files:
        print(f"\n  대기 파일:")
        for f in incoming_files[:10]:
            print(f"    - {f.name}")
    print()


def cmd_init_db(_args=None):
    """init-db 명령: file_uploads 테이블 생성"""
    _ensure_file_uploads_table()
    db = _get_db()
    count = db.execute("SELECT COUNT(*) FROM file_uploads").fetchone()[0]
    print(f"\n  [init-db] file_uploads 테이블 준비 완료")
    print(f"  DB: {DB_PATH}")
    print(f"  기존 레코드: {count}건\n")


def main():
    parser = argparse.ArgumentParser(
        description="BRIDGE Document Processor v2.6 — PII 삭제 + 강사번호"
    )
    sub = parser.add_subparsers(dest="command")

    p_proc = sub.add_parser("process", help="단일 파일/폴더 처리")
    p_proc.add_argument("input", help="파일 또는 폴더")
    p_proc.add_argument("--number", "-n", type=int, help="강사번호")
    p_proc.add_argument("--output", "-o", help="출력 폴더")
    p_proc.add_argument("--photo", help="사진 파일 경로 (이력서에 삽입)")
    p_proc.add_argument("--dry", action="store_true", help="미리보기 (수정 없음)")

    p_batch = sub.add_parser("batch", help="incoming/ 폴더 일괄 처리")
    p_batch.add_argument("--dry", action="store_true", help="미리보기 (수정 없음)")

    p_dl = sub.add_parser("download", help="S3에서 후보자 파일 다운로드 + 처리")
    p_dl.add_argument("number", type=int, help="강사번호")
    p_dl.add_argument("--output", "-o", help="다운로드 폴더 (기본: incoming/)")
    p_dl.add_argument("--no-process", action="store_true", help="다운로드만 (처리 안 함)")

    p_look = sub.add_parser("lookup", help="후보자 검색")
    p_look.add_argument("query", help="이름 또는 이메일")

    sub.add_parser("setup", help="폴더 구조 생성 + 상태 확인")
    sub.add_parser("init-db", help="file_uploads 테이블 생성 (IF NOT EXISTS)")

    args = parser.parse_args()
    if args.command == "process":
        cmd_process(args)
    elif args.command == "batch":
        cmd_batch(args)
    elif args.command == "download":
        cmd_download(args)
    elif args.command == "lookup":
        cmd_lookup(args)
    elif args.command == "setup":
        cmd_setup(args)
    elif args.command == "init-db":
        cmd_init_db(args)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
