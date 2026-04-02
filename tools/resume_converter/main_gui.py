"""
main_gui.py — BRIDGE Resume Converter GUI  v2.0
Tkinter + TkinterDnD2

레이아웃: 좌측(대기열+내비) / 가운데(편집+파일목록) / 우측(미리보기)
"""

from __future__ import annotations

import io
import logging
import os
import re
import threading
import time
import webbrowser
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    _DND_AVAILABLE = True
except ImportError:
    _DND_AVAILABLE = False
    TkinterDnD = tk.Tk

from PIL import Image, ImageTk

# ── 경로 설정 ─────────────────────────────────────────────────────────────
BASE_DIR   = Path(__file__).parent
INBOX_DIR  = BASE_DIR / "inbox"
OUTPUT_DIR = BASE_DIR / "output"
LOGS_DIR   = BASE_DIR / "logs"
CONFIG_PATH= BASE_DIR / "config.json"

# ── 폰트 (맑은 고딕 — Windows Korean 시스템 폰트) ─────────────────────────
F  = "맑은 고딕"   # Malgun Gothic
FM = "Consolas"    # monospace

# ── 색상 팔레트 ───────────────────────────────────────────────────────────
C_BG     = "#FFFFFF"
C_SIDE   = "#F8F9FA"
C_BORDER = "#E8EAED"
C_PRI    = "#1D9E75"
C_WARN   = "#EF9F27"
C_DANGER = "#E24B4A"
C_TEXT   = "#202124"
C_SUB    = "#5F6368"
C_HOVER  = "#E8F5F0"
C_PAUSE  = "#FFCDD2"   # 옅은 붉은 (일시정지)
C_STOP   = "#E53935"   # 강제중지

# 대기열 행 색
C_QWAIT  = "#FFF9C4"
C_QACT   = "#E8F5E9"
C_QDONE  = "#F1F8E9"
C_QERR   = "#FFEBEE"

# ── 단계 정의 ─────────────────────────────────────────────────────────────
STEPS = ["파일 확인", "사진 크롭", "PII 제거", "PDF 생성"]
STEP_ACTIONS = [
    "파일 목록 확인 중",
    "얼굴 감지 & 크롭 중",
    "개인정보 탐지 & 제거 중",
    "이력서 PDF 합치는 중",
]

QUEUE_MAX = 15  # 확장: 5 → 15 (여러 작업 동시 처리)
HOMEPAGE  = "https://www.bridgejob.co.kr"

# ── 파일 포맷 규칙 ─────────────────────────────────────────────────────────
# 형식: {ID}_{국가}_{성별}({생년}).pdf
# 예: 5325미국_여성(96born).pdf
FILE_FORMAT_PATTERN = re.compile(
    r"^(\d+)_([가-힣\w]+)_(남성|여성|기타)\((\d{2}born)\)\.pdf$", re.IGNORECASE
)

logging.basicConfig(level=logging.INFO)
log = logging.getLogger("main_gui")


@dataclass
class QueueItem:
    """대기열 항목 — 강사 1인 분량의 파일 묶음."""
    queue_id:    str
    candidate_id: str = ""
    files:       dict = field(default_factory=dict)
    status:      str  = "대기"     # 대기 | 처리중 | 완료 | 오류
    error:       str  = ""
    info:        dict = field(default_factory=dict)


class BridgeConverterApp:
    """메인 GUI 앱 v2.0."""

    def __init__(self):
        if _DND_AVAILABLE:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()

        self.root.title("BRIDGE Converter")
        self.root.geometry("1440x900")
        self.root.minsize(1100, 700)
        self.root.configure(bg=C_BG)

        # ── 상태 ─────────────────────────────────────────────────────
        self._mode          = tk.StringVar(value="자동")
        self._candidate_id  = tk.StringVar()
        self._files: dict[str, Path] = {}            # pipeline용 (ftype → 마지막 파일)
        self._files_multi: dict[str, list] = {}      # 표시용 (ftype → [Path, ...])
        self._pii_result    = None
        self._current_step  = 0
        self._paused        = False
        self._stopped       = False
        self._pipeline_thread: Optional[threading.Thread] = None
        self._sheet_connected = False

        # 대기열
        self._queue: list[QueueItem] = []
        self._active_queue_id: Optional[str] = None
        self._queue_counter = 0

        # 애니메이션
        self._blink_base   = ""
        self._blink_dots   = 0
        self._blink_after: Optional[str] = None
        self._blink_error  = False   # True → 빨간 깜박

        # PII 패널 펼침 상태
        self._pii_expanded = False

        # 단계 실행 중 중복 방지
        self._step_running = False

        # 사진 크롭 결과 (bytes) — 단계 간 전달
        self._photo_bytes: Optional[bytes] = None
        # 사진이 이력서 본문 내장인지 여부 (True = 커버에 사진 추가 안 함)
        self._photo_from_resume: bool = False

        # 파일목록 그룹 헤더 (ftype → Treeview group iid)
        self._group_iids: dict[str, str] = {}

        # 인박스 감시 상태
        self._watcher_active = False
        self._watcher_observer = None

        # 출력 폴더 + 마지막 생성 파일
        self._output_dir: Path = self._load_output_config()
        self._last_output_path: Optional[Path] = None

        # 완료 블링크
        self._done_blink_after: Optional[str] = None

        # 강사 메타 (시트에서 로드)
        self._nationality = ""
        self._gender      = ""
        self._birth_year  = ""
        self._id_load_timer: Optional[str] = None

        self._configure_styles()
        self._build_ui()
        # 강사번호 변경 시 자동 시트 조회
        self._candidate_id.trace_add("write", self._on_id_change)
        self._check_sheet_connection()

    # ══════════════════════════════════════════════════════════════════
    # UI 빌드
    # ══════════════════════════════════════════════════════════════════

    def _build_ui(self):
        # ── 툴바 ──────────────────────────────────────────────────
        bar = tk.Frame(self.root, bg=C_SIDE, height=52,
                       highlightbackground=C_BORDER, highlightthickness=1)
        bar.pack(fill="x", side="top")
        bar.pack_propagate(False)

        tk.Label(bar, text="  BRIDGE Converter",
                 font=(F, 15, "bold"), fg=C_PRI, bg=C_SIDE
                 ).pack(side="left", padx=14, pady=10)

        # ── 모드 토글 (pill 버튼) ──────────────────────────────
        self._mode_btns: dict[str, tk.Button] = {}
        pill = tk.Frame(bar, bg=C_BORDER, highlightbackground=C_PRI, highlightthickness=1)
        pill.pack(side="left", padx=18, pady=13)
        for m in ["자동", "반자동"]:
            b = tk.Button(pill, text=m, font=(F, 11),
                          relief="flat", bd=0, padx=16, pady=3,
                          cursor="hand2", command=lambda x=m: self._set_mode(x))
            b.pack(side="left")
            self._mode_btns[m] = b
        self._update_mode_btns()

        # 홈페이지 버튼
        tk.Button(bar, text="홈페이지",
                  command=lambda: webbrowser.open(HOMEPAGE),
                  bg=C_SIDE, fg=C_PRI, font=(F, 11),
                  relief="flat", padx=8, cursor="hand2"
                  ).pack(side="left", padx=4)

        # 인박스 감시 버튼
        self._watcher_btn = tk.Button(
            bar, text="● 인박스 감시",
            command=self._toggle_inbox_watcher,
            bg=C_BORDER, fg=C_TEXT, font=(F, 11),
            relief="flat", padx=8, cursor="hand2")
        self._watcher_btn.pack(side="left", padx=4)

        # 시트 상태 (우측)
        self._sheet_label = tk.Label(
            bar, text="● 시트 미연결",
            font=(F, 11), fg=C_DANGER, bg=C_SIDE)
        self._sheet_label.pack(side="right", padx=14)

        # ── 3패널 메인 ────────────────────────────────────────
        main = tk.Frame(self.root, bg=C_BG)
        main.pack(fill="both", expand=True)

        self._left_panel = self._build_left(main)
        self._left_panel.pack(side="left", fill="y")

        tk.Frame(main, width=1, bg=C_BORDER).pack(side="left", fill="y")

        paned = ttk.PanedWindow(main, orient="horizontal")
        paned.pack(side="left", fill="both", expand=True)

        self._center_panel = self._build_center(paned)
        paned.add(self._center_panel, weight=3)

        self._right_panel = self._build_right(paned)
        paned.add(self._right_panel, weight=1)

        # ── 루트 창 전체 드래그드롭 (어디서나) ───────────────
        if _DND_AVAILABLE:
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind("<<Drop>>", self._on_drop)
            main.drop_target_register(DND_FILES)
            main.dnd_bind("<<Drop>>", self._on_drop)
            self._left_panel.drop_target_register(DND_FILES)
            self._left_panel.dnd_bind("<<Drop>>", self._on_drop)
            self._right_panel.drop_target_register(DND_FILES)
            self._right_panel.dnd_bind("<<Drop>>", self._on_drop)

    # ── 좌측 패널 ─────────────────────────────────────────────────────
    def _build_left(self, parent) -> tk.Frame:
        panel = tk.Frame(parent, bg=C_SIDE, width=270)
        panel.pack_propagate(False)

        # 시트 연결하기
        tk.Button(panel, text="  시트 연결하기",
                  command=self._open_sheet_settings,
                  bg=C_PRI, fg="white",
                  font=(F, 11, "bold"), relief="flat",
                  padx=10, pady=5, cursor="hand2"
                  ).pack(fill="x", padx=12, pady=(12, 2))

        # API 키 설정 (이름 제거에 필요)
        api_status = self._get_api_key_status()
        self._api_btn = tk.Button(
            panel,
            text=f"  Claude API  {api_status}",
            command=self._open_api_key_dialog,
            bg=C_HOVER if api_status == "✓" else C_QERR,
            fg=C_TEXT,
            font=(F, 10), relief="flat",
            padx=10, pady=3, cursor="hand2")
        self._api_btn.pack(fill="x", padx=12, pady=(0, 4))

        # ── 강사 카드 ──────────────────────────────────────────
        card = tk.LabelFrame(panel, text="강사 정보", font=(F, 11, "bold"),
                              bg=C_SIDE, fg=C_TEXT, padx=8, pady=6,
                              relief="flat",
                              highlightbackground=C_BORDER, highlightthickness=1)
        card.pack(fill="x", padx=12, pady=(0, 6))

        vcmd = (panel.register(self._validate_id), "%P")
        self._id_entry = tk.Entry(
            card, textvariable=self._candidate_id,
            font=(F, 13, "bold"), width=10,
            validate="key", validatecommand=vcmd,
            relief="flat", bd=1,
            highlightbackground=C_BORDER, highlightthickness=1)
        self._id_entry.grid(row=0, column=0, sticky="ew", pady=2)
        card.columnconfigure(0, weight=1)

        self._meta_label = tk.Label(
            card, text="", bg=C_SIDE, fg=C_SUB,
            font=(F, 10), wraplength=220, justify="left")
        self._meta_label.grid(row=1, column=0, sticky="w")

        # ── 대기열 ─────────────────────────────────────────────
        qf = tk.LabelFrame(panel, text=f"처리 대기열  (최대 {QUEUE_MAX}명)",
                            font=(F, 11, "bold"),
                            bg=C_SIDE, fg=C_TEXT, padx=6, pady=4,
                            relief="flat",
                            highlightbackground=C_BORDER, highlightthickness=1)
        qf.pack(fill="x", padx=12, pady=4)

        q_cols = ("번호", "강사ID", "파일", "상태")
        self._queue_tv = ttk.Treeview(qf, columns=q_cols,
                                       show="headings", height=5,
                                       selectmode="browse")
        self._queue_tv.heading("번호",   text="#")
        self._queue_tv.heading("강사ID", text="강사번호")
        self._queue_tv.heading("파일",   text="파일")
        self._queue_tv.heading("상태",   text="상태")
        self._queue_tv.column("번호",   width=28,  minwidth=26, anchor="center")
        self._queue_tv.column("강사ID", width=70,  minwidth=60, anchor="center")
        self._queue_tv.column("파일",   width=36,  minwidth=30, anchor="center")
        self._queue_tv.column("상태",   width=76,  minwidth=60, anchor="center")
        self._queue_tv.tag_configure("wait",   background=C_QWAIT)
        self._queue_tv.tag_configure("active", background=C_QACT)
        self._queue_tv.tag_configure("done",   background=C_QDONE)
        self._queue_tv.tag_configure("error",  background=C_QERR)
        self._queue_tv.pack(fill="x")

        qbr = tk.Frame(qf, bg=C_SIDE)
        qbr.pack(fill="x", pady=(4, 0))
        tk.Button(qbr, text="+ 추가",
                  command=self._queue_add,
                  bg=C_PRI, fg="white",
                  font=(F, 10), relief="flat",
                  padx=8, pady=2, cursor="hand2").pack(side="left")
        tk.Button(qbr, text="× 제거",
                  command=self._queue_remove,
                  bg=C_PAUSE, fg=C_DANGER,
                  font=(F, 10), relief="flat",
                  padx=8, pady=2, cursor="hand2").pack(side="left", padx=4)
        tk.Button(qbr, text="전체 시작",
                  command=self._queue_start,
                  bg=C_PRI, fg="white",
                  font=(F, 10, "bold"), relief="flat",
                  padx=8, pady=2, cursor="hand2").pack(side="right")

        # ── 처리 단계 ──────────────────────────────────────────
        sf = tk.LabelFrame(panel, text="처리 단계", font=(F, 11, "bold"),
                            bg=C_SIDE, fg=C_TEXT, padx=8, pady=4,
                            relief="flat",
                            highlightbackground=C_BORDER, highlightthickness=1)
        sf.pack(fill="x", padx=12, pady=4)

        self._step_rows: list[tuple[tk.Label, tk.Label]] = []
        for step in STEPS:
            rf = tk.Frame(sf, bg=C_SIDE)
            rf.pack(fill="x", pady=1)
            dot = tk.Label(rf, text="○", fg=C_BORDER, bg=C_SIDE, font=(F, 13))
            dot.pack(side="left")
            lbl = tk.Label(rf, text=step, bg=C_SIDE, fg=C_TEXT, font=(F, 11))
            lbl.pack(side="left", padx=4)
            self._step_rows.append((dot, lbl))

        # 현재 작업 깜박 표시
        self._status_var = tk.StringVar(value="")
        self._status_lbl = tk.Label(
            sf, textvariable=self._status_var,
            bg=C_SIDE, fg=C_WARN,
            font=(F, 10, "bold"), anchor="w", wraplength=220)
        self._status_lbl.pack(fill="x", pady=(4, 0))

        # 전체 진행 바
        tk.Label(panel, text="전체 진행:", bg=C_SIDE, fg=C_SUB,
                 font=(F, 10)).pack(padx=12, anchor="w", pady=(6, 1))
        self._progress = ttk.Progressbar(panel, mode="determinate",
                                          maximum=len(STEPS))
        self._progress.pack(padx=12, fill="x")

        # ── 미처리 목록 (에러 펼치기) ───────────────────────
        self._build_unproc(panel)

        return panel

    def _build_unproc(self, parent):
        uf = tk.LabelFrame(parent, text="미처리 목록", font=(F, 11, "bold"),
                            bg=C_SIDE, fg=C_TEXT, padx=6, pady=4,
                            relief="flat",
                            highlightbackground=C_BORDER, highlightthickness=1)
        uf.pack(fill="x", padx=12, pady=6)

        u_cols = ("ID", "이름", "오류")
        self._unproc_tv = ttk.Treeview(
            uf, columns=u_cols, show="tree headings",
            height=8, selectmode="browse")
        self._unproc_tv.heading("ID",   text="ID")
        self._unproc_tv.heading("이름", text="이름")
        self._unproc_tv.heading("오류", text="오류 내용")
        self._unproc_tv.column("#0",   width=18, minwidth=18, stretch=False)
        self._unproc_tv.column("ID",   width=42, minwidth=38, anchor="center")
        self._unproc_tv.column("이름", width=68, minwidth=60)
        self._unproc_tv.column("오류", width=116, minwidth=80)
        self._unproc_tv.tag_configure("detail",
                                       foreground=C_DANGER,
                                       background="#FFF3F3",
                                       font=(FM, 9))
        self._unproc_tv.pack(fill="x")
        self._unproc_tv.bind("<<TreeviewOpen>>",  self._unproc_expand)
        self._unproc_tv.bind("<<TreeviewClose>>", self._unproc_collapse)

        ubr = tk.Frame(uf, bg=C_SIDE)
        ubr.pack(fill="x", pady=(4, 0))
        tk.Button(ubr, text="새로고침",
                  command=self._refresh_unproc,
                  bg=C_BG, fg=C_TEXT, font=(F, 10),
                  relief="flat", cursor="hand2").pack(side="left")
        tk.Button(ubr, text="선택 불러오기",
                  command=self._load_selected_unproc,
                  bg=C_HOVER, fg=C_PRI, font=(F, 10),
                  relief="flat", cursor="hand2").pack(side="left", padx=4)

    # ── 가운데 패널 ────────────────────────────────────────────────────
    def _build_center(self, parent) -> tk.Frame:
        panel = tk.Frame(parent, bg=C_BG)

        # 드롭 영역 (고정 높이 — 드래그 대상 아님)
        dz = tk.Frame(panel, bg=C_HOVER, height=100,
                       highlightbackground=C_PRI, highlightthickness=2,
                       cursor="hand2")
        dz.pack(fill="x", padx=14, pady=(14, 6))
        dz.pack_propagate(False)
        dl = tk.Label(dz,
                      text="파일을 드래그하거나 클릭하여 추가  ·  사진 / 이력서 / 커버레터 / 추천서",
                      bg=C_HOVER, fg=C_PRI, font=(F, 12), justify="center")
        dl.pack(expand=True)
        dz.bind("<Button-1>", self._browse_files)
        dl.bind("<Button-1>", self._browse_files)
        if _DND_AVAILABLE:
            dz.drop_target_register(DND_FILES)
            dz.dnd_bind("<<Drop>>", self._on_drop)

        # 하단 버튼 (항상 고정)
        btf = tk.Frame(panel, bg=C_BG,
                        highlightbackground=C_BORDER, highlightthickness=1)
        btf.pack(fill="x", padx=14, pady=8, side="bottom")

        self._prev_btn = tk.Button(btf, text="◀ 이전",
                                    command=self._prev_step,
                                    bg=C_BORDER, fg=C_TEXT,
                                    font=(F, 11), relief="flat",
                                    padx=10, pady=6, cursor="hand2")
        self._prev_btn.pack(side="left", padx=(8, 2), pady=6)

        tk.Button(btf, text="건너뜀",
                  command=self._skip_step,
                  bg=C_BORDER, fg=C_TEXT,
                  font=(F, 11), relief="flat",
                  padx=10, pady=6, cursor="hand2"
                  ).pack(side="left", padx=2, pady=6)

        self._pause_btn = tk.Button(btf, text="일시정지",
                                     command=self._toggle_pause,
                                     bg=C_PAUSE, fg=C_DANGER,
                                     font=(F, 11, "bold"), relief="flat",
                                     padx=10, pady=6, cursor="hand2")
        self._pause_btn.pack(side="left", padx=2, pady=6)

        self._stop_btn = tk.Button(btf, text="중단하기",
                                    command=self._force_stop,
                                    bg=C_STOP, fg="white",
                                    font=(F, 11, "bold"), relief="flat",
                                    padx=10, pady=6, cursor="hand2")
        self._stop_btn.pack(side="left", padx=2, pady=6)

        self._next_btn = tk.Button(btf, text="다음 단계  ▶",
                                    command=self._next_step,
                                    bg=C_PRI, fg="white",
                                    font=(F, 12, "bold"), relief="flat",
                                    padx=14, pady=6, cursor="hand2")
        self._next_btn.pack(side="right", padx=8, pady=6)

        # ── 세로 PanedWindow: 파일목록 ↕ PII영역 (드래그 자유) ──
        vpaned = ttk.PanedWindow(panel, orient="vertical")
        vpaned.pack(fill="both", expand=True, padx=14, pady=(0, 4))

        # 상단: 파일 목록 (드래그로 크기 조절)
        top_pane = tk.Frame(vpaned, bg=C_BG)
        vpaned.add(top_pane, weight=3)

        flf = tk.LabelFrame(top_pane, text="파일 목록", font=(F, 11, "bold"),
                             bg=C_BG, fg=C_TEXT, padx=6, pady=4,
                             relief="flat",
                             highlightbackground=C_BORDER, highlightthickness=1)
        flf.pack(fill="both", expand=True)

        # ── 그룹형 파일 목록 (5그룹, 각 그룹 여러 파일 가능) ──
        f_cols = ("파일명", "크기", "상태")
        self._file_tv = ttk.Treeview(flf, columns=f_cols,
                                      show="tree headings",
                                      selectmode="browse",
                                      height=20)  # 확장: 9 → 20 (여러 파일 표시)
        self._file_tv.heading("#0",    text="구분")
        self._file_tv.heading("파일명", text="파일명")
        self._file_tv.heading("크기",   text="크기")
        self._file_tv.heading("상태",   text="상태")
        self._file_tv.column("#0",    width=110, minwidth=90,  stretch=False)
        self._file_tv.column("파일명", width=240, minwidth=120)
        self._file_tv.column("크기",   width=68,  minwidth=56,  anchor="center")
        self._file_tv.column("상태",   width=86,  minwidth=64,  anchor="center")
        # 그룹 헤더 배경색
        self._file_tv.tag_configure("grp_photo",   background="#FFF3CD", font=(F, 11, "bold"))
        self._file_tv.tag_configure("grp_resume",  background="#D4EDDA", font=(F, 11, "bold"))
        self._file_tv.tag_configure("grp_cover",   background="#CCE5FF", font=(F, 11, "bold"))
        self._file_tv.tag_configure("grp_rec",     background="#F8D7DA", font=(F, 11, "bold"))
        self._file_tv.tag_configure("grp_unknown", background="#E2E3E5", font=(F, 11, "bold"))
        # 파일 행 배경색
        self._file_tv.tag_configure("photo",   background="#FFFDE7")
        self._file_tv.tag_configure("resume",  background="#F1F8F1")
        self._file_tv.tag_configure("cover",   background="#EDF5FF")
        self._file_tv.tag_configure("rec",     background="#FFF0F3")
        self._file_tv.tag_configure("unknown", background="#FAFAFA")
        self._file_tv.bind("<Delete>", self._delete_selected_file)

        sb = ttk.Scrollbar(flf, orient="vertical", command=self._file_tv.yview)
        self._file_tv.configure(yscrollcommand=sb.set)
        self._file_tv.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")

        # ── 파일 목록 전체에 드래그드롭 등록 ──────────────────────────────────
        if _DND_AVAILABLE:
            self._file_tv.drop_target_register(DND_FILES)
            self._file_tv.dnd_bind("<<Drop>>", self._on_drop)
            flf.drop_target_register(DND_FILES)
            flf.dnd_bind("<<Drop>>", self._on_drop)

        # 5개 그룹 헤더 생성 (항상 펼침)
        _groups = [
            ("photo",   "📷 사진"),
            ("resume",  "📄 이력서"),
            ("cover",   "📝 커버레터"),
            ("rec",     "📋 추천서"),
            ("unknown", "📎 기타"),
        ]
        for ftype, label in _groups:
            iid = self._file_tv.insert("", "end",
                                        text=label,
                                        values=("", "", ""),
                                        tags=(f"grp_{ftype}",),
                                        open=True)
            self._group_iids[ftype] = iid
            self._files_multi[ftype] = []

        tk.Button(flf, text="선택 파일 제거 (Del)",
                  command=self._delete_selected_file,
                  bg=C_BG, fg=C_SUB, font=(F, 11),
                  relief="flat", cursor="hand2").pack(anchor="e", pady=(2, 0))

        # 하단: PII 탐지 결과 (드래그로 크기 조절)
        bot_pane = tk.Frame(vpaned, bg=C_BG)
        vpaned.add(bot_pane, weight=1)

        pii_hdr = tk.Frame(bot_pane, bg=C_BG)
        pii_hdr.pack(fill="x", pady=(4, 0))

        self._pii_toggle_btn = tk.Button(
            pii_hdr, text="  PII 탐지 결과 보기",
            command=self._toggle_pii,
            bg=C_BG, fg=C_SUB, font=(F, 10),
            relief="flat", cursor="hand2", anchor="w")
        self._pii_toggle_btn.pack(side="left", fill="x", expand=True)

        self._focus_btn = tk.Button(
            pii_hdr, text="집중 점검",
            command=self._focus_check,
            bg=C_WARN, fg="white", font=(F, 10),
            relief="flat", padx=6, pady=1,
            cursor="hand2", state="disabled")
        self._focus_btn.pack(side="right")

        self._pii_container = tk.Frame(bot_pane, bg=C_BG)
        # 기본 숨김

        self._pii_text = scrolledtext.ScrolledText(
            self._pii_container, font=(FM, 11),
            bg="#FAFAFA", relief="flat", selectbackground=C_HOVER)
        self._pii_text.pack(fill="both", expand=True, pady=(0, 4))
        self._pii_text.tag_config("red",    foreground=C_DANGER, background="#FEF2F2")
        self._pii_text.tag_config("yellow", foreground=C_WARN,   background="#FFFBEB")
        self._pii_text.tag_config("green",  foreground=C_PRI,    background="#F0FDF4")
        self._pii_text.bind("<<Selection>>", self._on_pii_select)

        return panel

    # ── 우측 패널 (축소: 360→280px) ──────────────────────────────────────────
    def _build_right(self, parent) -> tk.Frame:
        panel = tk.Frame(parent, bg=C_SIDE, width=280)  # 확장: 미리보기 축소
        panel.pack_propagate(False)

        # ── PDF 미리보기 (꽉 채움) ─────────────────────────────────
        prf = tk.LabelFrame(panel, text="미리보기 (이력서/사진)", font=(F, 11, "bold"),
                             bg=C_SIDE, fg=C_TEXT, padx=8, pady=6,
                             relief="flat",
                             highlightbackground=C_BORDER, highlightthickness=1)
        prf.pack(fill="both", expand=True, padx=12, pady=(14, 4))

        self._preview_lbl = tk.Label(prf, bg="#E0E0E0",
                                      text="파일 추가 시\n자동 표시",
                                      fg=C_SUB, font=(F, 11))
        self._preview_lbl.pack(fill="both", expand=True)

        # ── 파일 정보 (컴팩트) ────────────────────────────────────
        mf = tk.Frame(panel, bg=C_SIDE, padx=12, pady=2)
        mf.pack(fill="x")

        self._meta_info_labels = {}
        info_row = tk.Frame(mf, bg=C_SIDE)
        info_row.pack(fill="x")
        for key, label in [("size", "용량"), ("pages", "페이지"), ("pii_count", "삭제항목")]:
            tk.Label(info_row, text=f"{label}:", bg=C_SIDE, fg=C_SUB,
                     font=(F, 10), anchor="w").pack(side="left", padx=(0, 2))
            lbl = tk.Label(info_row, text="—", bg=C_SIDE, fg=C_TEXT,
                           font=(F, 10, "bold"))
            lbl.pack(side="left", padx=(0, 8))
            self._meta_info_labels[key] = lbl

        self._fname_lbl = tk.Label(mf, text="", bg=C_SIDE,
                                    fg=C_TEXT, font=(F, 10, "bold"),
                                    wraplength=320, justify="left")
        self._fname_lbl.pack(anchor="w")

        # ── 저장 완료 블링크 라벨 ─────────────────────────────────
        self._save_done_lbl = tk.Label(
            panel, text="", bg=C_SIDE, fg=C_PRI,
            font=(F, 13, "bold"), pady=4)
        self._save_done_lbl.pack(fill="x", padx=12)

        # ── 출력 폴더 위젯 ────────────────────────────────────────
        off = tk.LabelFrame(panel, text="저장 폴더", font=(F, 11, "bold"),
                             bg=C_SIDE, fg=C_TEXT, padx=8, pady=6,
                             relief="flat",
                             highlightbackground=C_BORDER, highlightthickness=1)
        off.pack(fill="x", padx=12, pady=(4, 6))

        self._outdir_lbl = tk.Label(
            off, text=str(self._output_dir),
            bg=C_SIDE, fg=C_TEXT, font=(F, 9),
            wraplength=290, justify="left", anchor="w")
        self._outdir_lbl.pack(fill="x")

        obr = tk.Frame(off, bg=C_SIDE)
        obr.pack(fill="x", pady=(4, 0))
        tk.Button(obr, text="폴더 변경",
                  command=self._pick_output_dir,
                  bg=C_BORDER, fg=C_TEXT, font=(F, 10),
                  relief="flat", padx=6, cursor="hand2").pack(side="left")
        tk.Button(obr, text="폴더 열기 →",
                  command=self._open_output_dir,
                  bg=C_PRI, fg="white", font=(F, 10, "bold"),
                  relief="flat", padx=6, cursor="hand2").pack(side="left", padx=4)
        tk.Button(obr, text="최근 파일 열기",
                  command=self._open_last_output,
                  bg=C_HOVER, fg=C_PRI, font=(F, 10),
                  relief="flat", padx=6, cursor="hand2").pack(side="right")

        # ── 중복/재지원 카드 (숨김 상태) ─────────────────────────
        self._dup_card  = tk.Frame(panel, bg="#FFF8E1",
                                    highlightbackground=C_WARN, highlightthickness=1)
        self._dup_title = tk.Label(self._dup_card, text="⚠ 중복 제출 감지",
                                    bg="#FFF8E1", fg=C_WARN, font=(F, 11, "bold"))
        self._dup_body  = tk.Label(self._dup_card, text="", bg="#FFF8E1",
                                    fg=C_TEXT, font=(F, 11),
                                    wraplength=240, justify="left")
        self._dup_btn   = tk.Button(self._dup_card, text="비교",
                                     bg=C_WARN, fg="white", font=(F, 11),
                                     relief="flat", padx=6, cursor="hand2")

        self._reapp_card  = tk.Frame(panel, bg="#FEF2F2",
                                      highlightbackground=C_DANGER, highlightthickness=1)
        self._reapp_title = tk.Label(self._reapp_card, text="재지원 감지",
                                      bg="#FEF2F2", fg=C_DANGER,
                                      font=(F, 11, "bold"))
        self._reapp_body  = tk.Label(self._reapp_card, text="", bg="#FEF2F2",
                                      fg=C_TEXT, font=(F, 11),
                                      wraplength=240, justify="left")
        btn_row = tk.Frame(self._reapp_card, bg="#FEF2F2")
        self._reapp_view_btn = tk.Button(btn_row, text="이전 파일 보기",
                                          bg=C_BG, fg=C_DANGER, font=(F, 11),
                                          relief="flat", padx=4, cursor="hand2")
        self._reapp_del_btn  = tk.Button(btn_row, text="삭제",
                                          bg=C_DANGER, fg="white", font=(F, 11),
                                          relief="flat", padx=4, cursor="hand2")
        self._reapp_keep_btn = tk.Button(btn_row, text="유지",
                                          bg=C_BG, fg=C_TEXT, font=(F, 11),
                                          relief="flat", padx=4, cursor="hand2")
        for b in [self._reapp_view_btn, self._reapp_del_btn, self._reapp_keep_btn]:
            b.pack(side="left", padx=2)

        return panel

    # ══════════════════════════════════════════════════════════════════
    # 스타일 설정 (Treeview 폰트 크게)
    # ══════════════════════════════════════════════════════════════════

    def _configure_styles(self):
        style = ttk.Style()
        style.configure("Treeview",
                        font=(F, 12),
                        rowheight=28)
        style.configure("Treeview.Heading",
                        font=(F, 12, "bold"))

    # ══════════════════════════════════════════════════════════════════
    # 파일 상태 헬퍼
    # ══════════════════════════════════════════════════════════════════

    def _set_file_status(self, ftype: str | None, status: str):
        """그룹 내 모든 파일 행의 상태 컬럼 업데이트. ftype=None 이면 전체."""
        groups = [ftype] if ftype else list(self._group_iids.keys())
        for ft in groups:
            gid = self._group_iids.get(ft)
            if not gid:
                continue
            for child in self._file_tv.get_children(gid):
                vals = list(self._file_tv.item(child, "values"))
                if len(vals) >= 3:
                    vals[2] = status
                    self._file_tv.item(child, values=tuple(vals))

    # ══════════════════════════════════════════════════════════════════
    # 모드 토글
    # ══════════════════════════════════════════════════════════════════

    def _set_mode(self, mode: str):
        self._mode.set(mode)
        self._update_mode_btns()

    def _update_mode_btns(self):
        sel = self._mode.get()
        for m, btn in self._mode_btns.items():
            if m == sel:
                btn.config(bg=C_PRI, fg="white")
            else:
                btn.config(bg=C_SIDE, fg=C_PRI)

    # ══════════════════════════════════════════════════════════════════
    # 보안: candidate_id 검증 — 숫자만 / 최대 8자리
    # ══════════════════════════════════════════════════════════════════

    def _validate_id(self, value: str) -> bool:
        return value == "" or (value.isdigit() and len(value) <= 8)

    # ══════════════════════════════════════════════════════════════════
    # 파일 핸들러
    # ══════════════════════════════════════════════════════════════════

    def _browse_files(self, event=None):
        paths = filedialog.askopenfilenames(
            title="파일 선택",
            filetypes=[("지원 파일", "*.pdf *.docx *.doc *.jpg *.jpeg *.png *.webp"),
                       ("모든 파일", "*.*")])
        if paths:
            self._add_files([Path(p) for p in paths])

    def _on_drop(self, event):
        if _DND_AVAILABLE:
            self._add_files([Path(p) for p in self.root.tk.splitlist(event.data)])

    def _add_files(self, paths: list):
        from .file_classifier import detect_file_type, extract_number
        candidate_num_detected = None
        for p in paths:
            if not p.exists():
                continue
            ftype = detect_file_type(p)
            # 파일명에서 강사번호 자동 추출
            num = extract_number(p.name)
            if num and not candidate_num_detected:
                candidate_num_detected = num
            # 파이프라인용 (마지막 파일로 덮어씀)
            self._files[ftype] = p
            # 다중 파일 목록
            if ftype not in self._files_multi:
                self._files_multi[ftype] = []
            # 중복 방지 (같은 이름 파일 재추가 시 갱신)
            existing = [i for i, q in enumerate(self._files_multi[ftype])
                        if q.name == p.name]
            if existing:
                self._files_multi[ftype][existing[0]] = p
            else:
                self._files_multi[ftype].append(p)

            size_b   = p.stat().st_size
            size_str = (f"{size_b}B"            if size_b < 1024 else
                        f"{size_b // 1024}KB"   if size_b < 1024*1024 else
                        f"{size_b // (1024*1024):.1f}MB")

            # 그룹 헤더 아래 자식 행 추가 또는 갱신
            gid = self._group_iids.get(ftype, self._group_iids.get("unknown"))
            if gid is None:
                continue
            # 이미 같은 이름의 자식 행이 있으면 갱신, 없으면 추가
            updated = False
            for child in self._file_tv.get_children(gid):
                cv = self._file_tv.item(child, "values")
                if cv and cv[0] == p.name:
                    self._file_tv.item(child,
                                       values=(p.name, size_str, "대기"),
                                       tags=(ftype,))
                    updated = True
                    break
            if not updated:
                self._file_tv.insert(gid, "end",
                                     text="",
                                     values=(p.name, size_str, "대기"),
                                     tags=(ftype,))
            # 그룹 펼치기
            self._file_tv.item(gid, open=True)
        # 강사번호 자동 채우기 (비어 있을 때만)
        if candidate_num_detected and not self._candidate_id.get().strip():
            self._candidate_id.set(candidate_num_detected)
        self._update_fname_preview()
        # 로드된 PDF가 있으면 첫 페이지 미리보기 (백그라운드)
        for ftype in ["resume", "cover", "rec"]:
            p = self._files.get(ftype)
            if p and p.suffix.lower() == ".pdf":
                threading.Thread(target=lambda pp=p: self._render_pdf_page_preview(pp),
                                 daemon=True).start()
                break

    def _delete_selected_file(self, event=None):
        sel = self._file_tv.selection()
        if not sel:
            return
        iid = sel[0]
        # 그룹 헤더는 삭제 불가
        if iid in self._group_iids.values():
            return
        # 부모 그룹에서 ftype 역조회
        parent_iid = self._file_tv.parent(iid)
        ftype = None
        for ft, gid in self._group_iids.items():
            if gid == parent_iid:
                ftype = ft
                break
        # 파일 이름 조회
        vals = self._file_tv.item(iid, "values")
        fname = vals[0] if vals else ""
        # _files_multi 에서 제거
        if ftype and ftype in self._files_multi:
            self._files_multi[ftype] = [
                q for q in self._files_multi[ftype] if q.name != fname
            ]
            # 파이프라인 dict 갱신 (마지막 남은 파일로 대체)
            if self._files_multi[ftype]:
                self._files[ftype] = self._files_multi[ftype][-1]
            elif ftype in self._files:
                del self._files[ftype]
        # Treeview 행 삭제
        self._file_tv.delete(iid)

    # ══════════════════════════════════════════════════════════════════
    # 대기열
    # ══════════════════════════════════════════════════════════════════

    def _queue_add(self):
        if len(self._queue) >= QUEUE_MAX:
            messagebox.showwarning("대기열 가득 참",
                                   f"최대 {QUEUE_MAX}명까지만 추가 가능합니다.")
            return
        cid = self._candidate_id.get().strip()
        if not cid:
            messagebox.showwarning("입력 필요", "강사 번호를 입력하세요.")
            return
        if not self._files:
            messagebox.showwarning("파일 없음", "먼저 파일을 추가하세요.")
            return

        self._queue_counter += 1
        item = QueueItem(
            queue_id=f"Q{self._queue_counter:02d}",
            candidate_id=cid,
            files=dict(self._files),
            status="대기",
        )
        self._queue.append(item)
        self._refresh_queue_tv()

        # 입력 초기화
        self._candidate_id.set("")
        self._files.clear()
        self._photo_bytes = None
        # 그룹 내 모든 자식 행 삭제
        for ft, gid in self._group_iids.items():
            for child in self._file_tv.get_children(gid):
                self._file_tv.delete(child)
            self._files_multi[ft] = []
        self._meta_label.config(text="")
        messagebox.showinfo("추가 완료",
                             f"강사 {cid}번이 {len(self._queue)}번 슬롯에 추가됐습니다.")

    def _queue_remove(self):
        sel = self._queue_tv.selection()
        if not sel:
            return
        all_rows = list(self._queue_tv.get_children(""))
        idx = all_rows.index(sel[0]) if sel[0] in all_rows else -1
        if 0 <= idx < len(self._queue):
            del self._queue[idx]
        self._refresh_queue_tv()

    def _queue_start(self):
        if not self._queue:
            messagebox.showwarning("대기열 비어 있음", "처리할 항목이 없습니다.")
            return
        if self._pipeline_thread and self._pipeline_thread.is_alive():
            messagebox.showwarning("처리 중", "이미 처리가 진행 중입니다.")
            return
        self._stopped = False
        self._paused  = False
        self._pipeline_thread = threading.Thread(
            target=self._run_queue_pipeline, daemon=True)
        self._pipeline_thread.start()

    def _run_queue_pipeline(self):
        for item in self._queue:
            if self._stopped:
                break
            if item.status != "대기":
                continue
            item.status = "처리중"
            self._active_queue_id = item.queue_id
            self.root.after(0, self._refresh_queue_tv)
            try:
                self._process_item(item)
                item.status = "완료"
            except Exception as e:
                item.status = "오류"
                item.error  = str(e)
                log.error(f"Queue [{item.candidate_id}]: {e}")
            self.root.after(0, self._refresh_queue_tv)
        self._active_queue_id = None
        self.root.after(0, lambda: self._status_var.set("모든 항목 처리 완료"))

    def _process_item(self, item: QueueItem):
        """대기열 항목 처리 — 동기 실행 (worker thread 내부)."""
        self._files = item.files
        self.root.after(0, lambda: self._candidate_id.set(item.candidate_id))
        cid = item.candidate_id

        for step_idx in range(len(STEPS)):
            while self._paused and not self._stopped:
                time.sleep(0.3)
            if self._stopped:
                raise InterruptedError("강제 중단됨")
            self._current_step = step_idx
            self.root.after(0, self._update_step_ui)
            self.root.after(0, lambda s=STEP_ACTIONS[step_idx]: self._start_blink(s))
            self.root.after(0, lambda: self._set_file_status(None, "처리중..."))

            # 각 단계 직접 실행 (thread 내부에서 동기 호출)
            if step_idx == 1:
                self._exec_face_crop()
            elif step_idx == 2:
                self._exec_pii()
            elif step_idx == 3:
                self._exec_build_pdf(cid)
            # step 0 (파일 확인), step 4 (저장): 별도 로직 없음

            self.root.after(0, lambda si=step_idx: self._set_file_status(None, "✓ 완료"))
            time.sleep(0.4)

        self.root.after(0, self._stop_blink)
        self.root.after(0, lambda: (
            self._status_var.set("✓ 전체 처리 완료"),
            self._status_lbl.config(fg=C_PRI),
        ))
        self.root.after(200, self._run_save)

    def _refresh_queue_tv(self):
        for row in self._queue_tv.get_children():
            self._queue_tv.delete(row)
        s_icons = {"대기": "대기", "처리중": "처리중", "완료": "완료", "오류": "오류"}
        tag_map  = {"대기": "wait", "처리중": "active", "완료": "done", "오류": "error"}
        for i, item in enumerate(self._queue, 1):
            tag = tag_map.get(item.status, "wait")
            self._queue_tv.insert("", "end",
                                   values=(i, item.candidate_id,
                                           len(item.files),
                                           s_icons.get(item.status, item.status)),
                                   tags=(tag,))

    # ══════════════════════════════════════════════════════════════════
    # 애니메이션 (깜박 점)
    # ══════════════════════════════════════════════════════════════════

    def _start_blink(self, text: str):
        """초록 깜박 — 정상 처리 중."""
        self._stop_blink()
        self._blink_base  = text
        self._blink_dots  = 0
        self._blink_error = False
        self._do_blink()

    def _start_blink_error(self, text: str):
        """빨간 깜박 — 오류 발생."""
        self._stop_blink()
        self._blink_base  = text
        self._blink_dots  = 0
        self._blink_error = True
        self._do_blink()

    def _do_blink(self):
        dots = "." * (self._blink_dots % 4)
        icon = "✖" if self._blink_error else "⚙"
        self._status_var.set(f"{icon}  {self._blink_base}{dots}")
        self._blink_dots += 1
        color = C_DANGER if self._blink_error else C_WARN
        self._status_lbl.config(fg=color)
        self._blink_after = self.root.after(400, self._do_blink)

    def _stop_blink(self):
        if self._blink_after:
            self.root.after_cancel(self._blink_after)
            self._blink_after = None

    # ══════════════════════════════════════════════════════════════════
    # 컨트롤 버튼
    # ══════════════════════════════════════════════════════════════════

    def _toggle_pause(self):
        self._paused = not self._paused
        if self._paused:
            self._pause_btn.config(text="계속 진행")
            self._status_var.set("일시정지됨")
        else:
            self._pause_btn.config(text="일시정지")
            if self._blink_base:
                self._start_blink(self._blink_base)

    def _force_stop(self):
        ok = messagebox.askyesno("강제 중단",
                                  "처리를 중단하시겠습니까?\n현재 항목은 오류로 표시됩니다.")
        if ok:
            self._stopped = True
            self._paused  = False
            self._stop_blink()
            self._status_var.set("중단됨")
            for item in self._queue:
                if item.status == "처리중":
                    item.status = "오류"
                    item.error  = "강제 중단"
            self._refresh_queue_tv()

    def _next_step(self):
        if self._step_running:
            return
        if self._current_step < len(STEPS) - 1:
            self._current_step += 1
            self._update_step_ui()
            self._run_current_step()
        else:
            self._finish()

    def _prev_step(self):
        if self._step_running:
            return
        if self._current_step > 0:
            self._current_step -= 1
            self._update_step_ui()
            self._stop_blink()
            self._status_var.set("")

    def _skip_step(self):
        if self._step_running:
            return
        self._set_file_status(None, "건너뜀")
        self._stop_blink()
        self._status_var.set(f"↩ {STEPS[self._current_step]} 건너뜀")
        self._status_lbl.config(fg=C_SUB)
        if self._current_step < len(STEPS) - 1:
            self._current_step += 1
            self._update_step_ui()

    # ══════════════════════════════════════════════════════════════════
    # PII 패널 토글
    # ══════════════════════════════════════════════════════════════════

    def _toggle_pii(self):
        if self._pii_expanded:
            self._pii_container.pack_forget()
            self._pii_toggle_btn.config(text="  PII 탐지 결과 보기")
            self._pii_expanded = False
        else:
            self._pii_container.pack(fill="both", expand=True)
            self._pii_toggle_btn.config(text="  PII 탐지 결과 숨기기")
            self._pii_expanded = True

    # ══════════════════════════════════════════════════════════════════
    # API 키 설정
    # ══════════════════════════════════════════════════════════════════

    _API_KEY_FILE = BASE_DIR / ".api_key"

    def _get_api_key_status(self) -> str:
        """API 키 상태: '✓' 또는 '미설정'."""
        from .pii_engine import load_api_key
        return "✓" if load_api_key() else "미설정"

    def _open_api_key_dialog(self) -> None:
        """Anthropic API 키 입력/수정 다이얼로그."""
        from tkinter import simpledialog
        from .pii_engine import load_api_key

        current = load_api_key() or ""
        hint = f"현재: {current[:8]}...{current[-4:]}" if current else "미설정 — 이름 제거 기능 사용 불가"

        dlg = tk.Toplevel(self.root)
        dlg.title("Claude API 키 설정")
        dlg.geometry("500x220")
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.configure(bg=C_SIDE)

        tk.Label(dlg, text="Anthropic API 키 입력", font=(F, 13, "bold"),
                 bg=C_SIDE, fg=C_TEXT).pack(pady=(18, 2))
        tk.Label(dlg, text=hint, font=(F, 10), bg=C_SIDE, fg=C_SUB).pack(pady=(0, 10))

        entry_var = tk.StringVar()
        entry = tk.Entry(dlg, textvariable=entry_var, font=(FM, 11),
                         width=45, show="•",
                         relief="flat", bd=1,
                         highlightbackground=C_BORDER, highlightthickness=1)
        entry.pack(padx=24, ipady=5)
        entry.focus()

        show_var = tk.BooleanVar(value=False)
        def _toggle_show():
            entry.config(show="" if show_var.get() else "•")
        tk.Checkbutton(dlg, text="키 표시", variable=show_var,
                       command=_toggle_show, bg=C_SIDE,
                       font=(F, 10)).pack(pady=4)

        def _save():
            key = entry_var.get().strip()
            if not key:
                messagebox.showwarning("입력 오류", "API 키를 입력하세요.", parent=dlg)
                return
            if not key.startswith("sk-ant-"):
                if not messagebox.askyesno("확인", "일반적인 Anthropic 키는 'sk-ant-'로 시작합니다.\n그래도 저장하시겠습니까?", parent=dlg):
                    return
            self._API_KEY_FILE.write_text(key, encoding="utf-8")
            status = self._get_api_key_status()
            self._api_btn.config(
                text=f"  Claude API  {status}",
                bg=C_HOVER if status == "✓" else C_QERR)
            messagebox.showinfo("저장 완료", f"API 키가 저장되었습니다.\n경로: {self._API_KEY_FILE}", parent=dlg)
            dlg.destroy()

        def _clear():
            if self._API_KEY_FILE.exists():
                self._API_KEY_FILE.unlink()
            self._api_btn.config(text="  Claude API  미설정", bg=C_QERR)
            messagebox.showinfo("삭제 완료", "API 키가 삭제되었습니다.", parent=dlg)
            dlg.destroy()

        btn_row = tk.Frame(dlg, bg=C_SIDE)
        btn_row.pack(pady=12)
        tk.Button(btn_row, text="저장", command=_save,
                  bg=C_PRI, fg="white", font=(F, 11, "bold"),
                  relief="flat", padx=20, pady=4, cursor="hand2"
                  ).pack(side="left", padx=8)
        tk.Button(btn_row, text="삭제", command=_clear,
                  bg=C_DANGER, fg="white", font=(F, 11),
                  relief="flat", padx=20, pady=4, cursor="hand2"
                  ).pack(side="left", padx=8)
        tk.Button(btn_row, text="취소", command=dlg.destroy,
                  bg=C_BORDER, fg=C_TEXT, font=(F, 11),
                  relief="flat", padx=20, pady=4, cursor="hand2"
                  ).pack(side="left", padx=8)

    # ══════════════════════════════════════════════════════════════════
    # 시트 연결
    # ══════════════════════════════════════════════════════════════════

    def _open_sheet_settings(self):
        try:
            from .vault_import import VaultConfig
            vc = VaultConfig()
            vc.open_dialog(self.root, on_saved=self._check_sheet_connection)
        except ImportError:
            self._simple_sheet_dialog()

    def _simple_sheet_dialog(self):
        """연결 설정 다이얼로그 — Bridge API (기본) + Google Sheets (폴백)."""
        dlg = tk.Toplevel(self.root)
        dlg.title("데이터 소스 연결 설정")
        dlg.geometry("480x420")
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.configure(bg=C_SIDE)

        try:
            import keyring
            _kr = keyring
        except ImportError:
            messagebox.showerror("keyring 없음",
                                  "pip install keyring 을 실행하세요.", parent=dlg)
            dlg.destroy()
            return

        # 현재 저장값 로드
        def _get(key, default=""):
            try:
                return _kr.get_password("BRIDGE_RC_CONFIG_V1", key) or default
            except Exception:
                return default

        source_var = tk.StringVar(value=_get("source", "bridge"))

        # ── 소스 선택 ──────────────────────────────────────────────
        src_frame = tk.LabelFrame(dlg, text="데이터 소스", font=(F, 11, "bold"),
                                   bg=C_SIDE, fg=C_TEXT, padx=12, pady=6)
        src_frame.pack(fill="x", padx=18, pady=(14, 6))
        for val, lbl in [("bridge", "Bridge 관리자 API  (권장 — 실시간 연동)"),
                          ("sheets", "Google Sheets  (폴백)")]:
            tk.Radiobutton(src_frame, text=lbl, variable=source_var, value=val,
                           bg=C_SIDE, fg=C_TEXT, selectcolor=C_SIDE,
                           font=(F, 11), activebackground=C_HOVER
                           ).pack(anchor="w")

        # ── Bridge API ─────────────────────────────────────────────
        ba_frame = tk.LabelFrame(dlg, text="Bridge API 설정", font=(F, 11, "bold"),
                                  bg=C_SIDE, fg=C_TEXT, padx=12, pady=6)
        ba_frame.pack(fill="x", padx=18, pady=4)

        tk.Label(ba_frame, text="관리자 비밀번호:", font=(F, 10),
                 bg=C_SIDE).pack(anchor="w")
        pw_var = tk.StringVar()
        tk.Entry(ba_frame, textvariable=pw_var, font=(F, 11), show="*", width=38
                 ).pack(fill="x", pady=(0, 4))
        tk.Label(ba_frame,
                 text="API URL은 운영 도메인(api.bridgejob.co.kr)이 자동 적용됩니다.",
                 font=(F, 9), fg=C_SUB, bg=C_SIDE
                 ).pack(anchor="w")

        # ── Google Sheets ──────────────────────────────────────────
        gs_frame = tk.LabelFrame(dlg, text="Google Sheets 설정 (폴백용)",
                                  font=(F, 11, "bold"),
                                  bg=C_SIDE, fg=C_TEXT, padx=12, pady=6)
        gs_frame.pack(fill="x", padx=18, pady=4)

        tk.Label(gs_frame, text="Spreadsheet ID:", font=(F, 10),
                 bg=C_SIDE).pack(anchor="w")
        sid_var = tk.StringVar(value=_get("sheet_id"))
        tk.Entry(gs_frame, textvariable=sid_var, font=(F, 11), width=38
                 ).pack(fill="x", pady=(0, 4))
        tk.Label(gs_frame, text="Service Account JSON Key:", font=(F, 10),
                 bg=C_SIDE).pack(anchor="w")
        key_var = tk.StringVar()
        tk.Entry(gs_frame, textvariable=key_var, font=(F, 11), show="*", width=38
                 ).pack(fill="x")

        def _save():
            source = source_var.get()
            pw     = pw_var.get().strip()
            sid    = sid_var.get().strip()
            key    = key_var.get().strip()

            if source == "bridge" and not pw:
                messagebox.showwarning("입력 필요", "관리자 비밀번호를 입력하세요.", parent=dlg)
                return
            if source == "sheets" and (not sid):
                messagebox.showwarning("입력 필요", "Spreadsheet ID를 입력하세요.", parent=dlg)
                return
            try:
                _kr.set_password("BRIDGE_RC_CONFIG_V1", "source", source)
                if pw:
                    _kr.set_password("BRIDGE_RC_CONFIG_V1", "admin_pw", pw)
                if sid:
                    _kr.set_password("BRIDGE_RC_CONFIG_V1", "sheet_id", sid)
                if key:
                    _kr.set_password("BRIDGE_RC_CONFIG_V1", "sa_json", key)
                dlg.destroy()
                self._check_sheet_connection()
            except Exception as e:
                messagebox.showerror("저장 실패", str(e), parent=dlg)

        tk.Button(dlg, text="저장 & 연결 테스트", command=_save,
                  bg=C_PRI, fg="white", font=(F, 11, "bold"),
                  relief="flat", padx=14, pady=5).pack(pady=10)

    def _check_sheet_connection(self):
        def _check():
            try:
                from .sheets_connector import is_connected
                ok = is_connected()
            except Exception:
                ok = False
            color = C_PRI   if ok else C_DANGER
            text  = "● 시트 연결됨" if ok else "● 시트 미연결"
            self._sheet_connected = ok
            self.root.after(0, lambda: self._sheet_label.config(text=text, fg=color))
        threading.Thread(target=_check, daemon=True).start()

    def _toggle_inbox_watcher(self):
        """인박스 폴더 감시 시작/중지 토글."""
        if not self._watcher_active:
            try:
                from .file_classifier import start_watcher
                self._watcher_observer = start_watcher(
                    on_duplicate_fn=lambda num, old, new: self.root.after(
                        0, lambda: self.show_duplicate_alert(num, old, new)),
                    on_reapply_fn=lambda num, old: self.root.after(
                        0, lambda: self.show_reapply_alert(num, old)),
                )
                self._watcher_active = True
                self._watcher_btn.config(
                    text="● 감시중 ON", bg=C_PRI, fg="white")
                self._status_var.set(f"✓ 인박스 감시 시작 — {INBOX_DIR}")
                self._status_lbl.config(fg=C_PRI)
            except Exception as e:
                messagebox.showerror("감시 오류", f"인박스 감시 시작 실패:\n{e}")
        else:
            try:
                from .file_classifier import stop_watcher
                stop_watcher()
            except Exception:
                pass
            self._watcher_active = False
            self._watcher_observer = None
            self._watcher_btn.config(
                text="● 인박스 감시", bg=C_BORDER, fg=C_TEXT)
            self._status_var.set("인박스 감시 중지됨")
            self._status_lbl.config(fg=C_SUB)

    # ══════════════════════════════════════════════════════════════════
    # 시트 / 미처리 목록
    # ══════════════════════════════════════════════════════════════════

    def _load_from_sheet(self):
        cid = self._candidate_id.get().strip()
        if not cid:
            return

        def _fetch():
            try:
                from .sheets_connector import get_candidate_info
                info = get_candidate_info(cid)
                if info:
                    nat = info.get("nationality", "")
                    gen = info.get("gender", "")
                    yr  = info.get("birth_year", "")
                    def _apply(i=info, n=nat, g=gen, y=yr):
                        self._nationality = n
                        self._gender      = g
                        self._birth_year  = y
                        self._meta_label.config(
                            text=(f"국적: {n or '?'}  "
                                  f"성별: {g or '?'}  "
                                  f"생년: {y or '?'}"))
                        self._update_fname_preview(i)
                    self.root.after(0, _apply)
                else:
                    self.root.after(0, lambda: self._meta_label.config(text="시트 미등록"))
            except Exception as e:
                _e = str(e)
                self.root.after(0, lambda: self._meta_label.config(text=f"조회 실패: {_e}"))

        threading.Thread(target=_fetch, daemon=True).start()

    def _on_id_change(self, *args):
        """강사번호 입력 시 0.8초 후 시트 자동 조회."""
        cid = self._candidate_id.get().strip()
        if self._id_load_timer:
            self.root.after_cancel(self._id_load_timer)
            self._id_load_timer = None
        # 4자리 이상 + 시트 연결된 경우에만 자동 조회
        if len(cid) >= 4 and self._sheet_connected:
            self._id_load_timer = self.root.after(800, self._load_from_sheet)
        # 번호 바뀌면 기존 메타 초기화
        if len(cid) != len(self._candidate_id.get().strip()):
            self._nationality = ""
            self._gender = ""
            self._birth_year = ""

    def _refresh_unproc(self):
        for row in self._unproc_tv.get_children():
            self._unproc_tv.delete(row)
        try:
            from .sheets_connector import get_unprocessed_rows
            rows = get_unprocessed_rows(20)
            for r in rows:
                err = r.get("error", "")
                self._unproc_tv.insert(
                    "", "end",
                    values=(r["id"], r.get("name", ""), err or "정상"),
                    open=False)
        except Exception as e:
            log.warning(f"미처리 목록 갱신 실패: {e}")

    def _unproc_expand(self, event):
        item = self._unproc_tv.focus()
        if not item:
            return
        children = self._unproc_tv.get_children(item)
        if children:
            return  # 이미 자식 있음
        vals = self._unproc_tv.item(item, "values")
        err = vals[2] if len(vals) > 2 else ""
        if err and err != "정상":
            self._unproc_tv.insert(item, "end",
                                    values=("", f"오류 상세: {err}", ""),
                                    tags=("detail",))

    def _unproc_collapse(self, event):
        item = self._unproc_tv.focus()
        if not item:
            return
        for c in self._unproc_tv.get_children(item):
            self._unproc_tv.delete(c)

    def _load_selected_unproc(self):
        sel = self._unproc_tv.selection()
        if not sel:
            return
        vals = self._unproc_tv.item(sel[0], "values")
        if vals:
            self._candidate_id.set(str(vals[0]))
            self._load_from_sheet()

    # ══════════════════════════════════════════════════════════════════
    # 파이프라인 단계
    # ══════════════════════════════════════════════════════════════════

    def _run_current_step(self):
        """단계 실행 — 백그라운드 스레드 + 깜박 애니메이션."""
        step = self._current_step
        if step == 0:
            self._set_file_status(None, "확인")
            self._status_var.set("✓ 파일 확인 완료")
            self._status_lbl.config(fg=C_PRI)
            return

        if self._step_running:
            return
        self._step_running = True
        self._next_btn.config(state="disabled")
        self._prev_btn.config(state="disabled")

        # 로드된 모든 파일을 "처리중..." 으로 표시 (어떤 단계든 진행 중임을 보여줌)
        ftypes = list(self._files.keys())
        self._set_file_status(None, "처리중...")

        self._start_blink(STEP_ACTIONS[step])
        cid = self._candidate_id.get().strip() or "0000"

        def _worker():
            err = None
            try:
                if step == 1:
                    self._exec_face_crop()
                elif step == 2:
                    self._exec_pii()
                elif step == 3:
                    self._exec_build_pdf(cid)
            except Exception as e:
                err = str(e)
            self.root.after(0, lambda e=err: self._on_step_done(step, ftypes, e))

        threading.Thread(target=_worker, daemon=True).start()

    def _on_step_done(self, step: int, ftypes: list, err: str | None):
        """단계 완료 콜백 — UI 스레드에서 실행."""
        self._step_running = False
        self._stop_blink()
        self._next_btn.config(state="normal")
        self._prev_btn.config(state="normal")

        last_step = len(STEPS) - 1  # 마지막 단계 인덱스 (PDF 생성)
        if err:
            self._start_blink_error(f"{STEPS[step]} 오류")
            self._set_file_status(None, "✖ 오류")
            messagebox.showerror("처리 오류", f"{STEPS[step]} 실패:\n{err}")
        else:
            self._set_file_status(None, "✓ 완료")
            if step == last_step:
                # PDF 생성 완료 = 전체 완료
                self._status_var.set("✓ PDF 생성 완료")
                self._status_lbl.config(fg=C_PRI)
                self._run_save()   # 진행바 꽉 + 블링크
            elif self._mode.get() == "자동" and self._current_step < last_step:
                self._status_var.set(f"✓ {STEPS[step]} 완료 — 자동 진행 중...")
                self._status_lbl.config(fg=C_PRI)
                self.root.after(800, self._next_step)
            else:
                next_hint = f" — 다음 단계 ▶ 클릭" if self._current_step < last_step else ""
                self._status_var.set(f"✓ {STEPS[step]} 완료{next_hint}")
                self._status_lbl.config(fg=C_PRI)

    def _exec_face_crop(self):
        """[Thread] 사진 크롭 — self._photo_bytes 에 결과 저장."""
        from .face_crop import crop_face, FaceNotFoundError
        photo = self._files.get("photo")
        self._photo_from_resume = False  # 매번 초기화

        if not photo:
            # 별도 사진 없음 → 이력서 본문 내장 사진 탐색
            resume = self._files.get("resume")
            if resume and resume.suffix.lower() == ".pdf":
                from .pipeline import _extract_photo_from_pdf
                data = _extract_photo_from_pdf(resume)
                if data:
                    self._photo_bytes = data
                    self._photo_from_resume = True  # 본문 내장 사진
                    self.root.after(0, lambda d=data: self._show_photo_preview(d))
                    self.root.after(0, lambda: (
                        self._status_var.set("ℹ 이력서 내장 사진 감지 — 커버에 사진 생략, ID만 표시"),
                        self._status_lbl.config(fg=C_WARN),
                    ))
            return
        try:
            data = crop_face(photo)
        except FaceNotFoundError:
            # 얼굴 미감지 → 원본 사진 그대로 사용 (파이프라인 계속)
            data = photo.read_bytes()
            self.root.after(0, lambda: (
                self._status_var.set("⚠ 얼굴 미감지 — 원본 사진 사용"),
                self._status_lbl.config(fg=C_WARN),
            ))
        self._photo_bytes = data
        self._photo_from_resume = False
        self.root.after(0, lambda d=data: self._show_photo_preview(d))

    def _show_photo_preview(self, data: bytes):
        """우측 패널 크기에 맞춰 미리보기 이미지 동적 스케일."""
        self._preview_lbl.update_idletasks()
        w = self._preview_lbl.winfo_width()
        h = self._preview_lbl.winfo_height()
        # 최소 보장 크기
        if w < 80:
            w = 280
        if h < 80:
            h = 380
        # 3:4 비율 유지 (세로가 더 큰 경우 우선)
        target_w = min(w - 8, int((h - 8) * 3 / 4))
        target_h = int(target_w * 4 / 3)
        img    = Image.open(io.BytesIO(data)).resize(
                     (max(target_w, 60), max(target_h, 80)), Image.LANCZOS)
        tk_img = ImageTk.PhotoImage(img)
        self._preview_lbl.configure(image=tk_img, text="")
        self._preview_lbl.image = tk_img

    def _exec_pii(self):
        """[Thread] PII 탐지."""
        from .pii_engine import analyze_pii, load_api_key
        from .pipeline import _extract_text
        api_key  = load_api_key()
        all_text = ""
        for ftype in ["cover", "resume", "rec"]:
            p = self._files.get(ftype)
            if p:
                all_text += f"\n=== {ftype.upper()} ===\n{_extract_text(p)}\n"
        if not all_text.strip():
            return
        result = analyze_pii(all_text, api_key)
        self._pii_result = result
        self.root.after(0, lambda r=result: self._show_pii_result(r))

    def _show_pii_result(self, result):
        self._pii_text.config(state="normal")
        self._pii_text.delete("1.0", "end")
        text = result.cleaned_text
        self._pii_text.insert("end", text)
        for item in result.pii_found:
            ph = f"[{item.type.upper()}_REMOVED]"
            st = text.find(ph)
            if st >= 0:
                self._pii_text.tag_add(item.color,
                                        f"1.0+{st}c",
                                        f"1.0+{st + len(ph)}c")
        for u in result.uncertain:
            ut = u.get("text", "")
            if ut in text:
                st = text.find(ut)
                self._pii_text.tag_add("yellow",
                                        f"1.0+{st}c",
                                        f"1.0+{st + len(ut)}c")
        self._pii_text.config(state="disabled")
        self._meta_info_labels["pii_count"].config(text=str(len(result.pii_found)))
        if not self._pii_expanded:
            self._toggle_pii()

    def _exec_build_pdf(self, cid: str):
        """[Thread] PDF 생성 — 원본 파일 보존 + PII 인라인 제거.

        핵심 원칙:
          - 새 이력서 생성 X → 제출된 원본 파일을 그대로 유지
          - cover → resume → rec 순서로 병합
          - PyMuPDF apply_redactions(fill=None) = 텍스트 스트림 완전 삭제 (흰박스 아님)
          - 이력서 첫 페이지 상단: 이름 헤더 완전 삭제 + 강사 번호만 삽입
        """
        import fitz  # PyMuPDF
        from .pdf_builder import build_filename

        # ── 1. meta 정보: 인스턴스 변수 → 시트 조회 순 ──────────────
        nat = self._nationality or ""
        gen = self._gender      or ""
        yr  = self._birth_year  or ""
        if not (nat or gen or yr):
            try:
                from .sheets_connector import get_candidate_info
                info = get_candidate_info(cid)
                if info:
                    nat = info.get("nationality", "")
                    gen = info.get("gender", "")
                    yr  = info.get("birth_year", "")
                    self._nationality = nat
                    self._gender      = gen
                    self._birth_year  = yr
                    self.root.after(0, lambda: self._meta_label.config(
                        text=f"국적: {nat or '?'}  성별: {gen or '?'}  생년: {yr or '?'}"))
            except Exception:
                pass

        # ── 2. 출력 파일명 결정 ──────────────────────────────────────
        fname    = build_filename(cid, nat, gen, yr)
        out_path = self._output_dir / fname
        self._output_dir.mkdir(parents=True, exist_ok=True)

        # ── 3. Claude API PII 목록 (있으면 보조로 사용) ─────────────
        api_pii: list[str] = []
        if self._pii_result:
            api_pii = [
                m.original_value.strip()
                for m in self._pii_result.pii_found
                if m.original_value and len(m.original_value.strip()) >= 3
            ]

        # ── 4. PDF 조립 (원본 파일 순서대로 PII 제거 후 병합) ────────
        merged = fitz.open()

        # ── 5. 원본 파일 보존 + PII 직접 제거 ───────────────────────
        for ftype in ["cover", "resume", "rec"]:
            p = self._files.get(ftype)
            if not p or not p.exists():
                continue

            if p.suffix.lower() == ".pdf":
                is_resume = (ftype == "resume")
                src_doc = fitz.open(str(p))
                for page_idx, page in enumerate(src_doc):
                    # ① 섹션 전체 삭제 (contact / references / 상단 헤더)
                    _redact_pii_sections_pdf(page, is_first_page=(page_idx == 0))
                    # ② 개별 PII 값 삭제 (이메일·전화 등 잔류 텍스트)
                    page_text = page.get_text()
                    local_pii = _regex_pii(page_text)
                    all_pii = list(dict.fromkeys(api_pii + local_pii))
                    _redact_pdf_page(page, all_pii)
                    # ③ 이력서 첫 페이지: 강사 번호만 상단 삽입
                    if is_resume and page_idx == 0:
                        _overlay_id_on_resume_page(page, cid, self._photo_bytes)
                merged.insert_pdf(src_doc)
                src_doc.close()

            elif p.suffix.lower() in (".docx", ".doc"):
                # DOCX: run-level inline PII 교체 (포맷 보존) → PDF 변환
                try:
                    from docx import Document as DocxDoc
                    from .pdf_builder import text_to_pdf_bytes
                    doc = DocxDoc(str(p))

                    # ① 섹션 전체 삭제 (contact / references / 상단 이름 헤더)
                    _clear_pii_sections_docx(doc)

                    # ② 전체 문서 텍스트에서 개별 PII 목록 수집
                    full_text = "\n".join(
                        p2.text for p2 in doc.paragraphs
                    )
                    for tbl in doc.tables:
                        for row in tbl.rows:
                            for cell in row.cells:
                                full_text += "\n" + cell.text
                    local_pii = _regex_pii(full_text)
                    all_pii = list(dict.fromkeys(api_pii + local_pii))

                    def _clean_run_text(t: str) -> str:
                        for pii in all_pii:
                            t = t.replace(pii, "")
                        return t

                    # 1) 단락 run-level 교체
                    for para in doc.paragraphs:
                        for run in para.runs:
                            if run.text:
                                run.text = _clean_run_text(run.text)

                    # 2) 테이블 셀 run-level 교체 (기존 누락 부분)
                    for tbl in doc.tables:
                        for row in tbl.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if run.text:
                                            run.text = _clean_run_text(run.text)

                    # 3) 수정된 DOCX → PDF 변환 (원본 포맷 보존 — docx2pdf/Word COM 사용)
                    import tempfile as _tf
                    import os as _os
                    import threading as _th
                    _tmp = _tf.NamedTemporaryFile(suffix=".docx", delete=False)
                    doc.save(_tmp.name)
                    _tmp.close()
                    _tmp_pdf = _tmp.name[:-5] + ".pdf"
                    docx_pdf_bytes = None
                    # 상태 업데이트: Word 변환 중 안내
                    self.root.after(0, lambda: self._status_var.set(
                        "⏳ DOCX → PDF 변환 중 (Word 실행, 최대 45초)..."))
                    # docx2pdf를 별도 스레드 + 타임아웃 45초로 실행
                    _conv_result: list = [None, None]  # [bytes, error_str]
                    _conv_event = _th.Event()
                    _tmp_name_cap = _tmp.name
                    _tmp_pdf_cap  = _tmp_pdf
                    def _do_d2p():
                        try:
                            import pythoncom
                            pythoncom.CoInitialize()
                            try:
                                from docx2pdf import convert as _d2p
                                _d2p(_tmp_name_cap, _tmp_pdf_cap)
                                with open(_tmp_pdf_cap, "rb") as _f:
                                    _conv_result[0] = _f.read()
                                try: _os.unlink(_tmp_pdf_cap)
                                except: pass
                            finally:
                                pythoncom.CoUninitialize()
                        except Exception as _e:
                            _conv_result[1] = str(_e)
                        finally:
                            _conv_event.set()
                    _th.Thread(target=_do_d2p, daemon=True).start()
                    _finished = _conv_event.wait(timeout=45)
                    try: _os.unlink(_tmp.name)
                    except: pass
                    if not _finished:
                        log.warning("docx2pdf 타임아웃(45s) — 평문 폴백 사용")
                    elif _conv_result[1]:
                        log.warning(f"docx2pdf 실패 — 평문 폴백 사용: {_conv_result[1]}")
                    else:
                        docx_pdf_bytes = _conv_result[0]
                    if not docx_pdf_bytes:
                        _cparas = [_p2.text for _p2 in doc.paragraphs]
                        for _t in doc.tables:
                            for _r in _t.rows:
                                _rc = [_c.text.strip() for _c in _r.cells if _c.text.strip()]
                                if _rc:
                                    _cparas.append("  |  ".join(_rc))
                        docx_pdf_bytes = text_to_pdf_bytes("\n".join(_cparas), ftype.upper())
                    src_doc = fitz.open("pdf", docx_pdf_bytes)
                    # 이력서 첫 페이지: 잔류 PII 삭제 + 강사 번호/사진 오버레이
                    _is_docx_resume = (ftype == "resume")
                    if _is_docx_resume and len(src_doc) > 0:
                        _redact_pii_sections_pdf(src_doc[0], is_first_page=True)
                        _redact_pdf_page(src_doc[0], all_pii)
                        _overlay_id_on_resume_page(src_doc[0], cid, self._photo_bytes)
                    merged.insert_pdf(src_doc)
                    src_doc.close()
                except Exception as e:
                    log.warning(f"DOCX 처리 실패: {e}")

        # ── 6. 저장 ─────────────────────────────────────────────────
        merged.save(str(out_path), deflate=True,
                    garbage=4, clean=True)
        size = out_path.stat().st_size
        merged.close()

        pages_count = 0
        try:
            tmp = fitz.open(str(out_path))
            pages_count = len(tmp)
            tmp.close()
        except Exception:
            pass

        self._last_output_path = out_path
        self.root.after(0, lambda o=out_path, s=size, pg=pages_count: (
            self._meta_info_labels["size"].config(text=f"{s // 1024}KB"),
            self._meta_info_labels["pages"].config(text=str(pg) if pg else "—"),
            self._fname_lbl.config(text=o.name),
        ))
        # 생성된 PDF 첫 페이지 미리보기
        self._render_pdf_page_preview(out_path)

    def _run_save(self):
        """저장 완료 상태: 진행바 꽉 채우기 + 블링크."""
        # 진행바 꽉 채우기
        self._progress.config(maximum=len(STEPS))
        self._progress["value"] = len(STEPS)
        # 블링크 시작
        self._save_done_blink_state = True
        self._do_save_done_blink()

    def _do_save_done_blink(self):
        """저장 완료 라벨 초록 블링크."""
        if hasattr(self, "_save_done_blink_state") and self._save_done_blink_state:
            cur = self._save_done_lbl.cget("text")
            if "✓" in cur:
                self._save_done_lbl.config(text="", bg=C_SIDE)
            else:
                self._save_done_lbl.config(
                    text="  ✓  저장 완료!", bg="#E8F5E9", fg=C_PRI)
            self._done_blink_after = self.root.after(600, self._do_save_done_blink)

    def _stop_save_done_blink(self):
        self._save_done_blink_state = False
        if self._done_blink_after:
            self.root.after_cancel(self._done_blink_after)
            self._done_blink_after = None
        self._save_done_lbl.config(text="  ✓  저장 완료!", bg="#E8F5E9", fg=C_PRI)

    def _finish(self):
        self._run_save()
        self._current_step = 0
        self._update_step_ui()

    # ══════════════════════════════════════════════════════════════════
    # PDF 페이지 미리보기 렌더링 (PyMuPDF)
    # ══════════════════════════════════════════════════════════════════

    def _render_pdf_page_preview(self, pdf_path: Path):
        """[Thread 안전] PyMuPDF로 PDF 첫 페이지 → 우측 패널 표시."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(pdf_path))
            if len(doc) == 0:
                doc.close()
                return
            page = doc[0]
            # A4 비율 (210:297 ≈ 0.707) 로 렌더
            self._preview_lbl.update_idletasks()
            pw = self._preview_lbl.winfo_width()
            ph = self._preview_lbl.winfo_height()
            if pw < 80: pw = 300
            if ph < 80: ph = 420
            # 패널 크기에 맞게 스케일 결정
            scale_x = pw / page.rect.width
            scale_y = ph / page.rect.height
            scale = min(scale_x, scale_y) * 0.95
            mat = fitz.Matrix(scale, scale)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("png")
            doc.close()
            self.root.after(0, lambda d=img_data: self._show_doc_preview(d))
        except Exception as e:
            log.warning(f"PDF 미리보기 실패: {e}")

    def _show_doc_preview(self, data: bytes):
        """A4 비율 문서 미리보기 표시."""
        try:
            img    = Image.open(io.BytesIO(data))
            tk_img = ImageTk.PhotoImage(img)
            self._preview_lbl.configure(image=tk_img, text="")
            self._preview_lbl.image = tk_img
        except Exception as e:
            log.warning(f"미리보기 표시 실패: {e}")

    # ══════════════════════════════════════════════════════════════════
    # 출력 폴더 관리
    # ══════════════════════════════════════════════════════════════════

    def _load_output_config(self) -> Path:
        """config.json에서 output_dir 읽기. 없으면 기본 output/ 폴더."""
        try:
            if CONFIG_PATH.exists():
                import json
                cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
                d = cfg.get("output_dir")
                if d:
                    p = Path(d)
                    if p.exists() or True:  # 없어도 일단 사용
                        return p
        except Exception:
            pass
        return OUTPUT_DIR

    def _save_output_config(self):
        """output_dir을 config.json에 저장."""
        try:
            import json
            cfg = {}
            if CONFIG_PATH.exists():
                try:
                    cfg = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
                except Exception:
                    pass
            cfg["output_dir"] = str(self._output_dir)
            CONFIG_PATH.write_text(json.dumps(cfg, ensure_ascii=False, indent=2),
                                   encoding="utf-8")
        except Exception as e:
            log.warning(f"config 저장 실패: {e}")

    def _pick_output_dir(self):
        """폴더 선택 다이얼로그."""
        d = filedialog.askdirectory(
            title="저장 폴더 선택",
            initialdir=str(self._output_dir))
        if d:
            self._output_dir = Path(d)
            self._output_dir.mkdir(parents=True, exist_ok=True)
            self._outdir_lbl.config(text=str(self._output_dir))
            self._save_output_config()

    def _open_output_dir(self):
        """저장 폴더를 탐색기에서 열기."""
        self._output_dir.mkdir(parents=True, exist_ok=True)
        os.startfile(str(self._output_dir))

    def _open_last_output(self):
        """최근 생성된 파일 열기."""
        if self._last_output_path and self._last_output_path.exists():
            os.startfile(str(self._last_output_path))
        else:
            # 폴더의 가장 최근 PDF 열기
            pdfs = sorted(self._output_dir.glob("*.pdf"),
                          key=lambda p: p.stat().st_mtime, reverse=True)
            if pdfs:
                os.startfile(str(pdfs[0]))
            else:
                messagebox.showinfo("파일 없음", "저장된 파일이 없습니다.")

    # ══════════════════════════════════════════════════════════════════
    # UI 업데이트 헬퍼
    # ══════════════════════════════════════════════════════════════════

    def _update_step_ui(self):
        for i, (dot, lbl) in enumerate(self._step_rows):
            if i < self._current_step:
                dot.config(text="✓", fg=C_PRI)
                lbl.config(fg=C_SUB, font=(F, 11))
            elif i == self._current_step:
                dot.config(text="●", fg=C_WARN)
                lbl.config(fg=C_TEXT, font=(F, 11, "bold"))
            else:
                dot.config(text="○", fg=C_BORDER)
                lbl.config(fg=C_TEXT, font=(F, 11))
        self._progress["value"] = self._current_step
        # 처음으로 돌아가면 블링크 중지 + 초기화
        if self._current_step == 0:
            self._stop_save_done_blink()
            self._save_done_lbl.config(text="", bg=C_SIDE)

    def _update_fname_preview(self, info: dict | None = None):
        cid = self._candidate_id.get().strip() or "?"
        nat = info.get("nationality", "") if info else ""
        gen = info.get("gender",      "") if info else ""
        yr  = info.get("birth_year",  "") if info else ""
        try:
            from .pdf_builder import build_filename
            self._fname_lbl.config(text=build_filename(cid, nat, gen, yr))
        except Exception:
            pass

    def _focus_check(self):
        try:
            sel_text = self._pii_text.get("sel.first", "sel.last")
        except tk.TclError:
            sel_text = ""
        if not sel_text.strip():
            messagebox.showinfo("집중 점검", "텍스트를 먼저 선택하세요.")
            return
        from .pii_engine import analyze_pii_focus, load_api_key
        api_key = load_api_key()
        if not api_key:
            messagebox.showwarning("API 키 없음", "Anthropic API 키가 필요합니다.")
            return
        result = analyze_pii_focus(sel_text, api_key)
        self._show_pii_result(result)

    def _on_pii_select(self, event=None):
        try:
            sel   = self._pii_text.get("sel.first", "sel.last")
            state = "normal" if sel.strip() else "disabled"
        except tk.TclError:
            state = "disabled"
        self._focus_btn.config(state=state)

    # ══════════════════════════════════════════════════════════════════
    # 알림 카드 (외부 호출)
    # ══════════════════════════════════════════════════════════════════

    def show_duplicate_alert(self, candidate_id: str, old_path, new_path):
        self._dup_body.config(
            text=f"이전: {old_path.name}\n{candidate_id}번 7일 내 재제출 감지")
        for w in [self._dup_title, self._dup_body, self._dup_btn]:
            w.pack(anchor="w", padx=8, pady=2)
        self._dup_card.pack(fill="x", padx=12, pady=4)

    def show_reapply_alert(self, candidate_id: str, old_path):
        self._reapp_body.config(text=f"이전 지원: {old_path.name}")
        for w in [self._reapp_title, self._reapp_body]:
            w.pack(anchor="w", padx=8, pady=2)
        btn_row = self._reapp_view_btn.master
        btn_row.pack(anchor="w", padx=8, pady=4)
        self._reapp_view_btn.config(
            command=lambda: os.startfile(str(old_path.parent)))
        self._reapp_del_btn.config(
            command=lambda: self._confirm_delete_old(old_path))
        self._reapp_keep_btn.config(
            command=lambda: self._reapp_card.pack_forget())
        self._reapp_card.pack(fill="x", padx=12, pady=4)

    def _confirm_delete_old(self, old_path):
        ok = messagebox.askyesno(
            "이전 파일 삭제",
            f"이전 파일을 originals\\ 로 이동 후 삭제합니다.\n\n{old_path}"
            "\n\n복구 불가합니다. 계속하시겠습니까?")
        if ok:
            from .security import secure_delete, encrypt_file
            import shutil
            dest = BASE_DIR / "originals" / old_path.name
            dest.parent.mkdir(exist_ok=True)
            shutil.move(str(old_path), str(dest))
            encrypt_file(dest)
            secure_delete(dest)
            self._reapp_card.pack_forget()
            messagebox.showinfo("삭제 완료", "이전 파일이 삭제되었습니다.")

    # ══════════════════════════════════════════════════════════════════
    # 실행
    # ══════════════════════════════════════════════════════════════════

    def run(self):
        self.root.mainloop()


# ── 진입점 ─────────────────────────────────────────────────────────────────
def _regex_pii(text: str) -> list[str]:
    """PDF 페이지 텍스트에서 정규식으로 PII 직접 추출 (Claude API 불필요).

    추출 대상:
      - 이메일 주소
      - 전화번호 (한국 010, 국제 +82, 일반 국제)
      - 주소 (한국식)
      - SNS/LinkedIn URL, 카카오 ID
      - 한국 학원/학교명
      - 영문 University/College + Korea 패턴
    """
    import re
    found: list[str] = []

    # ── 이메일 ──────────────────────────────────────────────────────
    found += re.findall(
        r'[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}', text)

    # ── 전화번호 — 한국 (010-xxxx-xxxx, +82-10-xxxx-xxxx 등) ───────
    found += re.findall(
        r'(?:\+82|0082)[\s.\-]?1[0-9][\s.\-]?\d{3,4}[\s.\-]?\d{4}', text)
    found += re.findall(
        r'\+\d{10,15}', text)           # +821029006080 같은 붙여쓰기
    found += re.findall(
        r'0(?:10|11|16|17|18|19)[\s.\-]?\d{3,4}[\s.\-]?\d{4}', text)
    # 국제 전화 (일반 형식)
    found += re.findall(
        r'\+\d{1,3}[\s.\-]\(?\d{1,4}\)?[\s.\-]\d{3,4}[\s.\-]\d{3,4}', text)

    # ── 한국 주소 ────────────────────────────────────────────────────
    found += re.findall(
        r'(?:서울|부산|대구|인천|광주|대전|울산|세종|경기|강원|충북|충남|전북|전남|경북|경남|제주)'
        r'[가-힣\s\d,.\-]{2,40}'
        r'(?:시|구|동|로|길|번길|번지|읍|면|리)(?:\s*\d+)?', text)

    # ── SNS / 링크드인 ────────────────────────────────────────────────
    found += re.findall(
        r'(?:instagram\.com|linkedin\.com/in|facebook\.com|twitter\.com|t\.me)'
        r'[/\s]?\S{2,50}', text, re.IGNORECASE)
    # 카카오: 한글(카카오/카톡) + 영문(kakao/kakaotalk) + ID 접미사 포함
    found += re.findall(
        r'(?:카카오|카톡|kakao(?:talk)?(?:\s*(?:id|아이디))?)[:\s]+\S{2,30}',
        text, re.IGNORECASE)

    # ── 한국 학원/학교명 (한글 접미사) ──────────────────────────────
    found += re.findall(
        r'[A-Za-z가-힣]{2,30}'
        r'(?:어학원|학원|학교|유치원|교육원|교습소|영어학원|어린이집)', text)

    # ── 영문 교육기관 + 한국 도시/국가 패턴 (줄바꿈 금지) ──────────
    # 예: "FiE Korea University | Jeonju | South Korea"
    # 예: "Avalon Langcon Education | Cheonan | South Korea"
    found += re.findall(
        r'[A-Z][A-Za-z \t&]{2,35}'
        r'(?:University|College|Institute|Academy|School|Education|Langcon|Learning)'
        r'(?:[ \t]*\|[ \t]*[A-Za-z][A-Za-z \t]{2,25}'
        r'(?:[ \t]*\|[ \t]*[A-Za-z][A-Za-z \t]{2,25})?)?',
        text)

    # ── 정리: 공백 트림, 짧은 것 제거, 중복 제거 ────────────────────
    cleaned = []
    seen: set[str] = set()
    for item in found:
        s = item.strip()
        if len(s) >= 5 and s not in seen:
            cleaned.append(s)
            seen.add(s)
    return cleaned


def _pii_search_variants(text: str) -> list[str]:
    """PII 검색용 변형 목록 생성.

    전화번호가 PDF에 "+82 10 2900 6080" 또는 "+821029006080" 등으로
    저장될 수 있으므로 여러 형식을 시도.
    파이프(|) 구분자 포함 시 각 파트도 검색.
    """
    import re
    variants: set[str] = {text}

    # 구두점 제거 버전
    stripped = re.sub(r'[\s\-\.\(\)]', '', text)
    if stripped and stripped != text and len(stripped) >= 5:
        variants.add(stripped)

    # 대시→공백 정규화 버전
    spaced = re.sub(r'[\-\.]', ' ', text).strip()
    if spaced != text and len(spaced) >= 5:
        variants.add(spaced)

    # 파이프 구분자: 각 파트 별도 검색 (학교명 | 도시 | 국가 형식)
    if '|' in text:
        for part in text.split('|'):
            p = part.strip()
            if len(p) >= 4:
                variants.add(p)

    return [v for v in variants if v and len(v.strip()) >= 3]


def _redact_pdf_page(page, pii_list: list[str]) -> None:
    """PDF 페이지에서 PII 텍스트를 완전 삭제.

    PyMuPDF add_redact_annot + apply_redactions():
      - 텍스트 내용을 PDF 콘텐츠 스트림에서 실제 제거 (단순 오버레이 아님)
      - fill=None = 투명 (흰색 박스 없음, 빈 공간으로 대체)
    """
    if not pii_list:
        return
    annotated = False
    for pii in pii_list:
        for variant in _pii_search_variants(pii):
            if len(variant.strip()) < 3:
                continue
            try:
                rects = page.search_for(variant, quads=False)
            except Exception:
                continue
            for r in rects:
                page.add_redact_annot(r, fill=None)
                annotated = True
    if annotated:
        page.apply_redactions()


# ── PII 섹션 헤더 패턴 ────────────────────────────────────────────────────────
# 삭제 대상 섹션: contact / references / personal details 등
_RE_PII_SECTION = re.compile(
    r'^\s*(?:'
    r'contact(?:\s+(?:information|details?|info|us|me|number|numbers?))?'
    r'|personal(?:\s+(?:details?|information|info|data|profile|statement))?'
    r'|(?:email|e[-\s]?mail|phone|mobile|cell|telephone|fax|address|location)'
    r'(?:\s+(?:info(?:rmation)?|details?|number|address))?'
    r'|references?(?:\s*[&/]\s*referees?)?'
    r'|referees?'
    r'|emergency\s+contacts?'
    r'|연락처|개인\s*정보|참고인|추천인|레퍼런스'
    r')\s*:?\s*$',
    re.IGNORECASE,
)

# 콘텐츠 섹션 (이 섹션이 시작되면 삭제 종료)
_RE_CONTENT_SECTION = re.compile(
    r'^\s*(?:'
    r'education(?:al)?(?:\s+background)?|academic(?:\s+background)?'
    r'|qualifications?|academic\s+qualifications?'
    r'|(?:work\s+)?experience|employment(?:\s+history)?|career(?:\s+history)?'
    r'|professional\s+(?:experience|background|profile|development)'
    r'|skills?(?:\s+[&/]\s*(?:competencies?|abilities?))?'
    r'|competencies?|abilities?|core\s+competencies?'
    r'|summary|profile|career\s+(?:summary|objective)'
    r'|about(?:\s+me)?|objective|career\s+objective'
    r'|certifications?|credentials?|licenses?|accreditations?'
    r'|awards?|achievements?|accomplishments?'
    r'|languages?(?:\s+skills?)?|hobbies|interests?|extracurricular'
    r'|학력|경력|기술|자격증?|자기소개|소개|관심사'
    r')\s*:?\s*$',
    re.IGNORECASE,
)


def _redact_pii_sections_pdf(page, is_first_page: bool = False) -> None:
    """PDF 페이지에서 PII 섹션 전체를 완전 삭제.

    A) contact / references / personal details 섹션 헤더 탐지
       → 다음 콘텐츠 섹션 직전까지 삭제
       → 사이드바 레이아웃 자동 감지 (x < 40%): 사이드바 열만 삭제,
         메인 열 콘텐츠(Experience 등) 보호
    B) 첫 페이지 최상단 이름+연락처 헤더
       → 첫 콘텐츠 섹션 직전까지 전체 너비 삭제 (최대 50%)
    """
    import fitz

    raw_blocks = page.get_text("blocks")
    if not raw_blocks:
        return
    blocks = sorted(raw_blocks, key=lambda b: b[1])  # y0 오름차순
    page_h = page.rect.height
    page_w = page.rect.width

    # 사이드바 임계값: 블록 x1 < 40% = 사이드바 열
    SIDEBAR_X = page_w * 0.40

    # ── 블록 분류: (x0, y0, x1, y1, cls, in_sidebar) ─────────────────────
    classified: list[tuple[float, float, float, float, str, bool]] = []
    for b in blocks:
        x0, y0, x1, y1, text, *_ = b
        if not text.strip():
            continue
        first_line = text.strip().split('\n')[0].strip()
        in_sidebar = (x1 < SIDEBAR_X)
        if _RE_PII_SECTION.match(first_line):
            classified.append((x0, y0, x1, y1, 'pii', in_sidebar))
        elif _RE_CONTENT_SECTION.match(first_line):
            classified.append((x0, y0, x1, y1, 'content', in_sidebar))
        else:
            classified.append((x0, y0, x1, y1, 'body', in_sidebar))

    regions: list[tuple[float, float, float, float]] = []  # (rx0, ry0, rx1, ry1)

    # ── A: PII 섹션 → 같은 열의 다음 콘텐츠 섹션까지 ────────────────────
    n = len(classified)
    for i, (x0, y0, x1, y1, cls, in_sidebar) in enumerate(classified):
        if cls != 'pii':
            continue

        r_top = max(0.0, y0 - 8)
        r_bot = min(page_h, y1 + page_h * 0.6)  # 기본 fallback

        if in_sidebar:
            # 사이드바 PII: 사이드바 열(x=0~SIDEBAR_X)만 삭제
            rx0, rx1 = 0.0, SIDEBAR_X
            for j in range(i + 1, n):
                ny0, ny1, ncls, nside = classified[j][1], classified[j][3], classified[j][4], classified[j][5]
                if nside and ncls == 'content':
                    r_bot = ny0 - 4
                    break
                elif nside and ncls == 'pii':
                    r_bot = ny1  # 연속 PII 섹션 → 아래로 확장
        else:
            # 메인 열 PII (Reference 등): 전체 너비 삭제
            rx0, rx1 = 0.0, page_w
            for j in range(i + 1, n):
                ny0, ny1, ncls, nside = classified[j][1], classified[j][3], classified[j][4], classified[j][5]
                if (not nside) and ncls == 'content':
                    r_bot = ny0 - 4
                    break
                elif (not nside) and ncls == 'pii':
                    r_bot = ny1

        regions.append((rx0, r_top, rx1, r_bot))

    # ── B: 첫 페이지 최상단 이름+연락처 헤더 ────────────────────────────
    if is_first_page:
        first_content_y: float | None = None
        for x0, y0, x1, y1, cls, in_sidebar in classified:
            if cls == 'content':
                first_content_y = y0
                break
        if first_content_y is not None:
            cap = min(first_content_y - 4, page_h * 0.50)
            if cap > 10:
                regions.append((0.0, 0.0, page_w, cap))

    # ── 적용 (fill=None = 스트림 완전 삭제) ─────────────────────────────
    for rx0, r_top, rx1, r_bot in regions:
        if r_bot > r_top:
            rect = fitz.Rect(rx0, r_top, rx1, r_bot)
            page.add_redact_annot(rect, fill=None)
    if regions:
        page.apply_redactions()


def _overlay_id_on_resume_page(
    page, candidate_id: str, photo_bytes: bytes | None = None
) -> None:
    """이력서 첫 페이지 상단에 강사 번호(좌) + 증명사진(우) 삽입.

    _redact_pii_sections_pdf(is_first_page=True)가 이미 상단 이름 헤더를
    스트림에서 완전 삭제한 뒤, 그 자리에 번호와 사진을 삽입.

    배치:
      좌상단: 번호 (72pt Helvetica-Bold, 검정)
      우상단: 증명사진 (130×165pt)
    """
    import fitz
    page_w = page.rect.width
    # 강사 번호 — 좌상단 72pt
    try:
        page.insert_text(
            (20, 85),
            candidate_id,
            fontsize=72,
            fontname="helv",
            color=(0, 0, 0),
            render_mode=0,
        )
    except Exception as e:
        log.warning(f"ID 오버레이 실패: {e}")
    # 증명사진 — 우상단
    if photo_bytes:
        try:
            photo_rect = fitz.Rect(page_w - 150, 10, page_w - 20, 175)
            page.insert_image(photo_rect, stream=photo_bytes)
        except Exception as e:
            log.warning(f"사진 오버레이 실패: {e}")


def _clear_pii_sections_docx(doc) -> None:
    """DOCX 문서에서 PII 섹션 전체를 run-level로 삭제.

    A) contact / references / personal details 섹션 → 다음 콘텐츠 섹션까지
    B) 첫 콘텐츠 섹션 이전의 최상단 이름+연락처 헤더 전체
    """
    paras = list(doc.paragraphs)
    n = len(paras)
    if n == 0:
        return

    # 단락 분류
    flags: list[str] = []
    for para in paras:
        t = para.text.strip()
        if not t:
            flags.append('empty')
        elif _RE_PII_SECTION.match(t):
            flags.append('pii')
        elif _RE_CONTENT_SECTION.match(t):
            flags.append('content')
        else:
            flags.append('body')

    # 삭제할 인덱스 수집
    kill: set[int] = set()

    # ── A: PII 섹션 → 다음 콘텐츠 섹션까지 ─────────────────────────────
    in_pii = False
    for i, f in enumerate(flags):
        if f == 'pii':
            in_pii = True
        elif f == 'content':
            in_pii = False
        if in_pii:
            kill.add(i)

    # ── B: 첫 콘텐츠 섹션 이전 최상단 헤더 전체 ─────────────────────────
    first_content = next((i for i, f in enumerate(flags) if f == 'content'), None)
    if first_content is not None and first_content > 0:
        for i in range(first_content):
            kill.add(i)

    # ── run 텍스트 제거 ───────────────────────────────────────────────────
    for i in kill:
        for run in paras[i].runs:
            run.text = ''

    # ── 테이블도 동일 처리: 테이블 전체가 최상단 헤더 영역에 있으면 제거 ──
    # (간단 처리: 첫 콘텐츠 단락 인덱스 기준으로 테이블은 별도 판별 불가
    #  → 테이블 내 PII 섹션 헤더 셀 탐지로 처리)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if _RE_PII_SECTION.match(cell_text):
                    # 이 셀이 속한 테이블의 모든 셀 제거
                    for r2 in tbl.rows:
                        for c2 in r2.cells:
                            for para in c2.paragraphs:
                                for run in para.runs:
                                    run.text = ''
                    break


def _build_id_only_cover_page(
    buf: "io.BytesIO",
    candidate_id: str,
    page_w: float,
    page_h: float,
) -> None:
    """
    사진이 이력서 본문에 이미 있을 때 사용하는 커버 페이지.

    레이아웃:
      - 상단 가로선 아래 강사 ID를 좌상단 + 우상단 양쪽에 크게 표시
      - 중앙 BRIDGE 로고 텍스트 + 안내 문구
      - 하단 회색 가로선 + 처리일자

    Args:
        buf:          출력 BytesIO
        candidate_id: 강사 번호 (예: "3060")
        page_w, page_h: reportlab A4 기준 (pt)
    """
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import datetime

    c = rl_canvas.Canvas(buf, pagesize=(page_w, page_h))

    # ── 배경 ──────────────────────────────────────────────────────────
    c.setFillColor(colors.HexColor("#FFFFFF"))
    c.rect(0, 0, page_w, page_h, fill=1, stroke=0)

    # ── 상단 진한 헤더 바 ─────────────────────────────────────────────
    bar_h = 60
    c.setFillColor(colors.HexColor("#1D1D2E"))
    c.rect(0, page_h - bar_h, page_w, bar_h, fill=1, stroke=0)

    # BRIDGE 로고 (헤더 바 안, 가운데)
    c.setFillColor(colors.HexColor("#1D9E75"))
    c.setFont("Helvetica-Bold", 18)
    c.drawCentredString(page_w / 2, page_h - bar_h + 20, "BRIDGE")

    # ── 강사 ID — 좌상단 ──────────────────────────────────────────────
    # 헤더 바 바로 아래부터 시작
    top_y = page_h - bar_h - 70

    # 좌상단 ID 박스
    c.setFillColor(colors.HexColor("#F0FBF6"))
    c.roundRect(18*mm, top_y, 60*mm, 42, radius=6, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#1D9E75"))
    c.setFont("Helvetica-Bold", 9)
    c.drawString(20*mm, top_y + 28, "CANDIDATE ID")
    c.setFillColor(colors.HexColor("#1D1D2E"))
    c.setFont("Helvetica-Bold", 26)
    c.drawString(20*mm, top_y + 8, str(candidate_id))

    # 우상단 ID 박스 (미러)
    rx = page_w - 18*mm - 60*mm
    c.setFillColor(colors.HexColor("#F0FBF6"))
    c.roundRect(rx, top_y, 60*mm, 42, radius=6, fill=1, stroke=0)
    c.setFillColor(colors.HexColor("#1D9E75"))
    c.setFont("Helvetica-Bold", 9)
    c.drawString(rx + 4, top_y + 28, "CANDIDATE ID")
    c.setFillColor(colors.HexColor("#1D1D2E"))
    c.setFont("Helvetica-Bold", 26)
    c.drawString(rx + 4, top_y + 8, str(candidate_id))

    # ── 중앙 안내 텍스트 ──────────────────────────────────────────────
    mid_y = page_h / 2
    c.setFillColor(colors.HexColor("#5F6368"))
    c.setFont("Helvetica", 11)
    c.drawCentredString(page_w / 2, mid_y + 20,
                        "This document has been processed by BRIDGE.")
    c.setFont("Helvetica", 10)
    c.drawCentredString(page_w / 2, mid_y,
                        "Personal Identifiable Information (PII) has been removed.")
    c.setFont("Helvetica", 10)
    c.drawCentredString(page_w / 2, mid_y - 20,
                        "Photo is retained in the original resume pages below.")

    # ── 하단 처리일자 ─────────────────────────────────────────────────
    c.setStrokeColor(colors.HexColor("#E8EAED"))
    c.line(18*mm, 50, page_w - 18*mm, 50)
    today = datetime.date.today().strftime("%Y-%m-%d")
    c.setFillColor(colors.HexColor("#5F6368"))
    c.setFont("Helvetica", 9)
    c.drawString(18*mm, 34, f"Processed: {today}")
    c.drawRightString(page_w - 18*mm, 34, "BRIDGE ESL Recruitment")

    c.save()


def main():
    app = BridgeConverterApp()
    app.run()


if __name__ == "__main__":
    main()
