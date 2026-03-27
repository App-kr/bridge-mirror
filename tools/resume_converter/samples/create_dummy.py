"""
create_dummy.py — 테스트용 더미 파일 생성
실제 개인정보 미사용
"""
from pathlib import Path

BASE = Path(__file__).parent

# 더미 이력서 텍스트
RESUME_TEXT = """
John Doe
Email: john.doe.test@example.com
Phone: 010-0000-0000
Address: Seoul, South Korea

EDUCATION
- B.A. English Education, Test University (2015-2019)

EXPERIENCE
- English Teacher, ABC어학원 (2019-2021)
- English Teacher, XYZ English School (2021-2023)

SKILLS
- Native English speaker
- TEFL/TESOL certified

REFERENCES
Available upon request
"""

COVER_TEXT = """
Dear Hiring Manager,

I am writing to apply for the ESL teacher position at your school.
My name is John Doe and I can be reached at john.doe.test@example.com or 010-0000-0000.

I have 4 years of experience teaching English in Korea.

Best regards,
John Doe
"""

REC_TEXT = """
To Whom It May Concern,

I am pleased to recommend John Doe for your ESL program.
During his time at ABC어학원, he demonstrated excellent teaching skills.

Sincerely,
Jane Smith, Director
ABC어학원
Tel: 02-0000-0000
Email: director@abc-test.com
"""

def create_samples():
    # Word .docx 생성
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Test Resume', 0)
        for line in RESUME_TEXT.strip().split('\n'):
            doc.add_paragraph(line)
        doc.save(str(BASE / '3000_resume.docx'))
        print("OK: 3000_resume.docx")

        doc2 = Document()
        doc2.add_heading('Cover Letter', 0)
        for line in COVER_TEXT.strip().split('\n'):
            doc2.add_paragraph(line)
        doc2.save(str(BASE / '3000_cover_letter.docx'))
        print("OK: 3000_cover_letter.docx")

        doc3 = Document()
        doc3.add_heading('Recommendation Letter', 0)
        for line in REC_TEXT.strip().split('\n'):
            doc3.add_paragraph(line)
        doc3.save(str(BASE / '3000_recommendation.docx'))
        print("OK: 3000_recommendation.docx")
    except ImportError:
        print("python-docx 없음 — txt 파일로 대체")
        (BASE / '3000_resume.txt').write_text(RESUME_TEXT, encoding='utf-8')
        (BASE / '3000_cover.txt').write_text(COVER_TEXT, encoding='utf-8')
        (BASE / '3000_rec.txt').write_text(REC_TEXT, encoding='utf-8')

    # 더미 사진 (단색 JPEG)
    try:
        from PIL import Image, ImageDraw
        img = Image.new('RGB', (400, 500), color=(200, 210, 220))
        draw = ImageDraw.Draw(img)
        # 사람 얼굴 모양 원 그리기 (OpenCV 얼굴 인식용)
        draw.ellipse([150, 80, 250, 180], fill=(230, 200, 180))  # 얼굴
        draw.ellipse([170, 100, 200, 130], fill=(50, 50, 50))    # 눈1
        draw.ellipse([210, 100, 240, 130], fill=(50, 50, 50))    # 눈2
        draw.arc([180, 140, 230, 165], 0, 180, fill=(50, 50, 50), width=2)  # 입
        img.save(str(BASE / '3000_photo.jpg'), format='JPEG')
        print("OK: 3000_photo.jpg")
    except ImportError:
        print("Pillow 없음 — 사진 생성 건너뜀")

if __name__ == '__main__':
    create_samples()
    print("\n더미 파일 생성 완료:", list(BASE.glob('3000_*')))
